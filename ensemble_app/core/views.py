from django.conf import settings
from django.core.mail import EmailMessage, send_mail
from django.template.loader import render_to_string
from django.http import HttpResponse
from django.utils import timezone
import uuid
import json
import json
import uuid
from datetime import datetime, date, time, timedelta
from django.shortcuts import get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.timezone import make_aware
from django.core.mail import send_mail
from django.core.cache import cache
from django.views import View
from .models import VenueBooking, SessionDate, ProjectPlan, Venue
from .forms import VenueBookingForm
from .mixins import RolePermissionRequiredMixin
from itertools import groupby
from operator import attrgetter
from datetime import datetime, date, time, timedelta
from django.db.models import Q
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import TemplateView
from django.core.cache import cache
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from .models import (
    Learner, Group, ProjectPlan, Fingerprint, SessionDate,
    VenueBooking, WeeklySchedule, LearnerQualification, SLA
)
from django.http import FileResponse
from .models import AdobeForm
from .forms import AdobeFormUploadForm, AdobeFormFilledSubmissionForm
from django.views.generic import ListView, CreateView, FormView
from django.urls import reverse_lazy
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from PyPDF2 import PdfReader
from django.core.paginator import Paginator
from django.views.generic import ListView, DetailView, View, UpdateView
from django.views.generic import ListView, DetailView, View, TemplateView
from django.shortcuts import render, redirect, get_object_or_404
from dal import autocomplete
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.shortcuts import get_object_or_404, redirect
from django.db.models import Q
from django.views.generic import View
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Group, SETA
from .forms import GroupCreateForm, SETAForm
from django import forms
from django.http import JsonResponse
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.contrib import messages
from django.core.cache import cache
from django.utils import timezone
from .models import Role
from django.urls import reverse_lazy
from .forms import GroupCreateForm, SETAForm
from django.http import JsonResponse
from .models import (
    SLA, 
    SLA_Qualifications, 
    BillingHistory, 
    BillingHistoryItem, 
    LearnerQualification, 
    Learner, 
    VATRate, 
    Fingerprint, 
    WeeklySchedule,
    LearnerRole,  #
    Role, 
)
from datetime import datetime, date, time, timedelta
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.forms import formset_factory, modelformset_factory
from .models import SLA, SLA_Qualifications, BillingHistory, BillingHistoryItem, LearnerQualification, Learner, VATRate
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from .forms import SLAForm, SLAQualificationFormSet, TrancheBillingForm, QualificationLearnersForm
import calendar
import csv
import re
from num2words import num2words
import pandas as pd
from django.core.files.storage import FileSystemStorage
import logging
from django.http import JsonResponse
from .models import VenueBooking, SessionDate
from django.views.generic import TemplateView
from django.urls import reverse_lazy
from django.views.generic import (
    ListView, CreateView, UpdateView, DeleteView, DetailView
)
from .models import Service, Module, Service_Module, ModulePOETemplate, ModulePOEAnnexture
from .forms import ServiceForm, ModuleForm, ServiceModuleForm,ModulePOEAnnextureForm 
import hashlib
from .mixins import RolePermissionRequiredMixin

# Set up logging
logger = logging.getLogger(__name__)
from django.shortcuts import redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Learner  # Ensure your Learner model is imported

VENUE_COLOR_MAP = {
    "everest": "#e75480",           # Pink
    "buna": "#8B5C2A",              # Brown
    "kilimanjaro": "#8e44ad",       # Purple
    "offsite": "#3A7AFE",           # Blue
    "back office": "#FFA500",       # Orange
    "management office": "#888888", # Gray
    "hybrid": "#2acaea",            # Purple
}

def get_venue_color(venue_name):
    if not venue_name:
        return "#34C759"
    key = venue_name.strip().lower()
    return VENUE_COLOR_MAP.get(key, "#34C759")


def auto_book_virtual_sessions():
    from .models import SessionDate, VenueBooking, Venue
    from django.utils.timezone import make_aware
    from datetime import time, datetime

    virtual_sessions = SessionDate.objects.filter(
        preferred_training_methodology__icontains="virtual session"
    )
    virtual_venue = Venue.objects.filter(name__icontains="virtual session").first()
    if not virtual_venue:
        return

    for session in virtual_sessions:
        # Check if this session has any cancelled bookings
        has_cancelled_booking = VenueBooking.objects.filter(
            session_date=session,
            status='cancelled'
        ).exists()
        
        # Skip auto-booking if there are cancelled bookings for this session
        if has_cancelled_booking:
            continue
            
        start_dt = make_aware(datetime.combine(session.start_date, time(8, 0)))
        end_dt = make_aware(datetime.combine(session.end_date, time(17, 0)))
        
        # Check for any booking with same venue and time (regardless of session)
        exists = VenueBooking.objects.filter(
            venue=virtual_venue,
            start_datetime=start_dt,
            end_datetime=end_dt,
            status__in=['booked', 'rescheduled']
        ).exists()
        
        if not exists:
            # Create the virtual booking only if no cancelled bookings exist
            VenueBooking.objects.create(
                session_date=session,
                venue=virtual_venue,
                start_datetime=start_dt,
                end_datetime=end_dt,
                booking_purpose=f"Auto-booked Virtual Session - {session.project_plan.group.name if session.project_plan and session.project_plan.group else 'Unknown Group'}",
                status='booked',
                num_learners=0,  # Set appropriate default
                num_learners_lunch=0
            )

@login_required
def home_redirect_view(request):
    # Staff users (admin/finance) are redirected to the finance dashboard.
    if request.user.is_staff:
        return redirect('finance_dashboard')
    try:
        # Assumes you have a OneToOne relation or property 'learner_profile' on your user.
        learner = request.user.learner_profile
        return redirect('learner_home')
    except Learner.DoesNotExist:
        return redirect('finance_dashboard')

# --- Service CRUD ---
class ServiceListView(RolePermissionRequiredMixin, ListView):
    model = Service
    template_name = "core/service_list.html"
    context_object_name = "services"
    paginate_by = 10 
    def get_queryset(self):
        queryset = super().get_queryset()
        search = self.request.GET.get('search', '')
        if search:
            queryset = queryset.filter(
                Q(name__icontains=search) |
                Q(gl_code__icontains=search) |
                Q(saqa_id__icontains=search)
            )
        return queryset

class ServiceCreateView(RolePermissionRequiredMixin, CreateView):
    model = Service
    form_class = ServiceForm
    template_name = "core/service_form.html"
    success_url = reverse_lazy("service-list")

class ServiceUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = Service
    form_class = ServiceForm
    template_name = "core/service_form.html"
    success_url = reverse_lazy("service-list")

class ServiceDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = Service
    template_name = "core/service_confirm_delete.html"
    success_url = reverse_lazy("service-list")

# --- Module CRUD ---
class ModuleListView(RolePermissionRequiredMixin, ListView):
    model = Module
    template_name = "core/module_list.html"
    context_object_name = "modules"
    def get_queryset(self):
        queryset = super().get_queryset()
        search = self.request.GET.get('search', '')
        service = self.request.GET.get('service', '')
        if search:
            queryset = queryset.filter(Q(name__icontains=search) | Q(code__icontains=search))
        if service:
            queryset = queryset.filter(services__id=service)
        return queryset.distinct()

class ModuleCreateView(RolePermissionRequiredMixin, CreateView):
    model = Module
    form_class = ModuleForm
    template_name = "core/module_form.html"
    success_url = reverse_lazy("module-list")

class ModuleUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = Module
    form_class = ModuleForm
    template_name = "core/module_form.html"
    success_url = reverse_lazy("module-list")

class ModuleDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = Module
    template_name = "core/module_confirm_delete.html"
    success_url = reverse_lazy("module-list")

# --- Service_Module CRUD ---
class ServiceModuleListView(RolePermissionRequiredMixin, ListView):
    model = Service_Module
    template_name = "core/service_module_list.html"
    context_object_name = "service_modules"

class ServiceModuleCreateView(RolePermissionRequiredMixin, CreateView):
    model = Service_Module
    form_class = ServiceModuleForm
    template_name = "core/service_module_form.html"
    success_url = reverse_lazy("service-module-list")

class ServiceModuleUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = Service_Module
    form_class = ServiceModuleForm
    template_name = "core/service_module_form.html"
    success_url = reverse_lazy("service-module-list")

class ServiceModuleDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = Service_Module
    template_name = "core/service_module_confirm_delete.html"
    success_url = reverse_lazy("service-module-list")


class VenueBookingCalendarView(RolePermissionRequiredMixin, TemplateView):
    template_name = "core/venuebooking_calendar.html"
    def get_context_data(self, **kwargs):
        from .models import Venue
        # Auto-book virtual sessions before rendering the calendar
        auto_book_virtual_sessions()
        context = super().get_context_data(**kwargs)
        context['venues'] = Venue.objects.all().order_by('name')
        return context

from collections import defaultdict
from django.http import JsonResponse
from datetime import datetime, time
from django.http import JsonResponse
from django.utils.timezone import make_aware
from datetime import datetime, time
from django.http import JsonResponse

from django.http import JsonResponse
from django.utils.timezone import make_aware
from datetime import datetime, time
from django.http import JsonResponse
from django.utils.timezone import make_aware
from datetime import datetime, time
from django.core.cache import cache

from django.utils.timezone import make_aware
from datetime import datetime, time
from django.core.cache import cache
from django.http import JsonResponse

from django.utils.timezone import make_aware
from datetime import datetime, time
from django.core.cache import cache
from django.http import JsonResponse

def venuebooking_events_api(request):
    from .models import Venue, VenueBooking, SessionDate

    status_filter = request.GET.get('status')
    venue_filter = request.GET.get('venue')
    facilitator_filter = request.GET.get('facilitator')
    show_empty = request.GET.get('show_empty')
    date_filter = request.GET.get('date')

    bookings = VenueBooking.objects.select_related(
        'venue', 'session_date__project_plan__group', 'session_date__project_plan__module', 'facilitator', 'user'
    )
    
    # **KEY UPDATE**: Exclude cancelled bookings by default, show only when specifically requested
    if status_filter:
        bookings = bookings.filter(status=status_filter)
    else:
        # Default: exclude cancelled bookings
        bookings = bookings.exclude(status='cancelled')
    
    if venue_filter:
        if venue_filter.lower() == "virtual session":
            bookings = bookings.filter(venue__name__icontains="virtual session")
        else:
            bookings = bookings.filter(venue__name=venue_filter)
    if facilitator_filter:
        bookings = bookings.filter(facilitator__learner__FirstName__icontains=facilitator_filter)
    if date_filter:
        bookings = bookings.filter(start_datetime__date=date_filter)

    events = []
    booked_sessions = set()

    def get_venue_color(venue_name):
        VENUE_COLOR_MAP = {
            "everest": "#e75480",
            "buna": "#8B5C2A",
            "kilimanjaro": "#8e44ad",
            "offsite": "#3A7AFE",
            "back office": "#FFA500",
            "management office": "#888888",
            "hybrid": "#2acaea",
        }
        if not venue_name:
            return "#34C759"
        key = venue_name.strip().lower()
        return VENUE_COLOR_MAP.get(key, "#34C759")

    # --- Group combined bookings ---
    combined_map = {}
    for booking in bookings:
        if booking.combined_booking_id:
            combined_map.setdefault(booking.combined_booking_id, []).append(booking)
        else:
            combined_map.setdefault(f"single-{booking.id}", [booking])

    for group_id, group_bookings in combined_map.items():
        if len(group_bookings) == 1:
            booking = group_bookings[0]
            if booking.session_date_id:
                booked_sessions.add(booking.session_date_id)
            session = booking.session_date
            module = session.project_plan.module if session and session.project_plan else None
            group = session.project_plan.group if session and session.project_plan else None
            facilitator = ""
            if getattr(booking, "facilitator", None) and getattr(booking.facilitator, "learner", None):
                facilitator = f"{booking.facilitator.learner.FirstName} {booking.facilitator.learner.Surname}"

            # Booked by (user who made the booking)
            booked_by = ""
            if hasattr(booking, "user") and booking.user:
                booked_by = f"{booking.user.first_name} {booking.user.last_name}".strip() or booking.user.username
            num_learners = booking.num_learners if booking.num_learners is not None else ""

            if "virtual session" in booking.venue.name.lower():
                venue_color = "#34C759"
            else:
                venue_color = get_venue_color(booking.venue.name)
            color = venue_color
            if booking.status == 'rescheduled':
                color = "#FFA500"
            elif booking.status == 'cancelled':
                color = "#3AB0FF"

            # Check if this session has a change flag in cache
            session_changed = False
            if session:
                session_changed = cache.get(f'sessiondate_changed_{session.id}', False)
                if session_changed:
                    cache.delete(f'sessiondate_changed_{session.id}')

            events.append({
                "title": f"{str(group) if group else ''} - {module.name if module else ''}",
                "start": booking.start_datetime.isoformat(),
                "end": booking.end_datetime.isoformat(),
                "color": color,
                "status": booking.status,
                "session_changed": session_changed,
                "booked_by": booked_by,
                "num_learners": num_learners,
                "extendedProps": {
                    "venue": booking.venue.name,
                    "venue_color": venue_color,
                    "module": module.name if module else "",
                    "group": str(group) if group else "",
                    "facilitator": facilitator,
                    "booking_id": booking.id,
                    "status": booking.status,
                    "project_plan_id": session.project_plan.id if session and session.project_plan else None,
                    "combined_booking_id": booking.combined_booking_id,
                    "session_changed": session_changed,
                    "booked_by": booked_by,
                    "num_learners": num_learners,
                }
            })
        else:
            session_names = []
            module_names = []
            group_names = []
            facilitator_names = set()
            venue = group_bookings[0].venue
            status = group_bookings[0].status
            start = group_bookings[0].start_datetime
            end = group_bookings[0].end_datetime
            session_changed = False
            for b in group_bookings:
                s = b.session_date
                if s:
                    if s.project_plan and s.project_plan.group:
                        group_names.append(str(s.project_plan.group))
                    if s.project_plan and s.project_plan.module:
                        module_names.append(s.project_plan.module.name)
                    if cache.get(f'sessiondate_changed_{s.id}', False):
                        session_changed = True
                        cache.delete(f'sessiondate_changed_{s.id}')
                if b.facilitator and b.facilitator.learner:
                    facilitator_names.add(f"{b.facilitator.learner.FirstName} {b.facilitator.learner.Surname}")
                if b.session_date_id:
                    booked_sessions.add(b.session_date_id)
            venue_color = "#34C759" if "virtual session" in venue.name.lower() else get_venue_color(venue.name)
            color = venue_color
            if status == 'rescheduled':
                color = "#FFA500"
            elif status == 'cancelled':
                color = "#3AB0FF"
            events.append({
                "title": f"Multiple: {', '.join(group_names)} - {', '.join(module_names)}",
                "start": start.isoformat(),
                "end": end.isoformat(),
                "color": color,
                "status": status,
                "session_changed": session_changed,
                "booked_by": "",
                "num_learners": "",
                "extendedProps": {
                    "venue": venue.name,
                    "venue_color": venue_color,
                    "module": ", ".join(module_names),
                    "group": ", ".join(group_names),
                    "facilitator": ", ".join(facilitator_names),
                    "booking_id": group_bookings[0].id,
                    "status": status,
                    "project_plan_id": None,
                    "combined_booking_id": group_bookings[0].combined_booking_id,
                    "is_combined": True,
                    "booking_ids": [b.id for b in group_bookings],
                    "session_changed": session_changed,
                    "booked_by": "",
                    "num_learners": "",
                }
            })

    # --- AUTO-BOOK VIRTUAL SESSIONS ---
    virtual_sessions = SessionDate.objects.filter(
        preferred_training_methodology__icontains="virtual session"
    ).exclude(id__in=booked_sessions)

    for session in virtual_sessions:
        start_dt = make_aware(datetime.combine(session.start_date, time(8, 0)))
        end_dt = make_aware(datetime.combine(session.end_date, time(17, 0)))
        base_name = "Virtual Session"
        virtual_venues = list(Venue.objects.filter(name__istartswith=base_name).order_by('name'))
        venue_to_use = None
        for v in virtual_venues:
            conflict = VenueBooking.objects.filter(
                venue=v,
                start_datetime=start_dt,
                end_datetime=end_dt,
                status__in=['booked', 'rescheduled']
            ).exists()
            if not conflict:
                venue_to_use = v
                break
        if not venue_to_use:
            existing_names = [v.name for v in virtual_venues]
            idx = 2
            new_name = f"{base_name} {idx}"
            while new_name in existing_names:
                idx += 1
                new_name = f"{base_name} {idx}"
            venue_to_use = Venue.objects.create(name=new_name)
        if not VenueBooking.objects.filter(
            venue=venue_to_use,
            start_datetime=start_dt,
            end_datetime=end_dt,
            status__in=['booked', 'rescheduled']
        ).exists():
            VenueBooking.objects.create(
                venue=venue_to_use,
                session_date=session,
                start_datetime=start_dt,
                end_datetime=end_dt,
                booking_purpose="Virtual Session",
                facilitator=getattr(session, "facilitator", None),
                status="booked"
            )

    # --- Unbooked slots ---
    session_dates = SessionDate.objects.select_related('project_plan__group', 'project_plan__module')
    show_unbooked = (not status_filter or status_filter == "unbooked" or show_empty == "1")

    if show_unbooked:
        for session in session_dates:
            if session.id not in booked_sessions:
                module = session.project_plan.module if session.project_plan else None
                group = session.project_plan.group if session.project_plan else None
                session_changed = cache.get(f'sessiondate_changed_{session.id}', False)
                if session_changed:
                    cache.delete(f'sessiondate_changed_{session.id}')
                if session.preferred_training_methodology and "virtual session" in session.preferred_training_methodology.lower():
                    venue_name = "Virtual Session"
                    venue_color = "#34C759"
                    status = "booked"
                    start_dt = datetime.combine(session.start_date, time(8, 0))
                    end_dt = datetime.combine(session.end_date, time(17, 0))
                    if date_filter and str(session.start_date) != date_filter:
                        continue
                    events.append({
                        "title": f"{str(group) if group else ''} - {module.name if module else ''}",
                        "start": start_dt.isoformat(),
                        "end": end_dt.isoformat(),
                        "color": venue_color,
                        "status": status,
                        "session_changed": session_changed,
                        "booked_by": "",
                        "num_learners": "",
                        "extendedProps": {
                            "venue": venue_name,
                            "venue_color": venue_color,
                            "module": module.name if module else "",
                            "group": str(group) if group else "",
                            "facilitator": "",
                            "booking_id": "",
                            "status": status,
                            "session_id": session.id,
                            "project_plan_id": session.project_plan.id if session.project_plan else None,
                            "session_changed": session_changed,
                            "booked_by": "",
                            "num_learners": "",
                        }
                    })
                else:
                    if venue_filter:
                        if venue_filter.lower() == "virtual session":
                            continue
                        venue_obj = Venue.objects.filter(name=venue_filter).first()
                        if not venue_obj:
                            continue
                        if session.preferred_training_methodology and "virtual session" in session.preferred_training_methodology.lower():
                            continue
                        venue_name = venue_obj.name
                        venue_color = get_venue_color(venue_obj.name)
                    else:
                        venue_name = "Not booked"
                        venue_color = "#d32f2f"
                    status = "unbooked"
                    if date_filter and str(session.start_date) != date_filter:
                        continue
                    events.append({
                        "title": f"{str(group) if group else ''} - {module.name if module else ''}",
                        "start": session.start_date.isoformat(),
                        "end": session.end_date.isoformat(),
                        "color": venue_color,
                        "status": status,
                        "session_changed": session_changed,
                        "booked_by": "",
                        "num_learners": "",
                        "extendedProps": {
                            "venue": venue_name,
                            "venue_color": venue_color,
                            "module": module.name if module else "",
                            "group": str(group) if group else "",
                            "facilitator": "",
                            "booking_id": "",
                            "status": status,
                            "session_id": session.id,
                            "project_plan_id": session.project_plan.id if session.project_plan else None,
                            "session_changed": session_changed,
                            "booked_by": "",
                            "num_learners": "",
                        }
                    })

    if show_empty == "1":
        events = [e for e in events if e['status'] == 'unbooked']

    return JsonResponse(events, safe=False)

from django.views.decorators.http import require_GET
@require_GET
def venue_availability_api(request):
    from .models import VenueBooking, Venue
    import json
    venue_id = request.GET.get('venue_id')
    date_str = request.GET.get('date')
    virtual_group = request.GET.get('virtual_group')
    if not date_str:
        return JsonResponse({'available': False, 'reason': 'Missing venue or date'}, status=400)
    if virtual_group == "1":
        venues = Venue.objects.filter(name__icontains="virtual session")
        bookings = VenueBooking.objects.filter(
            venue__in=venues,
            start_datetime__date=date_str,
            status__in=['booked', 'rescheduled']
        ).order_by('start_datetime')
    else:
        if not venue_id:
            return JsonResponse({'available': False, 'reason': 'Missing venue or date'}, status=400)
        bookings = VenueBooking.objects.filter(
            venue_id=venue_id,
            start_datetime__date=date_str,
            status__in=['booked', 'rescheduled']
        ).order_by('start_datetime')
    available = not bookings.exists()
    booking_list = [
        {
            'start': b.start_datetime.strftime('%H:%M'),
            'end': b.end_datetime.strftime('%H:%M'),
            'purpose': b.booking_purpose,
            'status': b.status
        }
        for b in bookings
    ]
    return JsonResponse({'available': available, 'bookings': booking_list})

# Attendance analysis helper functions
def get_clock_in_category(clock_in_time, expected_clock_in):
    if not clock_in_time:
        return 'missing'
    if not expected_clock_in:
        return 'N/A'
    # Early is before 07:35, on_time is 07:35–08:05, late is after 08:05
    early_cutoff = time(7, 35)
    on_time_start = time(7, 35)
    on_time_end = time(8, 5)
    if clock_in_time < early_cutoff:
        return 'early'
    elif on_time_start <= clock_in_time <= on_time_end:
        return 'on_time'
    else:
        return 'late'

def get_expected_clock_out(clock_in_time, custom_clock_out=None):
    if custom_clock_out:
        return custom_clock_out
    if not clock_in_time:
        return time(16, 25)  # Default if no clock-in
    if clock_in_time <= time(7, 30):
        return time(15, 55)  # 3:55 PM
    return time(16, 25)  # 4:25 PM

def get_clock_out_category(clock_out_time, expected_clock_out):
    if not clock_out_time:
        return 'missing'
    if not expected_clock_out:
        return 'N/A'
    grace_period = timedelta(minutes=10)
    expected_datetime = datetime.combine(date.today(), expected_clock_out)
    clock_out_datetime = datetime.combine(date.today(), clock_out_time)
    if clock_out_datetime < expected_datetime - grace_period:
        return 'early_leave'
    elif clock_out_datetime <= expected_datetime + grace_period:
        return 'on_time'
    else:
        return 'overtime'

class LearnerAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):
        qs = Learner.objects.all()
        if self.q:
            qs = qs.filter(
                Q(FirstName__icontains=self.q) |
                Q(Surname__icontains=self.q)  |
                Q(IDNumber__icontains=self.q)
            )
        return qs

class FacilitatorAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):
        qs = LearnerRole.objects.filter(role__name='Facilitator').select_related('learner')
        if self.q:
            qs = qs.filter(
                Q(learner__FirstName__icontains=self.q) |
                Q(learner__Surname__icontains=self.q)
            )
        return qs

    def get_result_label(self, result):
        return f"{result.learner.FirstName} {result.learner.Surname}"

from django.contrib.auth.hashers import make_password

class SetLearnerPasswordView(RolePermissionRequiredMixin, View):
    template_name = 'core/set_password.html'

    def get(self, request, learner_id):
        learner = get_object_or_404(Learner, id=learner_id)
        return render(request, self.template_name, {'learner': learner})

    def post(self, request, learner_id):
        learner = get_object_or_404(Learner, id=learner_id)
        password = request.POST.get('password')
        
        if not learner.user:
            # Create new user
            username = f"learner_{learner.id}"
            user = User.objects.create(
                username=username,
                email=learner.EmailAddress,
                password=make_password(password)
            )
            learner.user = user
            learner.save()
        else:
            # Update existing user's password
            learner.user.set_password(password)
            learner.user.save()
        
        messages.success(request, f'Password set for {learner.FirstName} {learner.Surname}')
        return redirect('learner_list')
    

class SLADashboardView(ListView):
    model = SLA
    template_name = "core/sla_dashboard.html"
    context_object_name = "slas"

    def get_queryset(self):
        return SLA.objects.all().order_by('sla_reference')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        for sla in context["slas"]:
            bills = BillingHistory.objects.filter(sla=sla).order_by("invoice_date", "id")
            invoiced = bills.filter(billed=True).count()
            paid = bills.filter(payment_date__isnull=False).count()
            total_tranches = bills.count()
            for idx, bill in enumerate(bills, 1):
                bill.display_tranche = f"T{idx}"
            sla.invoiced_status = f"{invoiced}/{total_tranches} tranches"
            sla.paid_status = f"{paid}/{invoiced} paid" if invoiced > 0 else "None invoiced"
        return context

class SLADetailView(DetailView):
    model = SLA
    template_name = "core/sla_detail.html"
    context_object_name = "sla"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        qualifications = SLA_Qualifications.objects.filter(sla=self.object).select_related("service")
        context["quals"] = qualifications
        qual_learners = {}
        for qual in qualifications:
            learners = LearnerQualification.objects.filter(
                sla_qualification=qual
            ).select_related("learner")
            qual_learners[qual] = learners
        context["qual_learners"] = qual_learners
        bills = BillingHistory.objects.filter(sla=self.object).order_by("invoice_date", "id")
        for idx, bill in enumerate(bills, 1):
            bill.display_tranche = f"T{idx}"
            bill.display_label = bill.display_tranche
        context["bills"] = bills
        context["learners"] = LearnerQualification.objects.filter(sla_qualification__sla=self.object)
        if self.object.num_tranches == 99:
            context["display_num_tranches"] = "Progress Based"
        else:
            context["display_num_tranches"] = self.object.num_tranches
        return context

class AddLearnerQualificationView(RolePermissionRequiredMixin, View):
    def get(self, request, sla_id, qual_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        qual = get_object_or_404(SLA_Qualifications, pk=qual_id, sla=sla)
        search_query = request.GET.get('search', '')
        if search_query:
            learners = Learner.objects.filter(
                Q(FirstName__icontains=search_query) |
                Q(Surname__icontains=search_query) |
                Q(IDNumber__icontains=search_query)
            )
        else:
            learners = Learner.objects.all()[:50]
        return render(request, 'core/add_learner_qualification.html', {
            'sla': sla,
            'qual': qual,
            'learners': learners,
            'search_query': search_query
        })

    def post(self, request, sla_id, qual_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        qual = get_object_or_404(SLA_Qualifications, pk=qual_id, sla=sla)
        learner_id = request.POST.get('learner_id')
        learner = get_object_or_404(Learner, pk=learner_id)
        LearnerQualification.objects.get_or_create(sla_qualification=qual, learner=learner)
        return redirect('sla_detail', pk=sla_id)

class BillingExportView(RolePermissionRequiredMixin, View):
    template_name = "core/billing/export.html"

    def _build_months(self):
        today = datetime.today()
        months = []
        for i in range(-12, 13):
            d = today + relativedelta(months=i)
            months.append({
                'value': d.strftime('%Y-%m'),
                'label': d.strftime('%B %Y'),
            })
        return months

    def get(self, request):
        return render(request, self.template_name, {
            'available_months': self._build_months(),
            'selected_month': None,
            'selected_month_label': None,
            'export_rows': None,
        })

    def post(self, request):
        billing_month = request.POST.get('billing_month')
        do_csv = request.POST.get('generate_csv') == 'true'
        do_complete = request.POST.get('complete_billing') == 'true'
        months = self._build_months()

        if not billing_month:
            return render(request, self.template_name, {
                'available_months': months,
                'error': 'Please select a month',
                'selected_month': None,
                'selected_month_label': None,
                'export_rows': None,
            })

        year, month = map(int, billing_month.split('-'))
        selected_label = datetime(year, month, 1).strftime('%B %Y')
        accounting_period = month - 2 if month > 2 else month + 10
        last_day = calendar.monthrange(year, month)[1]
        last_day_str = datetime(year, month, last_day).strftime('%d/%m/%Y')
        invoices = BillingHistory.objects.filter(
            due_date__year=year,
            due_date__month=month
        ).order_by('sla__sla_reference', 'invoice_date', 'id')

        if do_complete:
            first_inv = request.POST.get('first_invoice', '').strip()
            billing_date = request.POST.get('billing_date')
            m = re.match(r'(\D*)(\d+)$', first_inv)
            if not m:
                messages.error(request, "Invalid invoice number format.")
                return redirect('billing_export')
            prefix, start_num = m.group(1), int(m.group(2))
            to_update = []
            for idx, inv in enumerate(invoices):
                inv.invoice_number = f"{prefix}{start_num + idx}"
                inv.invoice_date = billing_date
                inv.due_date = None
                inv.billed = True
                to_update.append(inv)
            BillingHistory.objects.bulk_update(
                to_update,
                ['invoice_number', 'invoice_date', 'due_date', 'billed']
            )
            messages.success(
                request,
                f"Marked {len(to_update)} invoices billed starting at {prefix}{start_num} on {billing_date}."
            )
            return redirect('billing_export')

        vat_obj = VATRate.objects.filter(active=True).first()
        if vat_obj:
            vat_code = vat_obj.code
            vat_multiplier = 1 + float(vat_obj.rate) / 100
        else:
            vat_code = '14'
            vat_multiplier = 1.15

        if do_csv:
            resp = HttpResponse(content_type='text/csv')
            resp['Content-Disposition'] = f'attachment; filename="billing_{billing_month}.csv"'
            writer = csv.writer(resp)
            header_middle_blanks = [''] * 6
            detail_trail_blanks = [''] * 13
            apost = "'"

            for inv in invoices:
                items_qs = BillingHistoryItem.objects.filter(billing_history=inv)
                cust_code = inv.sla.customer.code
                tranche_lbl = f"{inv.sla.sla_reference} {inv.invoice_type}"
                writer.writerow([
                    'Header',
                    inv.invoice_number or inv.invoice_type or '',
                    '',
                    '',
                    cust_code,
                    accounting_period,
                    inv.due_date.strftime('%d/%m/%Y'),
                    tranche_lbl,
                    'N',
                    '0',
                    'Ensemble Trading 460 PTY LTD',
                    'FNB ACC: 62430090647',
                    'Branch No: 251-655',
                    *header_middle_blanks,
                    '0',
                    last_day_str,
                    '', '', '',
                    '1'
                ])
                writer.writerow([
                    'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                    apost, '', 7,
                    *detail_trail_blanks
                ])
                for item in items_qs.select_related('sla_qualification__service'):
                    qual = item.sla_qualification
                    ex_vat = item.amount
                    inc_vat = ex_vat * vat_multiplier
                    writer.writerow([
                        'Detail', 0, 1,
                        f"{ex_vat:.2f}",
                        f"{inc_vat:.2f}",
                        '',
                        vat_code,
                        0,
                        0,
                        qual.service.gl_code,
                        qual.service.name,
                        6,
                        *detail_trail_blanks
                    ])
                    writer.writerow([
                        'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                        apost,
                        f"SAQA ID {qual.service.saqa_id or ''}",
                        7,
                        *detail_trail_blanks
                    ])
                    writer.writerow([
                        'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                        apost, '', 7,
                        *detail_trail_blanks
                    ])
                    active_qs = LearnerQualification.objects.filter(
                        sla_qualification=qual, status='active'
                    )
                    cnt = active_qs.count()
                    words = num2words(cnt, to='cardinal').upper()
                    writer.writerow([
                        'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                        apost,
                        f'Learners: {cnt} ({words})',
                        7,
                        *detail_trail_blanks
                    ])
                    for lq in active_qs.select_related('learner'):
                        ln = lq.learner
                        writer.writerow([
                            'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                            apost,
                            f'{ln.FirstName} {ln.Surname}',
                            7,
                            *detail_trail_blanks
                        ])
                    writer.writerow([
                        'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                        apost, '', 7,
                        *detail_trail_blanks
                    ])
                total_tranches = inv.sla.num_tranches
                tranche_num = int(inv.invoice_type.lstrip('T')) if inv.invoice_type and inv.invoice_type.startswith('T') else 1
                writer.writerow([
                    'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                    apost,
                    f'Ref: {inv.sla.sla_reference}',
                    7,
                    *detail_trail_blanks
                ])
                writer.writerow([
                    'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                    apost,
                    f'Tranche {tranche_num} ({tranche_num}/{total_tranches}) - {selected_label}',
                    7,
                    *detail_trail_blanks
                ])
                writer.writerow([
                    'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                    apost, '', 7,
                    *detail_trail_blanks
                ])
                writer.writerow([
                    'Detail', 0, 1, 0, 0, '', '0', 0, 0,
                    apost,
                    'Thank you for your valued business',
                    7,
                    *detail_trail_blanks
                ])
            return resp
        export_rows = []
        for inv in invoices:
            for item in BillingHistoryItem.objects.filter(billing_history=inv):
                qual = item.sla_qualification
                active_ct = LearnerQualification.objects.filter(
                    sla_qualification=qual,
                    status='active'
                ).count()
                export_rows.append({
                    'group': f"{inv.sla.sla_reference} – {inv.invoice_number or inv.invoice_type} – {inv.due_date:%b %e, %Y}",
                    'client': inv.sla.customer.name + (f" {inv.sla.end_client_name}" if inv.sla.end_client_name else ""),
                    'qualification': qual.service.name,
                    'active_learners': active_ct,
                    'amount': item.amount,
                })
        return render(request, self.template_name, {
            'available_months': months,
            'selected_month': billing_month,
            'selected_month_label': selected_label,
            'export_rows': export_rows,
        })

def upload_fingerprint(request):
    if request.method == 'POST':
        excel_files = request.FILES.getlist('excel_file')
        if not excel_files:
            messages.error(request, "No files selected. Please select at least one Excel file to upload.")
            return redirect('upload_fingerprint')

        fs = FileSystemStorage()
        total_files = len(excel_files)
        total_records_processed = 0
        all_unassigned_user_ids = set()

        for idx, excel_file in enumerate(excel_files, 1):
            filename = fs.save(excel_file.name, excel_file)
            file_path = fs.path(filename)
            records_processed = 0
            unassigned_user_ids = set()

            try:
                df = pd.read_excel(file_path)
                required_columns = ['User ID', 'Date/Time']
                if not all(col in df.columns for col in required_columns):
                    messages.error(request, f"File {excel_file.name}: Must contain 'User ID' and 'Date/Time' columns.")
                    continue

                for index, row in df.iterrows():
                    try:
                        user_id = int(row['User ID'])
                        datetime_value = row['Date/Time']

                        # Handle Excel date formats
                        if isinstance(datetime_value, (int, float)):
                            dt = datetime(1899, 12, 30) + pd.Timedelta(days=datetime_value)
                        else:
                            dt = pd.to_datetime(datetime_value, errors='raise')
                        date_value = dt.date()
                        time_value = dt.time()
                    except (ValueError, TypeError) as e:
                        logger.warning(f"Invalid data at row {index + 2} in {excel_file.name}: {str(e)}")
                        messages.warning(request, f"File {excel_file.name}: Invalid Date/Time at row {index + 2}. Skipping.")
                        continue

                    # Try to find learner, but save record even if none exists
                    learner = None
                    try:
                        learner = Learner.objects.get(UserID=user_id)
                    except Learner.DoesNotExist:
                        unassigned_user_ids.add(user_id)

                    # Alternate clock-in/clock-out based on existing records
                    existing_records = Fingerprint.objects.filter(user_id=user_id, date=date_value).count()
                    is_clock_in = existing_records % 2 == 0

                    try:
                        Fingerprint.objects.update_or_create(
                            user_id=user_id,
                            date=date_value,
                            time=time_value,
                            defaults={
                                'learner': learner,
                                'is_clock_in': is_clock_in
                            }
                        )
                        records_processed += 1
                        total_records_processed += 1
                    except Exception as e:
                        logger.error(f"Error saving fingerprint for User ID {user_id} in {excel_file.name}: {str(e)}")
                        messages.warning(request, f"File {excel_file.name}: Error saving record for User ID {user_id} at row {index + 2}. Skipping.")

                # Provide feedback for this file
                messages.success(request, f"File {excel_file.name}: Processed {records_processed} records successfully.")
                if unassigned_user_ids:
                    all_unassigned_user_ids.update(unassigned_user_ids)
                    messages.info(request, f"File {excel_file.name}: Saved records for unassigned User IDs: {', '.join(map(str, unassigned_user_ids))}. Assign learners to these User IDs to link records.")
            except Exception as e:
                logger.error(f"Error processing file {excel_file.name}: {str(e)}")
                messages.error(request, f"File {excel_file.name}: Error processing file - {str(e)}")
            finally:
                fs.delete(filename)

        # Summary message
        messages.success(request, f"Processed {total_files} file(s) with a total of {total_records_processed} records.")
        if all_unassigned_user_ids:
            messages.info(request, f"Total unassigned User IDs across all files: {', '.join(map(str, all_unassigned_user_ids))}. Consider assigning learners to these IDs.")

        return redirect('fingerprint_list')

    return render(request, 'core/upload_fingerprint.html')

class FingerprintListView(RolePermissionRequiredMixin, ListView):
    model = Fingerprint
    template_name = 'core/fingerprint_list.html'
    context_object_name = 'fingerprint_data'
    paginate_by = 10  # Number of records per page

    def get_queryset(self):
        queryset = Fingerprint.objects.select_related('learner').order_by('user_id', 'date', 'time')
        user_id = self.request.GET.get('user_id')
        date_filter = self.request.GET.get('date')
        sla_id = self.request.GET.get('sla_id')
        start_date = self.request.GET.get('start_date')
        end_date = self.request.GET.get('end_date')
        search = self.request.GET.get('search')
        qualification_search = self.request.GET.get('qualification_search', '')

        if user_id:
            try:
                user_id = int(user_id)
                queryset = queryset.filter(user_id=user_id)
            except ValueError:
                pass
        if date_filter:
            try:
                queryset = queryset.filter(date=date_filter)
            except ValueError:
                pass
        elif start_date and end_date:
            try:
                queryset = queryset.filter(date__range=[start_date, end_date])
            except ValueError:
                pass
        elif start_date:
            try:
                queryset = queryset.filter(date__gte=start_date)
            except ValueError:
                pass
        elif end_date:
            try:
                queryset = queryset.filter(date__lte=end_date)
            except ValueError:
                pass
        if sla_id:
            try:
                sla_id = int(sla_id)
                queryset = queryset.filter(
                    learner__learnerqualification__sla_qualification__sla_id=sla_id
                )
            except ValueError:
                pass
        if search:
            learners = Learner.objects.filter(
                Q(FirstName__icontains=search) |
                Q(Surname__icontains=search) |
                Q(IDNumber__icontains=search) |
                Q(UserID__icontains=search)
            )
            user_ids = learners.values_list('UserID', flat=True)
            queryset = queryset.filter(user_id__in=user_ids)
        if qualification_search:
            queryset = queryset.filter(
                learner__learnerqualification__sla_qualification__service__name__icontains=qualification_search
            )
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        fingerprints = self.get_queryset()

        # Cache key based on filters
        cache_key = f"fingerprint_list_{self.request.GET.urlencode()}"
        fingerprint_data = cache.get(cache_key)

        if fingerprint_data is None:
            fingerprint_data = []
            # Group fingerprints by user_id and date
            grouped = {}
            for fp in fingerprints:
                key = (fp.user_id, fp.date)
                if key not in grouped:
                    grouped[key] = {'clock_in': None, 'clock_out': None}
                if fp.is_clock_in:
                    grouped[key]['clock_in'] = fp
                else:
                    grouped[key]['clock_out'] = fp

            for (user_id, date_val), times in grouped.items():
                clock_in = times['clock_in']
                clock_out = times['clock_out']

                # Try to get learner by Fingerprint.learner, else by UserID
                learner = None
                if clock_in and clock_in.learner:
                    learner = clock_in.learner
                elif clock_out and clock_out.learner:
                    learner = clock_out.learner
                else:
                    # fallback: try to get by UserID
                    try:
                        learner = Learner.objects.get(UserID=user_id)
                    except Learner.DoesNotExist:
                        learner = None

                learner_name = f"{learner.FirstName} {learner.Surname}" if learner else "Unassigned"
                sla_ref = None
                qualification = "N/A"
                expected_clock_in = None
                expected_clock_out = None
                if learner:
                    lq = LearnerQualification.objects.filter(
                        learner=learner,
                        status='active'
                    ).select_related('sla_qualification__sla', 'sla_qualification__service').first()
                    if lq:
                        sla_ref = lq.sla_qualification.sla.sla_reference
                        qualification = lq.sla_qualification.service.name
                        # Check WeeklySchedule for the specific day
                        day_name = date_val.strftime('%A').lower()
                        weekly_schedule = WeeklySchedule.objects.filter(
                            learner_qualification=lq,
                            day=day_name
                        ).first()
                        if weekly_schedule:
                            expected_clock_in = weekly_schedule.clock_in
                            expected_clock_out = weekly_schedule.clock_out
                        # Fall back to existing defaults if no weekly schedule
                        if not expected_clock_in:
                            expected_clock_in = (
                                lq.expected_clock_in or
                                learner.expected_clock_in or
                                lq.sla_qualification.expected_clock_in or
                                time(7, 55)  # Default is 07:55
                            )
                        if not expected_clock_out:
                            expected_clock_out = (
                                lq.expected_clock_out or
                                learner.expected_clock_out or
                                lq.sla_qualification.expected_clock_out
                            )
                sla_group = sla_ref if sla_ref else 'N/A'

                clock_in_category = get_clock_in_category(
                    clock_in.time if clock_in else None,
                    expected_clock_in
                )
                clock_out_category = get_clock_out_category(
                    clock_out.time if clock_out else None,
                    get_expected_clock_out(
                        clock_in.time if clock_in else None,
                        expected_clock_out
                    )
                )

                fingerprint_data.append({
                    'user_id': user_id,
                    'learner_name': learner_name,
                    'sla_group': sla_group,
                    'qualification': qualification,
                    'date': date_val,
                    'clock_in_time': clock_in.time if clock_in else None,
                    'clock_out_time': clock_out.time if clock_out else None,
                    'clock_in_category': clock_in_category,
                    'clock_out_category': clock_out_category,
                    'expected_clock_in': expected_clock_in,
                    'expected_clock_out': get_expected_clock_out(
                        clock_in.time if clock_in else None,
                        expected_clock_out
                    ),
                })

            # Sort fingerprint_data by date for consistent grouping
            fingerprint_data.sort(key=lambda x: x['date'])
            cache.set(cache_key, fingerprint_data, timeout=3600)

        # Paginate the processed fingerprint_data
        paginator = Paginator(fingerprint_data, self.paginate_by)
        page_number = self.request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        # Update context with paginated data
        context['fingerprint_data'] = page_obj
        context['is_paginated'] = page_obj.has_other_pages()
        context['page_obj'] = page_obj

        # Compute attendance summary
        summary = {
            'clock_in': {'early': 0, 'on_time': 0, 'late': 0, 'missing': 0},
            'clock_out': {'early_leave': 0, 'on_time': 0, 'overtime': 0, 'missing': 0}
        }
        for record in fingerprint_data:  # Use full dataset for summary
            ci_cat = record['clock_in_category']
            if ci_cat in summary['clock_in']:
                summary['clock_in'][ci_cat] += 1
            co_cat = record['clock_out_category']
            if co_cat in summary['clock_out']:
                summary['clock_out'][co_cat] += 1
        context['summary'] = summary

        context['filters'] = {
            'user_id': self.request.GET.get('user_id', ''),
            'date': self.request.GET.get('date', ''),
            'sla_id': self.request.GET.get('sla_id', ''),
            'search': self.request.GET.get('search', ''),
            'qualification_search': self.request.GET.get('qualification_search', ''),
            'start_date': self.request.GET.get('start_date', ''),
            'end_date': self.request.GET.get('end_date', ''),
        }
        context['slas'] = SLA.objects.all().order_by('sla_reference')
        return context

class AssignUserIDView(RolePermissionRequiredMixin, View):
    template_name = 'core/assign_user_id.html'

    def get(self, request, learner_id):
        learner = get_object_or_404(Learner, pk=learner_id)
        unassigned_fingerprints = Fingerprint.objects.filter(learner__isnull=True).values('user_id').distinct()
        return render(request, self.template_name, {
            'learner': learner,
            'unassigned_user_ids': [fp['user_id'] for fp in unassigned_fingerprints]
        })

    def post(self, request, learner_id):
        learner = get_object_or_404(Learner, pk=learner_id)
        user_id = request.POST.get('user_id')
        try:
            user_id = int(user_id) if user_id else None
            if user_id and Learner.objects.filter(UserID=user_id).exclude(pk=learner.id).exists():
                messages.error(request, f"User ID {user_id} is already assigned to another learner.")
            else:
                learner.UserID = user_id
                learner.save()
                if user_id:
                    updated = Fingerprint.objects.filter(
                        user_id=user_id, 
                        learner__isnull=True
                    ).update(learner=learner)
                    if updated:
                        messages.info(request, f"Linked {updated} fingerprint records to {learner}.")
                messages.success(request, f"User ID {user_id or 'cleared'} assigned to {learner}.")
        except ValueError:
            messages.error(request, "Invalid User ID. Please enter a valid integer.")
        return redirect('assign_user_id', learner_id=learner_id)

class LearnerListView(RolePermissionRequiredMixin, ListView):
    model = Learner
    template_name = 'core/learner_list.html'
    context_object_name = 'learners'
    paginate_by = 10  # Number of learners per page

    def get_queryset(self):
        queryset = Learner.objects.all().order_by('Surname', 'FirstName')
        search_query = self.request.GET.get('search', '')
        qualification_search = self.request.GET.get('qualification_search', '')
        status = self.request.GET.get('status', '')

        if search_query:
            queryset = queryset.filter(
                Q(FirstName__icontains=search_query) |
                Q(Surname__icontains=search_query) |
                Q(IDNumber__icontains=search_query) |
                Q(UserID__icontains=search_query)
            )
        if qualification_search:
            queryset = queryset.filter(
                learnerqualification__sla_qualification__service__name__icontains=qualification_search,
                learnerqualification__status='active'
            ).distinct()
        if status:
            queryset = queryset.filter(learnerqualification__status=status).distinct()
        return queryset

    def post(self, request, *args, **kwargs):
        learner_id = request.POST.get('learner_id')
        action = request.POST.get('action')
        if action == 'deactivate':
            learner = get_object_or_404(Learner, pk=learner_id)
            updated = LearnerQualification.objects.filter(learner=learner, status='active').update(status='removed')
            if updated:
                messages.success(request, f"Learner {learner.FirstName} {learner.Surname} deactivated from {updated} qualification(s).")
            else:
                messages.info(request, f"Learner {learner.FirstName} {learner.Surname} has no active qualifications to deactivate.")
        return redirect('learner_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        queryset = self.get_queryset()
        paginator = Paginator(queryset, self.paginate_by)
        page_number = self.request.GET.get('page', 1)
        page_obj = paginator.get_page(page_number)

        for learner in page_obj:
            learner.qualifications = LearnerQualification.objects.filter(
                learner=learner
            ).select_related('sla_qualification__service', 'sla_qualification__sla')

        context['page_obj'] = page_obj
        context['is_paginated'] = page_obj.has_other_pages()
        context['filters'] = {
            'search': self.request.GET.get('search', ''),
            'qualification_search': self.request.GET.get('qualification_search', ''),
            'status': self.request.GET.get('status', ''),
        }
        return context

class LearnerScheduleListView(RolePermissionRequiredMixin, ListView):
    model = Learner
    template_name = 'core/learner_schedule_list.html'
    context_object_name = 'learners'

    def get_queryset(self):
        queryset = Learner.objects.all().order_by('Surname', 'FirstName')
        learner_id = self.request.GET.get('learner_id')
        if learner_id:
            try:
                learner_id = int(learner_id)
                queryset = queryset.filter(id=learner_id)
            except ValueError:
                pass
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        for learner in context['learners']:
            learner.qualifications = LearnerQualification.objects.filter(
                learner=learner
            ).select_related('sla_qualification__service', 'sla_qualification__sla')
            for lq in learner.qualifications:
                lq.schedules = WeeklySchedule.objects.filter(learner_qualification=lq)
        context['filters'] = {
            'learner_id': self.request.GET.get('learner_id', ''),
        }
        return context

class EditLearnerTimesView(RolePermissionRequiredMixin, View):
    template_name = 'core/edit_learner_weekly_times.html'

    def get(self, request, learner_qualification_id):
        lq = get_object_or_404(LearnerQualification, pk=learner_qualification_id)
        schedules = WeeklySchedule.objects.filter(learner_qualification=lq)
        schedule_data = {schedule.day: {'clock_in': schedule.clock_in, 'clock_out': schedule.clock_out} for schedule in schedules}
        days = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
        for day in days:
            if day not in schedule_data:
                schedule_data[day] = {'clock_in': None, 'clock_out': None}
        return render(request, self.template_name, {
            'learner_qualification': lq,
            'learner': lq.learner,
            'qualification': lq.sla_qualification.service.name,
            'sla': lq.sla_qualification.sla.sla_reference,
            'schedule_data': schedule_data,
            'days': days,
        })

    def post(self, request, learner_qualification_id):
        lq = get_object_or_404(LearnerQualification, pk=learner_qualification_id)
        WeeklySchedule.objects.filter(learner_qualification=lq).delete()
        days = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
        for day in days:
            clock_in = request.POST.get(f'{day}_clock_in')
            clock_out = request.POST.get(f'{day}_clock_out')
            if clock_in or clock_out:
                try:
                    clock_in_time = time.fromisoformat(clock_in) if clock_in else None
                    clock_out_time = time.fromisoformat(clock_out) if clock_out else None
                    WeeklySchedule.objects.create(
                        learner_qualification=lq,
                        day=day,
                        clock_in=clock_in_time,
                        clock_out=clock_out_time
                    )
                except ValueError:
                    messages.error(request, f"Invalid time format for {day}. Please use HH:MM format.")
                    return redirect('edit_learner_times', learner_qualification_id=lq.id)
        messages.success(request, f"Updated weekly schedule for {lq.learner} in {lq.sla_qualification.service.name}.")
        return redirect('learner_schedule_list')

class EditQualificationTimesView(RolePermissionRequiredMixin, View):
    template_name = 'core/edit_qualification_times.html'

    def get(self, request, qual_id):
        qual = get_object_or_404(SLA_Qualifications, pk=qual_id)
        return render(request, self.template_name, {
            'qual': qual,
            'expected_clock_in': qual.expected_clock_in,
            'expected_clock_out': qual.expected_clock_out,
        })

    def post(self, request, qual_id):
        qual = get_object_or_404(SLA_Qualifications, pk=qual_id)
        clock_in = request.POST.get('expected_clock_in')
        clock_out = request.POST.get('expected_clock_out')
        try:
            qual.expected_clock_in = clock_in if clock_in else None
            qual.expected_clock_out = clock_out if clock_out else None
            qual.save()
            messages.success(request, f"Updated attendance times for {qual.service.name}.")
        except ValueError:
            messages.error(request, "Invalid time format.")
        return redirect('qualification_times_list')

class QualificationTimesListView(RolePermissionRequiredMixin, ListView):
    model = SLA_Qualifications
    template_name = 'core/qualification_times_list.html'
    context_object_name = 'qualifications'
    paginate_by = 10  # Add this line

    def get_queryset(self):
        queryset = SLA_Qualifications.objects.select_related('sla', 'service').order_by('sla__sla_reference', 'service__name')
        sla_id = self.request.GET.get('sla_id')
        qualification_name = self.request.GET.get('qualification_name', '')
        status = self.request.GET.get('status', '')

        if sla_id:
            try:
                sla_id = int(sla_id)
                queryset = queryset.filter(sla_id=sla_id)
            except ValueError:
                logger.warning(f"Invalid sla_id: {sla_id}")
        if qualification_name:
            queryset = queryset.filter(service__name__icontains=qualification_name)
        if status:
            if status == 'active':
                queryset = queryset.filter(active=True)
            elif status == 'inactive':
                queryset = queryset.filter(active=False)
        logger.info(f"Retrieved {queryset.count()} qualifications for QualificationTimesListView")
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['slas'] = SLA.objects.all().order_by('sla_reference')
        context['filters'] = {
            'sla_id': self.request.GET.get('sla_id', ''),
            'qualification_name': self.request.GET.get('qualification_name', ''),
            'status': self.request.GET.get('status', ''),
        }
        context['qualification_count'] = context['paginator'].count if 'paginator' in context else context['qualifications'].count()
        return context

class EditLearnerView(RolePermissionRequiredMixin, UpdateView):
    model = Learner
    fields = ['FirstName', 'Surname', 'IDNumber', 'Gender', 'Equity', 'EmailAddress', 'UserID']
    template_name = 'core/edit_learner.html'
    success_url = reverse_lazy('learner_list')
    pk_url_kwarg = 'learner_id'

class LearnerDetailsView(RolePermissionRequiredMixin, DetailView):
    model = Learner
    template_name = 'core/learner_details.html'
    pk_url_kwarg = 'learner_id'

# ─── New Views for Added Models ──────────────────────────────────────

# Set up logging
logger = logging.getLogger(__name__)

class LearnerRoleAssignmentView(RolePermissionRequiredMixin, View):
    template_name = 'core/learner_role_assignment.html'
    paginate_by = 10  # Number of learners per page

    def get(self, request):
        # Get filter parameters
        search_query = request.GET.get('search', '')
        qualification_search = request.GET.get('qualification_search', '')
        status = request.GET.get('status', '')
        role_filter = request.GET.get('role', '')

        # Log filter parameters
        logger.debug(f"Filter parameters: search='{search_query}', qualification_search='{qualification_search}', status='{status}', role='{role_filter}'")

        # Base queryset with minimal fields for performance
        learners = Learner.objects.order_by('Surname', 'FirstName')

        # Apply filters
        if search_query:
            learners = learners.filter(
                Q(FirstName__icontains=search_query) |
                Q(Surname__icontains=search_query) |
                Q(IDNumber__icontains=search_query) |
                Q(UserID__icontains=search_query)
            ).distinct()
            logger.debug(f"After search filter ('{search_query}'): {learners.count()} learners")

        if qualification_search:
            learners = learners.filter(
                learnerqualification__sla_qualification__service__name__icontains=qualification_search,
                learnerqualification__status='active'
            ).distinct()
            logger.debug(f"After qualification filter ('{qualification_search}'): {learners.count()} learners")

        if status:
            learners = learners.filter(
                learnerqualification__status=status
            ).distinct()
            logger.debug(f"After status filter ('{status}'): {learners.count()} learners")

        if role_filter:
            learners = learners.filter(
                learnerrole__role__name=role_filter
            ).distinct()
            logger.debug(f"After role filter ('{role_filter}'): {learners.count()} learners")

        # Log total learners
        total_learners = learners.count()
        logger.debug(f"Total learners after filters: {total_learners}")

        # Paginate
        paginator = Paginator(learners, self.paginate_by)
        page_number = request.GET.get('page', 1)
        try:
            page_obj = paginator.page(page_number)
        except PageNotAnInteger:
            logger.warning(f"Invalid page number '{page_number}', defaulting to page 1")
            page_obj = paginator.page(1)
        except EmptyPage:
            logger.warning(f"Page {page_number} is empty, redirecting to last page {paginator.num_pages}")
            page_obj = paginator.page(paginator.num_pages)

        # Fetch roles and enrich learners with related data
        roles = Role.objects.all()
        for learner in page_obj:
            learner.roles = learner.learnerrole_set.all()
            learner.qualifications = learner.learnerqualification_set.select_related(
                'sla_qualification__service',
                'sla_qualification__sla'
            )

        # Get available qualifications for the dropdown
        available_qualifications = LearnerQualification.objects.filter(
            learner__in=learners
        ).values_list('sla_qualification__service__name', flat=True).distinct()

        # Prepare context
        context = {
            'page_obj': page_obj,
            'filters': {
                'search': search_query,
                'qualification_search': qualification_search,
                'status': status,
                'role': role_filter,
            },
            'roles': roles,
            'no_results': total_learners == 0 and any([search_query, qualification_search, status, role_filter]),
            'available_qualifications': available_qualifications,
        }

        # Log rendering details
        logger.debug(f"Rendering page {page_obj.number} of {paginator.num_pages}, showing {len(page_obj)} learners")

        return render(request, self.template_name, context)

    def post(self, request):
        action = request.POST.get('action')
        logger.debug(f"POST data: {request.POST}")

        if action == 'add_role':
            learner_id = request.POST.get('learner_id')
            role_id = request.POST.get('role_id')
            if learner_id and role_id:
                try:
                    learner = get_object_or_404(Learner, pk=learner_id)
                    role = get_object_or_404(Role, pk=role_id)
                    LearnerRole.objects.get_or_create(learner=learner, role=role)
                    messages.success(request, f"Added role '{role.name}' to {learner.FirstName} {learner.Surname}.")
                except Exception as e:
                    logger.error(f"Error adding role: {str(e)}")
                    messages.error(request, f"Error adding role: {str(e)}")
            else:
                messages.error(request, "Please select a learner and a role.")
            return redirect('learner_role_assignment')

        elif action == 'remove_role':
            learner_id = request.POST.get('learner_id')
            role_id = request.POST.get('role_id')
            if learner_id and role_id:
                try:
                    learner = get_object_or_404(Learner, pk=learner_id)
                    role = get_object_or_404(Role, pk=role_id)
                    LearnerRole.objects.filter(learner=learner, role=role).delete()
                    messages.success(request, f"Removed role '{role.name}' from {learner.FirstName} {learner.Surname}.")
                except Exception as e:
                    logger.error(f"Error removing role: {str(e)}")
                    messages.error(request, f"Error removing role: {str(e)}")
            else:
                messages.error(request, "Please select a learner and a role.")
            return redirect('learner_role_assignment')

        elif action == 'assign_role_to_selected':
            learner_ids = request.POST.getlist('learner_ids')
            role_id = request.POST.get('role_id')
            logger.debug(f"Bulk assignment - learner_ids: {learner_ids}, role_id: {role_id}")
            if learner_ids and role_id:
                try:
                    role = get_object_or_404(Role, pk=role_id)
                    for learner_id in learner_ids:
                        learner = get_object_or_404(Learner, pk=learner_id)
                        LearnerRole.objects.get_or_create(learner=learner, role=role)
                    messages.success(request, f"Assigned role '{role.name}' to {len(learner_ids)} learner(s).")
                except Exception as e:
                    logger.error(f"Error assigning role {role_id} to learners {learner_ids}: {str(e)}")
                    messages.error(request, f"Error assigning role: {str(e)}")
            else:
                logger.warning(f"Bulk role assignment failed: learner_ids={learner_ids}, role_id={role_id}")
                messages.error(request, "Please select at least one learner and a role.")
            return redirect('learner_role_assignment')

        elif action == 'assign_role_to_all':
            role_id = request.POST.get('role_id')
            logger.debug(f"Assign to all - role_id: {role_id}")
            if role_id:
                try:
                    role = get_object_or_404(Role, pk=role_id)
                    learners = Learner.objects.all()
                    assignment_count = 0
                    for learner in learners:
                        created = LearnerRole.objects.get_or_create(learner=learner, role=role)[1]
                        if created:
                            assignment_count += 1
                    messages.success(request, f"Assigned role '{role.name}' to {assignment_count} learner(s).")
                except Exception as e:
                    logger.error(f"Error assigning role {role_id} to all learners: {str(e)}")
                    messages.error(request, f"Error assigning role: {str(e)}")
            else:
                messages.error(request, "Please select a role.")
            return redirect('learner_role_assignment')

        else:
            messages.error(request, "Invalid action.")
            return redirect('learner_role_assignment')

def assign_learner_qualification_to_group(request):
    if request.method == 'POST':
        group_id = request.POST.get('group')
        lq_id = request.POST.get('learner_qualification')
        group = get_object_or_404(Group, pk=group_id)
        lq = get_object_or_404(LearnerQualification, pk=lq_id)
        lq.group = group
        lq.save()
        messages.success(request, f"Assigned {lq} to {group}.")
    return redirect('group_management')

def get_sla_qualifications(request):
    service_id = request.GET.get('service_id')
    if not service_id:
        return JsonResponse({'results': []})
        
    qualifications = SLA_Qualifications.objects.filter(
        service_id=service_id
    ).select_related('sla', 'service')
    
    results = [{
        'id': qual.id,
        'label': f"{qual.service.name} - {qual.sla.sla_reference} ({qual.learner_count} learners)"
    } for qual in qualifications]
    
    return JsonResponse({'results': results})

from django.views import View
from django.shortcuts import render, get_object_or_404, redirect
from django.core.paginator import Paginator
from django.db.models import Q
from .models import Group, Service, SETA, SLA_Qualifications, LearnerRole
from .forms import GroupCreateForm, SETAForm

class GroupManagementView(RolePermissionRequiredMixin, View):
    template_name = 'core/group_management.html'
    paginate_by = 15

    def get(self, request):
        search = request.GET.get('search', '').strip()
        service = request.GET.get('service', '')
        seta = request.GET.get('seta', '')

        groups = Group.objects.select_related('service', 'seta', 'project_lead', 'etqa_lead') \
            .prefetch_related('sla_qualifications', 'sla_qualifications__service')

        if search:
            groups = groups.filter(
                Q(name__icontains=search) |
                Q(projectcode__icontains=search)
            )
        # Filter by qualification (service)
        if service:
            groups = groups.filter(sla_qualifications__service_id=service).distinct()
        if seta:
            groups = groups.filter(seta_id=seta)

        groups = groups.order_by('-start_date', 'name')

        paginator = Paginator(groups, self.paginate_by)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        services = Service.objects.all().order_by('name')
        setas = SETA.objects.all().order_by('name')

        return render(request, self.template_name, {
            'groups': page_obj,
            'services': services,
            'setas': setas,
            'filters': {
                'search': search,
                'service': service,
                'seta': seta,
            },
            'page_obj': page_obj,
        })

from .models import LIF, LIFTemplate
from django.urls import reverse
class GroupDetailView(RolePermissionRequiredMixin, View):
    template_name = 'core/group_detail.html'

    def get(self, request, pk):
        group = get_object_or_404(
            Group.objects.select_related('service', 'seta', 'project_lead', 'etqa_lead').prefetch_related('sla_qualifications'),
            pk=pk
        )
        sla_qualifications = group.sla_qualifications.select_related('service', 'sla').all()
        group_learners = Learner.objects.filter(
            learnerqualification__sla_qualification__groups=group
        ).distinct().prefetch_related(
            'learnerqualification_set__sla_qualification__service',
            'learnerqualification_set__sla_qualification__sla'
        )
        session_dates = SessionDate.objects.filter(
            project_plan__group=group
        ).select_related('project_plan__module', 'facilitator')

        # LIF status for each learner
        lif_map = {lif.learner_id: lif for lif in LIF.objects.filter(learner__in=group_learners)}
        lif_templates = LIFTemplate.objects.all()

        return render(request, self.template_name, {
            'group': group,
            'sla_qualifications': sla_qualifications,
            'group_learners': group_learners,
            'session_dates': session_dates,
            'lif_map': lif_map,
            'lif_templates': lif_templates,
        })

    def post(self, request, pk):
        learner_id = request.POST.get('lif_learner_id')
        template_id = request.POST.get('lif_template_id')
        if learner_id and template_id:
            url = reverse('generate_lif_word') + f'?learner_id={learner_id}&template_id={template_id}'
            return redirect(url)
        return redirect('group_detail', pk=pk)

class GroupCreateView(RolePermissionRequiredMixin, View):
    template_name = 'core/group_detail.html'

    def get(self, request, pk):
        group = get_object_or_404(
            Group.objects.select_related('service', 'seta', 'project_lead', 'etqa_lead').prefetch_related('sla_qualifications'),
            pk=pk
        )
        sla_qualifications = group.sla_qualifications.select_related('service', 'sla').all()
        # Get all learners linked to this group via SLA_Qualifications
        group_learners = Learner.objects.filter(
            learnerqualification__sla_qualification__groups=group
        ).distinct().prefetch_related('learnerqualification_set__sla_qualification__service', 'learnerqualification_set__sla_qualification__sla')
        # Get all session dates for this group via ProjectPlan
        session_dates = SessionDate.objects.filter(
            project_plan__group=group
        ).select_related('project_plan__module', 'facilitator')
        # Get all project plans for this group
        project_plans = ProjectPlan.objects.filter(group=group).select_related('module')
        return render(request, self.template_name, {
            'group': group,
            'sla_qualifications': sla_qualifications,
            'group_learners': group_learners,
            'session_dates': session_dates,
            'project_plans': project_plans,
        })

    def post(self, request):
        service_id = request.POST.get('service')
        form = GroupCreateForm(request.POST, service_id=service_id)
        seta_form = SETAForm()
        if 'add_seta' in request.POST:
            seta_form = SETAForm(request.POST)
            if seta_form.is_valid():
                seta_form.save()
                messages.success(request, "SETA added successfully.")
                return redirect('group_create')
        elif form.is_valid():
            group = form.save()
            messages.success(request, "Group created successfully.")
            return redirect('group_list')
        return render(request, self.template_name, {
            'form': form,
            'seta_form': seta_form,
        })

# AJAX endpoint for dynamic LearnerQualification filtering
def get_learner_qualifications(request):
    service_id = request.GET.get('service_id')
    if not service_id:
        return JsonResponse({'results': []})
    lqs = LearnerQualification.objects.filter(sla_qualification__service_id=service_id)
    results = [
        {'id': lq.id, 'label': str(lq)}
        for lq in lqs
    ]
    return JsonResponse({'results': results})

class BillingPaymentsView(RolePermissionRequiredMixin, View):
    template_name = "core/billing/payments.html"

    def get(self, request):
        # month comes in as "YYYY-MM"
        month = request.GET.get("month")
        context = {"selected_month": month}
        if month:
            year, m = map(int, month.split("-"))
            last_day = calendar.monthrange(year, m)[1]
            end_of_month = date(year, m, last_day)

            # Section 1: outstanding (unpaid) invoices
            unpaid = BillingHistory.objects.filter(
                invoice_date__lte=end_of_month,
                payment_date__isnull=True,
                billed=True,
            )

            # Section 2: any due_date ≤ end-of-month but not yet invoiced
            missed = BillingHistory.objects.filter(
                due_date__lte=end_of_month,
                invoice_date__isnull=True,
            )

            context.update({"unpaid": unpaid, "missed": missed})
        return render(request, self.template_name, context)

    def post(self, request):
        month = request.POST.get("selected_month")
        if not month:
            return redirect(request.path)

        year, m = map(int, month.split("-"))
        last_day = calendar.monthrange(year, m)[1]
        end_of_month = date(year, m, last_day)

        unpaid_qs = BillingHistory.objects.filter(
            invoice_date__lte=end_of_month,
            payment_date__isnull=True,
            billed=True,
        )
        for inv in unpaid_qs:
            # each input named payment_date_<pk>
            pd = request.POST.get(f"payment_date_{inv.pk}")
            if pd:
                inv.payment_date = pd
                inv.save()

        # redirect with the same month in the querystring
        return redirect(f"{request.path}?month={month}")
        
class FinanceDashboardView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    """
    A simple landing page for all finance-related tools.
    Only staff users may access.
    """
    template_name = "core/finance/dashboard.html"
    login_url = "login"          # or whatever your login URL name is

    def test_func(self):
        # only staff can access—for more granular control you could test group membership or permissions
        return self.request.user.is_staff
        
class SLAWizardStep1View(LoginRequiredMixin, UserPassesTestMixin, View):
    template_name = "core/finance/sla_add_step1.html"

    def test_func(self):
        return self.request.user.is_staff

    def get(self, request):
        form = SLAForm()
        return render(request, self.template_name, {"form": form})

    def post(self, request):
        form = SLAForm(request.POST)
        if form.is_valid():
            sla = form.save()
            return redirect("add_sla_qualifications", sla_id=sla.id)
        return render(request, self.template_name, {"form": form})
        
class SLAWizardStep2View(LoginRequiredMixin, UserPassesTestMixin, View):
    template_name = "core/finance/sla_add_step2.html"

    def test_func(self):
        return self.request.user.is_staff

    def get(self, request, sla_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        formset = SLAQualificationFormSet(
            queryset=SLA_Qualifications.objects.filter(sla=sla)
        )
        return render(
            request,
            self.template_name,
            {
                "sla": sla,
                "formset": formset,
            }
        )

    def post(self, request, sla_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        formset = SLAQualificationFormSet(
            request.POST,
            queryset=SLA_Qualifications.objects.filter(sla=sla)
        )
        if formset.is_valid():
            instances = formset.save(commit=False)
            for inst in instances:
                inst.sla = sla
                inst.save()
            formset.save_m2m()

            if sla.num_tranches == 99:
                return redirect("add_sla_learners", sla_id=sla.id)

            return redirect("add_sla_billing", sla_id=sla.id)

        return render(
            request,
            self.template_name,
            {
                "sla": sla,
                "formset": formset,
            }
        )
        
class SLAWizardStep3View(LoginRequiredMixin, UserPassesTestMixin, View):
    template_name = "core/finance/sla_add_step3.html"
    formset_class = formset_factory(TrancheBillingForm, extra=0)

    def test_func(self):
        return self.request.user.is_staff

    def _get_qualifications(self, sla):
        return SLA_Qualifications.objects.filter(sla=sla).select_related('service')

    def _build_initial(self, sla):
        initial = []
        for idx in range(sla.num_tranches):
            invoice_type = f"T{idx+1}"
            row = {'invoice_type': invoice_type}
            try:
                bh = BillingHistory.objects.get(sla=sla, invoice_type=invoice_type)
                row.update({
                    'due_date':     bh.due_date,
                    'total_amount': bh.amount,
                })
                for item in bh.items.all():
                    key = f"qual_{item.sla_qualification_id}"
                    row[key] = item.amount
            except BillingHistory.DoesNotExist:
                pass
            initial.append(row)
        return initial

    def _annotate_qual_fields(self, formset, quals):
        """
        Attach a list of BoundField objects for each SLA_Qualifications on the SLA,
        with a more descriptive label so duplicates are distinguishable.
        """
        for form in formset:
            form.qual_fields = []
            for qual in quals:
                label = (
                    f"{qual.service.name} "
                    f"(Count: {qual.learner_count}"
                    f"{', ' + qual.employment_status if qual.employment_status else ''}"
                    f") [#{qual.id}]"
                )
                form.qual_fields.append({
                    'label': label,
                    'field': form[f"qual_{qual.id}"]
                })

    def get(self, request, sla_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        if sla.num_tranches == 99:
            return redirect("add_sla_learners", sla_id=sla.id)

        quals = self._get_qualifications(sla)
        initial = self._build_initial(sla)
        formset = self.formset_class(initial=initial, form_kwargs={'sla': sla})
        self._annotate_qual_fields(formset, quals)

        return render(request, self.template_name, {
            'sla': sla,
            'formset': formset,
            'quals': quals,
        })

    def post(self, request, sla_id):
        sla = get_object_or_404(SLA, pk=sla_id)
        if sla.num_tranches == 99:
            return redirect("add_sla_learners", sla_id=sla.id)

        quals = self._get_qualifications(sla)
        formset = self.formset_class(request.POST, form_kwargs={'sla': sla})

        if formset.is_valid():
            # Remove old entries
            BillingHistoryItem.objects.filter(billing_history__sla=sla).delete()
            BillingHistory.objects.filter(sla=sla).delete()

            # Create new tranches and items
            for form in formset:
                cd = form.cleaned_data
                bh = BillingHistory.objects.create(
                    sla=sla,
                    invoice_type=cd['invoice_type'],
                    due_date=cd.get('due_date'),
                    amount=cd['total_amount']
                )
                for qual in quals:
                    amt = cd.get(f"qual_{qual.id}") or 0
                    if amt > 0:
                        BillingHistoryItem.objects.create(
                            billing_history=bh,
                            sla_qualification=qual,
                            amount=amt
                        )
            return redirect("add_sla_learners", sla_id=sla.id)

        # On error, re-annotate and re-render
        self._annotate_qual_fields(formset, quals)
        return render(request, self.template_name, {
            'sla': sla,
            'formset': formset,
            'quals': quals,
        })

class SLAWizardStep4View(LoginRequiredMixin, UserPassesTestMixin, View):
    template_name = "core/finance/sla_add_step4.html"

    def test_func(self):
        return self.request.user.is_staff

    def get(self, request, sla_id):
        sla   = get_object_or_404(SLA, pk=sla_id)
        quals = SLA_Qualifications.objects.filter(sla=sla).select_related('service')

        qual_forms = []
        for qual in quals:
            existing_ids = LearnerQualification.objects.filter(
                sla_qualification=qual, status='active'
            ).values_list('learner_id', flat=True)
            form = QualificationLearnersForm(
                prefix=f"qual_{qual.id}",
                sla_qual=qual,
                initial={'learners': existing_ids}
            )
            qual_forms.append((qual, form))

        # Grab the media from the first form (for CSS/JS injection)
        media = qual_forms[0][1].media if qual_forms else None

        return render(request, self.template_name, {
            'sla': sla,
            'qual_forms': qual_forms,
            'media': media,
        })

    def post(self, request, sla_id):
        sla   = get_object_or_404(SLA, pk=sla_id)
        quals = SLA_Qualifications.objects.filter(sla=sla).select_related('service')

        qual_forms = []
        all_valid = True

        for qual in quals:
            form = QualificationLearnersForm(
                request.POST,
                prefix=f"qual_{qual.id}",
                sla_qual=qual
            )
            qual_forms.append((qual, form))
            if not form.is_valid():
                all_valid = False

        media = qual_forms[0][1].media if qual_forms else None

        if all_valid:
            # clear out any old associations for this SLA
            LearnerQualification.objects.filter(
                sla_qualification__sla=sla
            ).delete()

            # create fresh ones
            for qual, form in qual_forms:
                for learner in form.cleaned_data['learners']:
                    LearnerQualification.objects.create(
                        sla_qualification=qual,
                        learner=learner,
                        status='active'
                    )

            return redirect('finance_dashboard')

        # if errors, re-render with forms + errors (and media)
        return render(request, self.template_name, {
            'sla': sla,
            'qual_forms': qual_forms,
            'media': media,
        })
class LearnerPortalView(LoginRequiredMixin, TemplateView):
    template_name = "portal/welcome.html"
    login_url = 'login'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        user = self.request.user

        # Grab the learner profile linked to this User
        try:
            learner = Learner.objects.get(user=user)
        except Learner.DoesNotExist:
            learner = None
        context['learner'] = learner

        # NEW: If learner exists, add current Group and schedule details.
        if learner:
            # Assuming your Learner model is related to Group through a reverse relation named 'group_set'
            current_groups = learner.group_set.all()
            context['current_groups'] = current_groups

            # Example: Get weekly schedules for each group (adjust if your relationship is different)
            from datetime import date
            today = date.today()
            schedules = {}
            for group in current_groups:
                # This example assumes you can get schedules via a LearnerQualification relation;
                # adjust the query as needed.
                weekly = WeeklySchedule.objects.filter(learner_qualification__learner=learner)
                schedules[group.id] = weekly
            context['schedules'] = schedules

        # Preserve existing exam button logic
        show_exam = LearnerQualification.objects.filter(
            learner__user=user,
            status='active',
            sla_qualification__service__requires_summative_exam=True
        ).exists()
        context['show_exam'] = show_exam

        return context
from django.views.generic.edit import CreateView
from .models import Learner

class AddLearnerForm(RolePermissionRequiredMixin, forms.ModelForm):
    GENDER_CHOICES = [
        ('Male', 'Male'),
        ('Female', 'Female'),
        ('Other', 'Other'),
    ]
    EQUITY_CHOICES = [
        ('African', 'African'),
        ('Coloured', 'Coloured'),
        ('Indian', 'Indian'),
        ('White', 'White'),
        ('Other', 'Other'),
    ]
    Gender = forms.ChoiceField(choices=GENDER_CHOICES)
    Equity = forms.ChoiceField(choices=EQUITY_CHOICES)
    expected_clock_in = forms.TimeField(
        initial="08:00",
        widget=forms.TimeInput(format='%H:%M', attrs={'type': 'time'})
    )
    expected_clock_out = forms.TimeField(
        initial="16:30",
        widget=forms.TimeInput(format='%H:%M', attrs={'type': 'time'})
    )

    class Meta:
        model = Learner
        fields = [
            'id', 'FirstName', 'Surname', 'IDNumber', 'Gender', 'Equity',
            'EmailAddress', 'UserID', 'expected_clock_in', 'expected_clock_out'
        ]

class AddLearnerView(RolePermissionRequiredMixin, CreateView):
    model = Learner
    form_class = AddLearnerForm
    template_name = 'core/add_learner.html'
    success_url = reverse_lazy('learner_list')

from django.views.generic.edit import DeleteView

class DeleteLearnerView(RolePermissionRequiredMixin, DeleteView):
    model = Learner
    template_name = 'core/delete_learner_confirm.html'
    success_url = reverse_lazy('learner_list')
    pk_url_kwarg = 'learner_id'

from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from django.urls import reverse_lazy
from .models import ProjectPlan, SessionDate, Venue, VenueBooking, Group, Module
from .forms import ProjectPlanForm, SessionDateForm, VenueForm, VenueBookingForm
from django.contrib import messages
from django.db.models import Q

class ProjectPlanListView(RolePermissionRequiredMixin, ListView):
    model = ProjectPlan
    template_name = 'core/projectplan_list.html'
    context_object_name = 'projectplans'
    paginate_by = 10  
    
    def get_queryset(self):
        queryset = super().get_queryset()
        search = self.request.GET.get('search', '')
        if search:
            queryset = queryset.filter(
                Q(group__name__icontains=search) |
                Q(module__name__icontains=search)
            )
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['is_paginated'] = context['page_obj'].has_other_pages()
        return context

class ProjectPlanCreateView(RolePermissionRequiredMixin, CreateView):
    model = ProjectPlan
    form_class = ProjectPlanForm
    template_name = 'core/projectplan_form.html'
    success_url = reverse_lazy('projectplan_list')

class ProjectPlanUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = ProjectPlan
    form_class = ProjectPlanForm
    template_name = 'core/projectplan_form.html'
    success_url = reverse_lazy('projectplan_list')

class ProjectPlanDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = ProjectPlan
    template_name = 'core/projectplan_confirm_delete.html'
    success_url = reverse_lazy('projectplan_list')

class SessionDateListView(RolePermissionRequiredMixin, ListView):
    model = SessionDate
    template_name = 'core/sessiondate_list.html'
    context_object_name = 'sessiondates'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['is_paginated'] = context['page_obj'].has_other_pages()
        return context

class SessionDateCreateView(RolePermissionRequiredMixin, CreateView):
    model = SessionDate
    form_class = SessionDateForm
    template_name = 'core/sessiondate_form.html'
    success_url = reverse_lazy('sessiondate_list')

class SessionDateUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = SessionDate
    form_class = SessionDateForm
    template_name = 'core/sessiondate_form.html'
    success_url = reverse_lazy('sessiondate_list')

class SessionDateDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = SessionDate
    template_name = 'core/sessiondate_confirm_delete.html'
    success_url = reverse_lazy('sessiondate_list')

class VenueListView(RolePermissionRequiredMixin, ListView):
    model = Venue
    template_name = 'core/venue_list.html'
    context_object_name = 'venues'

class VenueCreateView(RolePermissionRequiredMixin, CreateView):
    model = Venue
    form_class = VenueForm
    template_name = 'core/venue_form.html'
    success_url = reverse_lazy('venue_list')

from django.shortcuts import render
from .models import Venue, VenueBooking
from datetime import date, timedelta


class VenueBookingListView(RolePermissionRequiredMixin, ListView):
    model = VenueBooking
    template_name = 'core/venuebooking_list.html'
    context_object_name = 'venuebookings'

    def get_queryset(self):
        queryset = super().get_queryset().select_related(
            'venue',
            'session_date__project_plan__group',
            'session_date__project_plan__module',
            'facilitator__learner',
            'user'
        )
        
        # **KEY UPDATE**: Handle status filtering and exclude cancelled by default
        status_filter = self.request.GET.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        else:
            # Default: exclude cancelled bookings
            queryset = queryset.exclude(status='cancelled')
        
        date_str = self.request.GET.get('date')
        if date_str:
            try:
                from datetime import datetime
                selected = datetime.strptime(date_str, "%Y-%m-%d").date()
            except Exception:
                selected = date.today()
        else:
            selected = date.today()
        queryset = queryset.filter(start_datetime__date=selected)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        today = date.today()
        tomorrow = today + timedelta(days=1)
        next7 = today + timedelta(days=7)
        selected_date = self.request.GET.get('date')
        status_filter = self.request.GET.get('status')
        
        if selected_date:
            try:
                from datetime import datetime
                selected_date = datetime.strptime(selected_date, "%Y-%m-%d").date()
            except Exception:
                selected_date = today
        else:
            selected_date = today

        context['today'] = today
        context['tomorrow'] = tomorrow
        context['next7'] = next7
        context['selected_date'] = selected_date
        context['status_filter'] = status_filter  # Add status filter to context
        
        # Status choices for filtering
        context['status_choices'] = VenueBooking.STATUS_CHOICES

        # Enforced venue color logic
        venue_color_map = {}
        for booking in context['venuebookings']:
            venue_name = booking.venue.name
            if venue_name not in venue_color_map:
                if "virtual session" in venue_name.lower():
                    venue_color_map[venue_name] = "#3A7AFE"
                else:
                    venue_color_map[venue_name] = get_venue_color(venue_name)
        context['venue_color_map'] = venue_color_map

        # --- UNBOOKED VENUES LOGIC ---
        venues = Venue.objects.exclude(name__iexact="mafadi").order_by('name')
        venues = Venue.objects.exclude(name__iexact="hogsback").exclude(name__iexact="mafadi").order_by('name')
        context['venues'] = venues
        booked_venue_ids = set(
            b.venue_id for b in context['venuebookings']
            if b.venue and b.venue.name.lower() != "virtual session"
        )
        unbooked_venues = [
            v for v in venues
            if "virtual session" not in v.name.lower() and v.id not in booked_venue_ids
        ]
        context['unbooked_venues'] = unbooked_venues

        # --- Group bookings by combined_booking_id for display ---
        bookings = list(context['venuebookings'])
        grouped = {}
        for b in bookings:
            key = b.combined_booking_id or f"single-{b.id}"
            grouped.setdefault(key, []).append(b)
        context['grouped_bookings'] = list(grouped.values())

        # --- Separate physical and virtual bookings for template display ---
        physical_bookings = []
        virtual_bookings = []
        for booking in context['venuebookings']:
            if booking.venue.name.lower() == "virtual session":
                virtual_bookings.append(booking)
            else:
                physical_bookings.append(booking)
        context['physical_bookings'] = physical_bookings
        context['virtual_bookings'] = virtual_bookings
        context['virtual_count'] = len(virtual_bookings)

        return context

from django.views.generic.edit import CreateView
from .models import VenueBooking, SessionDate, ProjectPlan
from .forms import VenueBookingForm
from django.urls import reverse_lazy
from django.shortcuts import redirect
from django.contrib import messages
import uuid

from django.views.generic.edit import CreateView
from .models import VenueBooking, SessionDate, ProjectPlan
from .forms import VenueBookingForm
from django.urls import reverse_lazy
from django.shortcuts import redirect
from django.contrib import messages
import uuid

class VenueBookingCreateView(RolePermissionRequiredMixin, CreateView):
    model = VenueBooking
    form_class = VenueBookingForm
    template_name = 'core/venuebooking_form.html'
    success_url = reverse_lazy('venuebooking_list')

    def get_initial(self):
        initial = super().get_initial()
        date_str = self.request.GET.get('date')
        project = self.request.GET.get('project')
        session = self.request.GET.get('session')
        if date_str:
            try:
                from datetime import datetime, time
                date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                initial['start_datetime'] = datetime.combine(date_obj, time(8, 0))
                initial['end_datetime'] = datetime.combine(date_obj, time(17, 0))
            except Exception:
                pass
        if session:
            try:
                session_obj = SessionDate.objects.select_related('project_plan').get(pk=session)
                initial['session_date'] = session_obj.id
                initial['project_plan'] = session_obj.project_plan.id
                from .views import get_default_num_learners_for_session
                default_learners = get_default_num_learners_for_session(session_obj.id)
                if default_learners is not None:
                    initial['num_learners'] = default_learners
            except Exception:
                pass
        elif project:
            try:
                initial['project_plan'] = int(project)
                first_session = SessionDate.objects.filter(project_plan_id=project).order_by('start_date').first()
                if first_session:
                    initial['session_date'] = first_session.id
                    from .views import get_default_num_learners_for_session
                    default_learners = get_default_num_learners_for_session(first_session.id)
                    if default_learners is not None:
                        initial['num_learners'] = default_learners
            except Exception:
                pass
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        session_date_filter = self.request.GET.get('session_date_filter')
        project_plan_filter = self.request.GET.get('project_plan_filter')
        venue = self.request.GET.get('venue')

        session_dates = SessionDate.objects.select_related('project_plan', 'project_plan__group', 'project_plan__module')
        if session_date_filter:
            session_dates = session_dates.filter(start_date=session_date_filter)
        if project_plan_filter:
            # Fix: Search within session data instead of project_plan__name
            session_dates = session_dates.filter(
                Q(project_plan__group__name__icontains=project_plan_filter) |
                Q(project_plan__module__name__icontains=project_plan_filter) |
                Q(project_plan__group__projectcode__icontains=project_plan_filter)
            )
        if venue:
            # Filter by venues that have bookings for these sessions
            session_dates = session_dates.filter(
                venuebooking__venue__name__icontains=venue
            ).distinct()
        
        context['session_dates'] = session_dates.distinct().order_by('-start_date')[:50]
        context['project_plans'] = {pp.id: str(pp) for pp in ProjectPlan.objects.all().order_by('id')}
        return context

    def form_valid(self, form):
        # Handle combined booking for multiple sessions
        session_dates = form.cleaned_data.get('session_dates')
        if session_dates and len(session_dates) > 1:
            combined_id = str(uuid.uuid4())
            for session in session_dates:
                VenueBooking.objects.create(
                    session_date=session,
                    venue=form.cleaned_data['venue'],
                    start_datetime=form.cleaned_data['start_datetime'],
                    end_datetime=form.cleaned_data['end_datetime'],
                    booking_purpose=form.cleaned_data['booking_purpose'],
                    facilitator=form.cleaned_data['facilitator'],
                    status=form.cleaned_data['status'],
                    num_learners=form.cleaned_data['num_learners'],
                    num_learners_lunch=form.cleaned_data['num_learners_lunch'],
                    combined_booking_id=combined_id,
                    user=self.request.user  # Ensure the booking is linked to the current user
                )
            messages.success(self.request, "Venue booked for multiple sessions.")
            return redirect(self.success_url)
        # Single booking fallback
        session_date = form.cleaned_data.get('session_date')
        if not session_date:
            session_id = self.request.GET.get('session')
            if session_id:
                try:
                    form.instance.session_date = SessionDate.objects.get(pk=session_id)
                except SessionDate.DoesNotExist:
                    form.instance.session_date = None
        if not form.cleaned_data.get('num_learners') and form.instance.session_date:
            from .views import get_default_num_learners_for_session
            default_learners = get_default_num_learners_for_session(form.instance.session_date.id)
            if default_learners is not None:
                form.instance.num_learners = default_learners
        form.instance.user = self.request.user  # Ensure the booking is linked to the current user
        return super().form_valid(form)

class VenueBookingUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = VenueBooking
    form_class = VenueBookingForm
    template_name = 'core/venuebooking_form.html'
    success_url = reverse_lazy('venuebooking_list')

    def form_valid(self, form):
        form.instance.status = 'rescheduled'
        # If num_learners is not set, try to set it from helper
        if not form.cleaned_data.get('num_learners') and form.instance.session_date:
            default_learners = get_default_num_learners_for_session(form.instance.session_date.id)
            if default_learners is not None:
                form.instance.num_learners = default_learners
        return super().form_valid(form)


class VenueBookingDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = VenueBooking
    template_name = 'core/venuebooking_confirm_delete.html'
    success_url = reverse_lazy('venuebooking_calendar')  # FIXED typo

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.object.status = 'cancelled'
        self.object.save()
        messages.success(request, "Booking cancelled (status updated).")
        return redirect(self.success_url)

    def get(self, request, *args, **kwargs):
        booking = get_object_or_404(VenueBooking, pk=kwargs['pk'])
        return render(request, self.template_name, {'object': booking})

    def post(self, request, *args, **kwargs):
        booking = get_object_or_404(VenueBooking, pk=kwargs['pk'])
        booking.status = 'cancelled'
        booking.save()
        messages.success(request, "Booking cancelled (status updated).")
        return redirect(self.success_url)

# Add this class after the VenueBookingModalFormView class (around line 5950)

class CancelVenueBookingView(RolePermissionRequiredMixin, View):
    """
    AJAX view to handle venue booking cancellation with email notifications
    """
    
    def post(self, request):
        booking_id = request.POST.get('booking_id')
        
        if not booking_id:
            return JsonResponse({'success': False, 'message': 'No booking ID provided'})
        
        try:
            booking = get_object_or_404(VenueBooking, pk=booking_id)
            
            # Check if user has permission to cancel
            if booking.user != request.user and not request.user.is_staff:
                return JsonResponse({'success': False, 'message': 'You do not have permission to cancel this booking'})
            
            # Store original booking user for email notification
            original_booking_user = booking.user
            cancelled_by_user = request.user
            
            # Get booking details before cancellation for email
            venue_name = booking.venue.name
            session_info = str(booking.session_date) if booking.session_date else "N/A"
            start_datetime = booking.start_datetime
            end_datetime = booking.end_datetime
            booking_purpose = booking.booking_purpose
            
            # Cancel the booking
            booking.status = 'cancelled'
            booking.save()
            
            # If this is a combined booking, cancel all related bookings
            if hasattr(booking, 'combined_booking_id') and booking.combined_booking_id:
                related_bookings = VenueBooking.objects.filter(
                    combined_booking_id=booking.combined_booking_id
                ).exclude(id=booking.id)
                
                for related_booking in related_bookings:
                    related_booking.status = 'cancelled'
                    related_booking.save()
            
            # Send email notification
            self.send_cancellation_email(
                venue_name=venue_name,
                session_info=session_info,
                start_datetime=start_datetime,
                end_datetime=end_datetime,
                booking_purpose=booking_purpose,
                original_user=original_booking_user,
                cancelled_by_user=cancelled_by_user
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Booking cancelled successfully. Notification email sent.'
            })
            
        except Exception as e:
            logger.error(f"Error cancelling booking {booking_id}: {str(e)}")
            return JsonResponse({'success': False, 'message': f'Error cancelling booking: {str(e)}'})
    
    def send_cancellation_email(self, venue_name, session_info, start_datetime, end_datetime, 
                               booking_purpose, original_user, cancelled_by_user):
        """Send email notification about booking cancellation"""
        try:
            # Determine recipient email
            recipient_email = None
            if original_user and original_user.email:
                recipient_email = original_user.email
            
            if not recipient_email:
                logger.warning("No email address found for booking cancellation notification")
                return
            
            # Prepare email content
            subject = f"Venue Booking Cancelled - {venue_name}"
            
            # Email body
            message = f"""
Dear {original_user.get_full_name() if original_user else 'User'},

Your venue booking has been cancelled.

Booking Details:
- Venue: {venue_name}
- Session: {session_info}
- Date: {start_datetime.strftime('%B %d, %Y')}
- Time: {start_datetime.strftime('%I:%M %p')} - {end_datetime.strftime('%I:%M %p')}
- Purpose: {booking_purpose}

Cancelled by: {cancelled_by_user.get_full_name() if cancelled_by_user.get_full_name() else cancelled_by_user.username}
Cancelled on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}

If you have any questions about this cancellation, please contact the person who cancelled the booking or your administrator.

Best regards,
The Learning Organisation
            """
            
            # Send email
            send_mail(
                subject=subject,
                message=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[recipient_email],
                fail_silently=True,
            )
            
            logger.info(f"Cancellation email sent to {recipient_email}")
            
        except Exception as e:
            # Log the error but don't fail the cancellation
            logger.error(f"Failed to send cancellation email: {str(e)}")


class ProjectPlanDetailView(RolePermissionRequiredMixin, DetailView):
    model = ProjectPlan
    template_name = "core/projectplan_detail.html"
    context_object_name = "projectplan"

import pandas as pd
import csv
import io
from django.db import transaction

def upload_group_excel(request):
    import pandas as pd
    import csv
    import io
    from django.db import transaction

    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "No file selected.")
            return redirect('upload_group_excel')

        errors = []
        created, updated = 0, 0

        def process_row(row, index):
            # Always use normalized keys (lowercase, underscores)
            group_id = row.get('id') or row.get('ID')
            name = row.get('title') or row.get('name')
            seta_name = row.get('seta')
            start_date = row.get('startdate') or row.get('start_date')
            end_date = row.get('enddate') or row.get('end_date')
            projectcode = row.get('projectco') or row.get('projectcode')
            service_name = row.get('service')
            project_lead_id = row.get('projectlead') or row.get('project_lead')
            etqa_admin_id = row.get('etqaadministratorid') or row.get('etqa_lead')

            # Validate and convert group_id
            if group_id is None or str(group_id).strip() == '':
                raise Exception("Missing group id")
            try:
                group_id = int(str(group_id).strip())
            except Exception:
                raise Exception(f"Invalid group id: {group_id}")

            # Clean up other fields
            name = str(name).strip() if name else None
            seta_name = str(seta_name).strip() if seta_name else None
            projectcode = str(projectcode).strip() if projectcode else None
            service_name = str(service_name).strip() if service_name else None
            project_lead_id = str(project_lead_id).strip() if project_lead_id else None
            etqa_admin_id = str(etqa_admin_id).strip() if etqa_admin_id else None

            # Parse dates
            def parse_date(val):
                if val is None or str(val).strip() == '' or str(val).strip() == '00:00.0':
                    return None
                try:
                    return pd.to_datetime(val, errors='coerce').date()
                except Exception:
                    return None

            start_date = parse_date(start_date)
            end_date = parse_date(end_date)

            # SETA
            seta = None
            if seta_name:
                seta, _ = SETA.objects.get_or_create(name=seta_name)

            # Service
            service = None
            if service_name:
                service = Service.objects.filter(name__iexact=service_name).first()

            # Project Lead (LearnerRole)
            project_lead = None
            if project_lead_id:
                learner = Learner.objects.filter(IDNumber=project_lead_id).first()
                if learner:
                    role = Role.objects.get_or_create(name="Project")[0]
                    project_lead, _ = LearnerRole.objects.get_or_create(learner=learner, role=role)

            # ETQA Lead (LearnerRole)
            etqa_lead = None
            if etqa_admin_id:
                learner = Learner.objects.filter(IDNumber=etqa_admin_id).first()
                if learner:
                    role = Role.objects.get_or_create(name="ETQA")[0]
                    etqa_lead, _ = LearnerRole.objects.get_or_create(learner=learner, role=role)

            group_obj, created_flag = Group.objects.update_or_create(
                id=group_id,
                defaults={
                    'name': name,
                    'seta': seta,
                    'start_date': start_date,
                    'end_date': end_date,
                    'projectcode': projectcode,
                    'service': service,
                    'project_lead': project_lead,
                    'etqa_lead': etqa_lead,
                }
            )
            return created_flag

        try:
            if excel_file.name.lower().endswith('.csv'):
                decoded = io.TextIOWrapper(excel_file, encoding='utf-8')
                reader = csv.DictReader(decoded)
                with transaction.atomic():
                    for idx, row in enumerate(reader, 2):
                        # Normalize keys to lowercase and underscores!
                        row = {k.strip().lower().replace(' ', '_'): (v if v != 'NULL' else None) for k, v in row.items()}
                        try:
                            created_flag = process_row(row, idx)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx}: {e}")
            else:
                df = pd.read_excel(excel_file)
                df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns]
                with transaction.atomic():
                    for idx, row in df.iterrows():
                        # Convert row to dict and replace 'NULL' with None
                        row_dict = {k: (v if v != 'NULL' else None) for k, v in row.items()}
                        try:
                            created_flag = process_row(row_dict, idx + 2)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx + 2}: {e}")
        except Exception as e:
            messages.error(request, f"Error reading file: {e}")
            return redirect('upload_group_excel')

        if created or updated:
            messages.success(request, f"Imported {created} new groups, updated {updated} existing groups.")
        if errors:
            for error in errors[:10]:
                messages.error(request, error)
            if len(errors) > 10:
                messages.warning(request, f"... and {len(errors) - 10} more errors. Check logs for details.")
        return redirect('group_management')
    return render(request, 'core/upload_group_excel.html')

import pandas as pd
from .models import Module
from django.contrib import messages
from django.shortcuts import render, redirect
from django.db import transaction

def upload_module_excel(request):
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "No file selected.")
            return redirect('upload_module_excel')

        errors = []
        created, updated = 0, 0

        def process_row(row, index):
            module_id = row.get('id')
            module_title = row.get('module title') or row.get('name')
            project_code = row.get('projectcode') or row.get('modulecode')
            if pd.isna(module_id) or pd.isna(module_title):
                raise Exception("Missing required fields")
            module_title_str = str(module_title).strip()[:100]
            code_int = None
            if not pd.isna(project_code):
                try:
                    code_int = int(project_code)
                except Exception:
                    code_int = None
            obj, created_flag = Module.objects.update_or_create(
                id=int(module_id),
                defaults={'name': module_title_str, 'code': code_int}
            )
            return created_flag

        try:
            if excel_file.name.lower().endswith('.csv'):
                decoded = io.TextIOWrapper(excel_file, encoding='utf-8')
                reader = csv.DictReader(decoded)
                with transaction.atomic():
                    for idx, row in enumerate(reader, 2):
                        try:
                            created_flag = process_row(row, idx)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx}: {e}")
            else:
                df = pd.read_excel(excel_file)
                df.columns = [c.strip().lower() for c in df.columns]
                with transaction.atomic():
                    for idx, row in df.iterrows():
                        try:
                            created_flag = process_row(row, idx + 2)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx + 2}: {e}")
        except Exception as e:
            messages.error(request, f"Error reading file: {e}")
            return redirect('upload_module_excel')

        if created or updated:
            messages.success(request, f"Imported {created} new modules, updated {updated} existing modules.")
        if errors:
            for error in errors[:10]:
                messages.error(request, error)
            if len(errors) > 10:
                messages.warning(request, f"... and {len(errors) - 10} more errors. Check logs for details.")
        return redirect('module-list')
    return render(request, 'core/upload_module_excel.html')

def upload_project_plan_excel(request):
    import pandas as pd
    import csv
    import io
    from django.db import transaction

    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "No file selected.")
            return redirect('upload_project_plan_excel')

        errors = []
        created, updated = 0, 0

        def get_learner_role(role_id, role_name):
            if pd.isna(role_id) or role_id is None:
                return None
            try:
                return LearnerRole.objects.get(id=int(role_id))
            except (ValueError, LearnerRole.DoesNotExist):
                logger.warning(f"LearnerRole ID {role_id} for {role_name} not found.")
                return None

        def parse_date(date_value):
            if pd.isna(date_value) or str(date_value).strip() == "00:00.0":
                return None
            try:
                if isinstance(date_value, (int, float)):
                    return (datetime(1899, 12, 30) + pd.Timedelta(days=date_value)).date()
                return pd.to_datetime(date_value, errors='coerce').date()
            except Exception as e:
                logger.warning(f"Invalid date format for {date_value}: {str(e)}")
                return None

        def process_row(row, index):
            # Extract all fields
            project_plan_id = row.get('id')
            group_id = row.get('group_id')
            module_id = row.get('module_id')
            module_briefing_session_person_id = row.get('module_briefing_session_person_id')
            remedial_briefing_session_start_date = row.get('remedial_briefing_session_start_date')
            complete_formative_poe_date = row.get('complete_formative_poe_date')
            summative_qa_session_person_id = row.get('summative_qa_session_person_id')
            summative_qa_session_start_date = row.get('summative_qa_session_start_date')
            complete_summative_poe_date = row.get('complete_summative_poe_date')
            assessment_person_id = row.get('assessment_person_id')
            assessment_date = row.get('assessment_date')
            remedial_briefing_session_person_id = row.get('remedial_briefing_session_person_id')
            summative_qa_session_end_date = row.get('summative_qa_session_end_date')
            remediation_submission_date = row.get('remediation_submission_date')
            remediation_assessment_person_id = row.get('remediation_assessment_person_id')
            remediation_assessment_date = row.get('remediation_assessment_date')
            client_report_person_id = row.get('client_report_person_id')
            client_report_date = row.get('client_report_date')
            module_briefing_session_end_date = row.get('module_briefing_session_end_date')
            remedial_briefing_session_end_date = row.get('remedial_briefing_session_end_date')
            assessment_book_in_date = row.get('assessment_book_in_date')
            assessment_book_out_date = row.get('assessment_book_out_date')
            results_release_due_date = row.get('results_release_due_date')
            remediation_booked_in_date = row.get('remediation_booked_in_date')
            remediation_booked_out_date = row.get('remediation_booked_out_date')
            remediation_results_release_due_date = row.get('remediation_results_release_due_date')

            # Validate required fields
            if pd.isna(project_plan_id) or pd.isna(group_id) or pd.isna(module_id):
                raise Exception("Missing required fields: id, group_id, or module_id")

            try:
                project_plan_id = int(project_plan_id)
                group_id = int(group_id)
                module_id = int(module_id)
            except (ValueError, TypeError):
                raise Exception("Invalid ID format for id, group_id, or module_id")

            # Foreign keys
            try:
                group = Group.objects.get(id=group_id)
            except Group.DoesNotExist:
                raise Exception(f"Group ID {group_id} not found")
            try:
                module = Module.objects.get(id=module_id)
            except Module.DoesNotExist:
                raise Exception(f"Module ID {module_id} not found")

            module_briefing_session_person = get_learner_role(module_briefing_session_person_id, "module_briefing_session_person")
            summative_qa_session_person = get_learner_role(summative_qa_session_person_id, "summative_qa_session_person")
            assessment_person = get_learner_role(assessment_person_id, "assessment_person")
            remedial_briefing_session_person = get_learner_role(remedial_briefing_session_person_id, "remedial_briefing_session_person")
            remediation_assessment_person = get_learner_role(remediation_assessment_person_id, "remediation_assessment_person")
            client_report_person = get_learner_role(client_report_person_id, "client_report_person")

            # Dates
            remedial_briefing_session_start_date = parse_date(remedial_briefing_session_start_date)
            complete_formative_poe_date = parse_date(complete_formative_poe_date)
            summative_qa_session_start_date = parse_date(summative_qa_session_start_date)
            complete_summative_poe_date = parse_date(complete_summative_poe_date)
            assessment_date = parse_date(assessment_date)
            summative_qa_session_end_date = parse_date(summative_qa_session_end_date)
            remediation_submission_date = parse_date(remediation_submission_date)
            remediation_assessment_date = parse_date(remediation_assessment_date)
            client_report_date = parse_date(client_report_date)
            module_briefing_session_end_date = parse_date(module_briefing_session_end_date)
            remedial_briefing_session_end_date = parse_date(remedial_briefing_session_end_date)
            assessment_book_in_date = parse_date(assessment_book_in_date)
            assessment_book_out_date = parse_date(assessment_book_out_date)
            results_release_due_date = parse_date(results_release_due_date)
            remediation_booked_in_date = parse_date(remediation_booked_in_date)
            remediation_booked_out_date = parse_date(remediation_booked_out_date)
            remediation_results_release_due_date = parse_date(remediation_results_release_due_date)

            obj, created_flag = ProjectPlan.objects.update_or_create(
                id=project_plan_id,
                defaults={
                    'group': group,
                    'module': module,
                    'module_briefing_session_person': module_briefing_session_person,
                    'remedial_briefing_session_start_date': remedial_briefing_session_start_date,
                    'complete_formative_poe_date': complete_formative_poe_date,
                    'summative_qa_session_person': summative_qa_session_person,
                    'summative_qa_session_start_date': summative_qa_session_start_date,
                    'complete_summative_poe_date': complete_summative_poe_date,
                    'assessment_person': assessment_person,
                    'assessment_date': assessment_date,
                    'remedial_briefing_session_person': remedial_briefing_session_person,
                    'summative_qa_session_end_date': summative_qa_session_end_date,
                    'remediation_submission_date': remediation_submission_date,
                    'remediation_assessment_person': remediation_assessment_person,
                    'remediation_assessment_date': remediation_assessment_date,
                    'client_report_person': client_report_person,
                    'client_report_date': client_report_date,
                    'module_briefing_session_end_date': module_briefing_session_end_date,
                    'remedial_briefing_session_end_date': remedial_briefing_session_end_date,
                    'assessment_book_in_date': assessment_book_in_date,
                    'assessment_book_out_date': assessment_book_out_date,
                    'results_release_due_date': results_release_due_date,
                    'remediation_booked_in_date': remediation_booked_in_date,
                    'remediation_booked_out_date': remediation_booked_out_date,
                    'remediation_results_release_due_date': remediation_results_release_due_date,
                }
            )
            return created_flag

        try:
            if excel_file.name.lower().endswith('.csv'):
                decoded = io.TextIOWrapper(excel_file, encoding='utf-8')
                reader = csv.DictReader(decoded)
                with transaction.atomic():
                    for idx, row in enumerate(reader, 2):
                        try:
                            created_flag = process_row(row, idx)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx}: {e}")
            else:
                df = pd.read_excel(excel_file)
                df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns]
                if 'prefered_taining_methology' in df.columns:
                    df = df.rename(columns={'prefered_taining_methology': 'preferred_training_methodology'})
                with transaction.atomic():
                    for idx, row in df.iterrows():
                        try:
                            created_flag = process_row(row, idx + 2)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx + 2}: {e}")
        except Exception as e:
            messages.error(request, f"Error reading file: {e}")
            return redirect('upload_project_plan_excel')

        if created or updated:
            messages.success(request, f"Imported {created} new project plans, updated {updated} existing project plans.")
        if errors:
            for error in errors[:10]:
                messages.error(request, error)
            if len(errors) > 10:
                messages.warning(request, f"... and {len(errors) - 10} more errors. Check logs for details.")
        return redirect('projectplan_list')
    return render(request, 'core/upload_project_plan_excel.html')

def upload_session_date_excel(request):
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "No file selected. Please select a file to upload.")
            return redirect('upload_session_date_excel')

        errors = []
        created, updated = 0, 0

        def process_row(row, index):
            session_id = row.get('id')
            project_plan_id = row.get('project_plan_id')
            start_date = row.get('start_date')
            end_date = row.get('end_date')
            preferred_training_methodology = row.get('preferred_training_methodology')
            if pd.isna(session_id) or pd.isna(project_plan_id) or pd.isna(start_date) or pd.isna(end_date):
                raise Exception("Missing required fields (id, project_plan_id, start_date, or end_date)")
            try:
                session_id = int(session_id)
                project_plan_id = int(project_plan_id)
            except (ValueError, TypeError):
                raise Exception("Invalid ID format for session_id or project_plan_id")
            try:
                project_plan = ProjectPlan.objects.get(id=project_plan_id)
            except ProjectPlan.DoesNotExist:
                raise Exception(f"ProjectPlan ID {project_plan_id} not found")
            def parse_date(date_value):
                if pd.isna(date_value):
                    return None
                try:
                    if isinstance(date_value, (int, float)):
                        return (datetime(1899, 12, 30) + pd.Timedelta(days=date_value)).date()
                    return pd.to_datetime(date_value, errors='coerce').date()
                except Exception as e:
                    raise Exception(f"Invalid date format for {date_value}: {str(e)}")
            start_date_parsed = parse_date(start_date)
            end_date_parsed = parse_date(end_date)
            if not start_date_parsed or not end_date_parsed:
                raise Exception("Invalid start_date or end_date")
            if pd.notna(preferred_training_methodology):
                preferred_training_methodology = str(preferred_training_methodology).strip()[:50]
            else:
                preferred_training_methodology = None
            obj, created_flag = SessionDate.objects.update_or_create(
                id=session_id,
                defaults={
                    'project_plan': project_plan,
                    'start_date': start_date_parsed,
                    'end_date': end_date_parsed,
                    'preferred_training_methodology': preferred_training_methodology,
                }
            )
            return created_flag

        try:
            if excel_file.name.lower().endswith('.csv'):
                decoded = io.TextIOWrapper(excel_file, encoding='utf-8')
                reader = csv.DictReader(decoded)
                with transaction.atomic():
                    for idx, row in enumerate(reader, 2):
                        try:
                            created_flag = process_row(row, idx)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx}: {e}")
            else:
                df = pd.read_excel(excel_file)
                df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns]
                if 'prefered_taining_methology' in df.columns:
                    df = df.rename(columns={'prefered_taining_methology': 'preferred_training_methodology'})
                with transaction.atomic():
                    for idx, row in df.iterrows():
                        try:
                            created_flag = process_row(row, idx + 2)
                            if created_flag:
                                created += 1
                            else:
                                updated += 1
                        except Exception as e:
                            errors.append(f"Row {idx + 2}: {e}")
        except Exception as e:
            messages.error(request, f"Error reading file: {e}")
            return redirect('upload_session_date_excel')

        if created or updated:
            messages.success(request, f"Imported {created} new session dates, updated {updated} existing session dates.")
        if errors:
            for error in errors[:10]:
                messages.error(request, error)
            if len(errors) > 10:
                messages.warning(request, f"... and {len(errors) - 10} more errors. Check logs for details.")
        return redirect('sessiondate_list')
    return render(request, 'core/upload_session_date_excel.html')

class GroupQualificationAssignmentView(RolePermissionRequiredMixin, View):
    template_name = 'core/group_qualification_assignment.html'
    paginate_by = 10
    def get(self, request):
        groups = Group.objects.all().order_by('name')
        services = Service.objects.all().order_by('name')
        # Pagination
        paginator = Paginator(groups, self.paginate_by)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, self.template_name, {
            'groups': groups,  # Keep this for the dropdown
            'page_obj': page_obj,  # This is for the table
            'services': services
        })

    def post(self, request):
        group_id = request.POST.get('group')
        sla_qualification_ids = request.POST.getlist('sla_qualifications')
        
        if not group_id or not sla_qualification_ids:
            messages.error(request, "Please select both a group and at least one SLA qualification")
            return redirect('group_qualification_assignment')
            
        try:
            group = Group.objects.get(pk=group_id)
            for sla_qual_id in sla_qualification_ids:
                sla_qual = SLA_Qualifications.objects.get(pk=sla_qual_id)
                group.sla_qualifications.add(sla_qual)
            
            messages.success(request, f"Successfully assigned qualifications to {group.name}")
            return redirect('group_qualification_assignment')
            
        except (Group.DoesNotExist, SLA_Qualifications.DoesNotExist) as e:
            messages.error(request, f"Error assigning qualifications: {str(e)}")
            return redirect('group_qualification_assignment')


class AdobeFormListView(RolePermissionRequiredMixin, ListView):
    model = AdobeForm
    template_name = "core/adobe_form_list.html"
    context_object_name = "forms"

class AdobeFormUploadView(RolePermissionRequiredMixin, CreateView):
    model = AdobeForm
    form_class = AdobeFormUploadForm
    template_name = "core/adobe_form_upload.html"
    success_url = reverse_lazy("adobe_form_list")

class AdobeFormDownloadView(RolePermissionRequiredMixin, View):
    def get(self, request, pk):
        adobe_form = get_object_or_404(AdobeForm, pk=pk)
        return FileResponse(adobe_form.pdf.open('rb'), as_attachment=True, filename=adobe_form.pdf.name.split('/')[-1])

class AdobeFormFilledSubmitView(RolePermissionRequiredMixin, FormView):
    template_name = "core/adobe_form_submit.html"
    form_class = AdobeFormFilledSubmissionForm
    success_url = reverse_lazy("adobe_form_list")

    def dispatch(self, request, *args, **kwargs):
        self.form_obj = get_object_or_404(AdobeForm, pk=kwargs['pk'])
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['adobe_form'] = self.form_obj
        return context

    def form_valid(self, form):
        filled_pdf = form.cleaned_data['filled_pdf']
        your_email = form.cleaned_data['your_email']
        recipient_email = form.cleaned_data['recipient_email']

        # Validate required fields in PDF (AcroForms only)
        try:
            reader = PdfReader(filled_pdf)
            required_unfilled = []
            if "/AcroForm" in reader.trailer["/Root"]:
                fields = reader.get_fields()
                for name, field in fields.items():
                    if field.get("/Ff", 0) & 2:
                        value = field.get("/V", "")
                        if not value:
                            required_unfilled.append(name)
            if required_unfilled:
                from django.contrib import messages
                messages.error(self.request, f"Required fields not filled: {', '.join(required_unfilled)}")
                return self.form_invalid(form)
        except Exception as e:
            from django.contrib import messages
            messages.error(self.request, f"Could not process PDF: {e}")
            return self.form_invalid(form)

        # Send email with PDF attached
        try:
            subject = f"Filled Adobe Form: {self.form_obj.name}"
            body = f"Form submitted by {your_email}."
            email = EmailMessage(
                subject=subject,
                body=body,
                from_email=None,
                to=[recipient_email],
                cc=[your_email],
            )
            filled_pdf.seek(0)
            email.attach(filled_pdf.name, filled_pdf.read(), "application/pdf")
            email.send()
            from django.contrib import messages
            messages.success(self.request, "Form sent successfully!")
        except Exception as e:
            from django.contrib import messages
            messages.error(self.request, f"Failed to send email: {e}")
            return self.form_invalid(form)

        return super().form_valid(form)


from django.core.mail import EmailMessage
from django.conf import settings
from django.urls import reverse_lazy
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.views.generic.edit import CreateView
from .models import (
    LearnerModulePOE, POEAnnexture, AdobeFormAnnexture,
    Learner, Module, LearnerQualification
)
from .forms import POESubmissionForm
from .mixins import RolePermissionRequiredMixin
from PyPDF2 import PdfReader
import fitz

def check_signatures(pdf_file):
    """
    Checks if signature fields are filled in the PDF.
    Accepts an in-memory file (e.g., Django UploadedFile).
    Returns a list of issues found.
    """
    signature_issues = []
    pdf_file.seek(0)
    raw_data = pdf_file.read()
    signature_tag_found = any(tag in raw_data for tag in [
        b"<</Subtype/page/Type/FillSignData>>", b"/Sig", b"/Signature", b"/FillSignData"
    ])
    pdf_file.seek(0)
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for i in range(len(doc)):
        page = doc[i]
        text = page.get_text().lower()
        page_number = i + 1
        if "signature" in text and not signature_tag_found:
            signature_issues.append(f"Signature mentioned but not filled on Page {page_number}")
    doc.close()
    pdf_file.seek(0)
    return signature_issues


class POESubmissionView(RolePermissionRequiredMixin, CreateView):
    model = LearnerModulePOE
    form_class = POESubmissionForm
    template_name = 'core/poe_submission.html'
    success_url = reverse_lazy('poe_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        learner_id = self.kwargs.get('learner_id')
        module_id = self.kwargs.get('module_id')
        context['learner'] = get_object_or_404(Learner, id=learner_id)
        context['module'] = get_object_or_404(Module, id=module_id)
        context['qualification'] = get_object_or_404(
            LearnerQualification,
            learner_id=learner_id,
            module__id=module_id
        )
        context['annexture_templates'] = AdobeFormAnnexture.objects.filter(
            adobe_form__module=context['module']
        )
        return context

    def form_valid(self, form):
        # Set POE metadata
        form.instance.learner_id = self.kwargs['learner_id']
        form.instance.module_id = self.kwargs['module_id']
        form.instance.learner_qualification_id = self.kwargs['qualification_id']

        # --- PDF AcroForm and signature validation ---
        poe_file = self.request.FILES.get('poe_file')
        if poe_file:
            try:
                poe_file.seek(0)
                reader = PdfReader(poe_file)
                required_unfilled = []
                if "/AcroForm" in reader.trailer["/Root"]:
                    fields = reader.get_fields()
                    for name, field in fields.items():
                        # /Ff bit 1 (value 2) means required
                        if field.get("/Ff", 0) & 2:
                            value = field.get("/V", "")
                            if not value:
                                required_unfilled.append(name)
                if required_unfilled:
                    form.add_error(None, "Required PDF fields not filled: " + ", ".join(required_unfilled))
                    return self.form_invalid(form)
                # --- Signature check ---
                signature_issues = check_signatures(poe_file)
                if signature_issues:
                    form.add_error(None, "Signature issues: " + "; ".join(signature_issues))
                    return self.form_invalid(form)
            except Exception as e:
                form.add_error(None, f"Could not process PDF: {e}")
                return self.form_invalid(form)

        # Save POE
        self.object = form.save()

        # Handle annextures
        annexture_templates = AdobeFormAnnexture.objects.filter(
            adobe_form__module_id=self.kwargs['module_id']
        )
        for i, template in enumerate(annexture_templates):
            annexture_file = self.request.FILES.get(f'annexture_{i}')
            if not annexture_file and template.required:
                form.add_error(None, f'Annexture {template.name} is required')
                return self.form_invalid(form)
            if annexture_file:
                POEAnnexture.objects.create(
                    poe=self.object,
                    annexture_template=template,
                    file=annexture_file
                )

        # Send email notification
        self.send_notification_email()

        messages.success(self.request, 'POE submitted successfully')
        return redirect(self.get_success_url())

    def send_notification_email(self):
        poe = self.object
        facilitator_email = (
            poe.module.project_plans.first().assessment_person.learner.EmailAddress
            if poe.module.project_plans.exists() and poe.module.project_plans.first().assessment_person
            else None
        )
        if facilitator_email:
            subject = f'New POE Submission - {poe.module.name}'
            message = f'''
            A new POE has been submitted:

            Learner: {poe.learner}
            Module: {poe.module.name}
            Submitted: {poe.submission_date}

            Please login to review the submission.
            '''
            email = EmailMessage(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [facilitator_email]
            )
            email.send()
            
class POEListView(RolePermissionRequiredMixin, ListView):
    model = LearnerModulePOE
    template_name = 'core/poe_list.html'
    context_object_name = 'poes'

    def get_queryset(self):
        learner_id = self.kwargs.get('learner_id')
        return LearnerModulePOE.objects.filter(learner_id=learner_id)
from django.contrib import messages
from django.shortcuts import redirect
from django.views.generic import CreateView
from .forms import FacilitatorPOEUploadForm

class FacilitatorPOEUploadView(RolePermissionRequiredMixin, CreateView):
    model = AdobeForm
    form_class = FacilitatorPOEUploadForm
    template_name = 'core/facilitator_poe_upload.html'
    success_url = reverse_lazy('facilitator_poe_list')

    def form_valid(self, form):
        response = super().form_valid(form)
        
        # Create annextures
        for idx, name in enumerate(form.cleaned_data['annextures']):
            AdobeFormAnnexture.objects.create(
                adobe_form=self.object,
                name=name,
                required=True,
                order=idx
            )
        
        # Link to module
        module_id = self.kwargs.get('module_id')
        if module_id:
            module = Module.objects.get(id=module_id)
            module.adobe_forms.add(self.object)

        messages.success(self.request, 'POE template uploaded successfully')
        return response

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        module_id = self.kwargs.get('module_id')
        if module_id:
            context['module'] = Module.objects.get(id=module_id)
        return context



# Attendance analysis helper functions (already defined elsewhere, included for reference)
def get_clock_in_category(clock_in_time, expected_clock_in):
    if not clock_in_time:
        return 'missing'
    if not expected_clock_in:
        return 'N/A'
    early_cutoff = time(7, 35)
    on_time_start = time(7, 35)
    on_time_end = time(8, 5)
    if clock_in_time < early_cutoff:
        return 'early'
    elif on_time_start <= clock_in_time <= on_time_end:
        return 'on_time'
    else:
        return 'late'

def get_expected_clock_out(clock_in_time, custom_clock_out=None):
    if custom_clock_out:
        return custom_clock_out
    if not clock_in_time:
        return time(16, 25)  # Default if no clock-in
    if clock_in_time <= time(7, 30):
        return time(15, 55)  # 3:55 PM
    return time(16, 25)  # 4:25 PM

def get_clock_out_category(clock_out_time, expected_clock_out):
    if not clock_out_time:
        return 'missing'
    if not expected_clock_out:
        return 'N/A'
    grace_period = timedelta(minutes=10)
    expected_datetime = datetime.combine(date.today(), expected_clock_out)
    clock_out_datetime = datetime.combine(date.today(), clock_out_time)
    if clock_out_datetime < expected_datetime - grace_period:
        return 'early_leave'
    elif clock_out_datetime <= expected_datetime + grace_period:
        return 'on_time'
    else:
        return 'overtime'

from itertools import groupby
from operator import attrgetter
from datetime import datetime, date, time, timedelta
from django.db.models import Q
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import TemplateView
from django.core.cache import cache
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from .models import (
    Learner, Group, ProjectPlan, Fingerprint, SessionDate,
    VenueBooking, WeeklySchedule, LearnerQualification, SLA
)

# Attendance analysis helper functions (unchanged)
def get_clock_in_category(clock_in_time, expected_clock_in):
    if not clock_in_time:
        return 'missing'
    if not expected_clock_in:
        return 'N/A'
    early_cutoff = time(7, 35)
    on_time_start = time(7, 35)
    on_time_end = time(8, 5)
    if clock_in_time < early_cutoff:
        return 'early'
    elif on_time_start <= clock_in_time <= on_time_end:
        return 'on_time'
    else:
        return 'late'

def get_expected_clock_out(clock_in_time, custom_clock_out=None):
    if custom_clock_out:
        return custom_clock_out
    if not clock_in_time:
        return time(16, 25)  # Default if no clock-in
    if clock_in_time <= time(7, 30):
        return time(15, 55)  # 3:55 PM
    return time(16, 25)  # 4:25 PM

def get_clock_out_category(clock_out_time, expected_clock_out):
    if not clock_out_time:
        return 'missing'
    if not expected_clock_out:
        return 'N/A'
    grace_period = timedelta(minutes=10)
    expected_datetime = datetime.combine(date.today(), expected_clock_out)
    clock_out_datetime = datetime.combine(date.today(), clock_out_time)
    if clock_out_datetime < expected_datetime - grace_period:
        return 'early_leave'
    elif clock_out_datetime <= expected_datetime + grace_period:
        return 'on_time'
    else:
        return 'overtime'


from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import TemplateView
from django.contrib import messages
from datetime import date, timedelta
from .models import (
    Learner, LearnerQualification, LearnerModulePOE, Group, SessionDate, Fingerprint, LIF
)

class LearnerHomeView(LoginRequiredMixin, TemplateView):
    template_name = "portal/learner_home.html"
    login_url = 'login'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            learner = self.request.user.learner_profile
        except Learner.DoesNotExist:
            messages.error(self.request, "No learner profile found.")
            return context

        # Get active qualifications with related data
        active_quals = LearnerQualification.objects.filter(
            learner=learner,
            status='active'
        ).select_related(
            'sla_qualification__service',
            'sla_qualification__sla'
        )

        # Recent POE submissions
        recent_poes = LearnerModulePOE.objects.filter(
            learner=learner
        ).select_related('module').order_by('-submission_date')[:5]

        # Get groups through SLA qualifications
        groups = Group.objects.filter(
            sla_qualifications__learnerqualification__learner=learner,
            sla_qualifications__learnerqualification__status='active'
        ).distinct()

        # Get upcoming sessions
        today = date.today()
        upcoming_sessions = SessionDate.objects.filter(
            project_plan__group__in=groups,
            start_date__gte=today
        ).select_related(
            'project_plan__group',
            'project_plan__module'
        ).order_by('start_date')[:5]

        # Get recent attendance records
        recent_attendance = []
        if learner.UserID:
            recent_attendance = Fingerprint.objects.filter(
                user_id=learner.UserID,
                date__gte=today - timedelta(days=7)
            ).order_by('-date', 'time')

        # Check if LIF form is missing for this learner
        lif_missing = not hasattr(learner, 'lif_form')

        context.update({
            'learner_details': {
                'full_name': f"{learner.FirstName} {learner.Surname}",
                'email': learner.EmailAddress,
                'gender': learner.Gender,
                'equity': learner.Equity
            },
            'active_qualifications': active_quals,
            'recent_poes': recent_poes,
            'upcoming_sessions': upcoming_sessions,
            'groups': groups,
            'recent_attendance': recent_attendance,
            'lif_missing': lif_missing,
        })

        return context
    

from django.shortcuts import render, redirect, get_object_or_404
from django.views import View
from django.views.generic import ListView, DetailView, CreateView, UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from .models import LearnerModulePOE, LearnerRole, LearnerQualification, Module
from .forms import POESubmissionForm, FacilitatorPOEFeedbackForm
from django.core.mail import send_mail
from django.conf import settings
from django.urls import reverse_lazy

# ...existing code...

class FacilitatorPOEListView(LoginRequiredMixin, UserPassesTestMixin, ListView):
    model = LearnerModulePOE
    template_name = 'core/facilitator_poe_list.html'
    context_object_name = 'poes'

    def test_func(self):
        # Check if the user is a facilitator
        return LearnerRole.objects.filter(learner=self.request.user.learner_profile, role__name='Facilitator').exists()

    def get_queryset(self):
        # Get the facilitator's learner role
        facilitator_role = LearnerRole.objects.get(learner=self.request.user.learner_profile, role__name='Facilitator')
        # Filter POEs that are under review and assigned to the facilitator
        return LearnerModulePOE.objects.filter(facilitator=facilitator_role, status='under_review')

class FacilitatorPOEFeedbackView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = LearnerModulePOE
    form_class = FacilitatorPOEFeedbackForm
    template_name = 'core/facilitator_poe_feedback.html'
    pk_url_kwarg = 'poe_id'
    success_url = reverse_lazy('facilitator_poe_list')

    def test_func(self):
        # Check if the user is a facilitator
        return LearnerRole.objects.filter(learner=self.request.user.learner_profile, role__name='Facilitator').exists()

    def form_valid(self, form):
        form.instance.facilitator = LearnerRole.objects.get(learner=self.request.user.learner_profile, role__name='Facilitator')
        form.instance.review_date = timezone.now()
        form.instance.status = form.cleaned_data['status']  # Set status based on form input
        form.save()

        # Send email to learner with feedback
        send_mail(
            'POE Feedback',
            f'Your POE for {form.instance.module} has been reviewed. Status: {form.instance.status}. Feedback: {form.instance.feedback}',
            settings.DEFAULT_FROM_EMAIL,
            [form.instance.learner.EmailAddress],
            fail_silently=True,
        )
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('facilitator_poe_list')
    
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView, CreateView, View
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.contrib import messages
from .models import ModulePOETemplate, ModulePOEAnnexture, Module
from .forms import ModulePOETemplateForm, ModulePOEAnnextureForm

from django.core.paginator import Paginator
from django.db.models import Q

class AdminPOETemplateUploadView(LoginRequiredMixin, CreateView):
    model = ModulePOETemplate
    form_class = ModulePOETemplateForm
    template_name = 'core/poe_template_form.html'
    success_url = reverse_lazy('poe_template_list')

    def get(self, request, *args, **kwargs):
        if request.path.endswith('/add/'):
            return super().get(request, *args, **kwargs)
        else:
            # --- Filtering ---
            templates = ModulePOETemplate.objects.select_related('module', 'uploaded_by').order_by('-uploaded_at')
            module_id = request.GET.get('module')
            search = request.GET.get('search', '')
            if module_id:
                templates = templates.filter(module_id=module_id)
            if search:
                templates = templates.filter(
                    Q(module__name__icontains=search) |
                    Q(uploaded_by__username__icontains=search)
                )
            # --- Pagination ---
            paginator = Paginator(templates, 10)
            page_number = request.GET.get('page')
            page_obj = paginator.get_page(page_number)
            # --- Modules for filter dropdown ---
            modules = Module.objects.all().order_by('name')
            filters = {
                'module': module_id or '',
                'search': search or '',
            }
            return render(
                request,
                'core/poe_template_list.html',
                {
                    'templates': page_obj,
                    'modules': modules,
                    'filters': filters,
                    'is_paginated': page_obj.has_other_pages(),
                    'page_obj': page_obj,
                }
            )

    def form_valid(self, form):
        form.instance.uploaded_by = self.request.user
        response = super().form_valid(form)
        messages.success(self.request, 'POE template uploaded successfully.')
        return response
    
class AdminPOEAnnextureConfigView(LoginRequiredMixin, View):
    template_name = 'core/annexture_config.html'

    def get(self, request, template_id):
        template = get_object_or_404(ModulePOETemplate, id=template_id)
        return render(request, self.template_name, {
            'template': template,
            'all_required': all(a.required for a in template.required_annextures.all()),
        })

    def post(self, request, template_id):
        template = get_object_or_404(ModulePOETemplate, id=template_id)
        annexture_names = request.POST.get('annexture_names', '').splitlines()
        annexture_names = [n.strip() for n in annexture_names if n.strip()]
        all_required = bool(request.POST.get('all_required'))

        # Remove old annextures not in the new list
        existing = {a.name: a for a in template.required_annextures.all()}
        for name in existing:
            if name not in annexture_names:
                existing[name].delete()

        # Add/update annextures
        for i, name in enumerate(annexture_names):
            ann, created = ModulePOEAnnexture.objects.get_or_create(
                template=template, name=name,
                defaults={'required': all_required}
            )
            if not created:
                ann.required = all_required
                ann.save()

        messages.success(request, "Annextures updated.")
        return redirect(request.path)
    

class POESubmissionView(LoginRequiredMixin, View):
    template_name = 'core/poe_submission.html'

    def get(self, request):
        try:
            learner = request.user.learner_profile
        except Learner.DoesNotExist:
            messages.error(request, "No learner profile found.")
            return redirect('learner_home')

        # 1. Get all active learner qualifications for this learner
        learner_quals = LearnerQualification.objects.filter(
            learner=learner, status='active'
        ).select_related('sla_qualification__service')

        # 2. For each qualification, get all groups linked to its sla_qualification
        group_qual_map = {}
        for lq in learner_quals:
            groups = Group.objects.filter(sla_qualifications=lq.sla_qualification)
            for group in groups:
                if group.id not in group_qual_map:
                    group_qual_map[group.id] = {'group': group, 'qualifications': []}
                group_qual_map[group.id]['qualifications'].append(lq)

        # 3. For each group, for each qualification, get modules via ProjectPlan for that group
        display_data = []
        for group_id, gq in group_qual_map.items():
            group = gq['group']
            group_entry = {
                'group': group,
                'qualifications': []
            }
            for lq in gq['qualifications']:
                modules = []
                project_plans = ProjectPlan.objects.filter(group=group)
                for plan in project_plans:
                    module = plan.module
                    template = ModulePOETemplate.objects.filter(module=module).first()
                    existing_poe = LearnerModulePOE.objects.filter(
                        learner=learner,
                        module=module,
                        learner_qualification=lq
                    ).order_by('-submission_date').first()
                    modules.append({
                        'module': module,
                        'template': template,
                        'existing_poe': existing_poe,
                        'qualification_id': lq.id,
                    })
                group_entry['qualifications'].append({
                    'qualification': lq,
                    'qualification_name': lq.sla_qualification.service.name,
                    'modules': modules
                })
            display_data.append(group_entry)

        context = {
            'learner': learner,
            'display_data': display_data,
        }
        return render(request, self.template_name, context)

    def post(self, request):
        try:
            learner = request.user.learner_profile
        except Learner.DoesNotExist:
            messages.error(request, "No learner profile found.")
            return redirect('learner_home')

        module_id = request.POST.get('module_id')
        qualification_id = request.POST.get('qualification_id')

        if not all([module_id, qualification_id]):
            messages.error(request, 'Missing required information')
            return redirect('poe_submission')

        qualification = get_object_or_404(LearnerQualification, id=qualification_id)
        module = get_object_or_404(Module, id=module_id)
        template = ModulePOETemplate.objects.filter(module=module).first()  # May be None

        missing_fields = []

        # Check main POE file
        poe_file = request.FILES.get('poe_file')
        if not poe_file:
            missing_fields.append('POE file')

        # Check required annextures if template exists
        if template:
            required_annextures = template.required_annextures.filter(required=True)
            for annexture in required_annextures:
                if f'annexture_{annexture.id}' not in request.FILES:
                    missing_fields.append(f"Annexture: {annexture.name}")

        if missing_fields:
            messages.error(
                request,
                "Please upload all required files: " + ", ".join(missing_fields)
            )
            return redirect('poe_submission')

        # --- PDF AcroForm validation ---
        try:
            poe_file.seek(0)
            reader = PdfReader(poe_file)
            required_unfilled = []
            if "/AcroForm" in reader.trailer["/Root"]:
                fields = reader.get_fields()
                for name, field in fields.items():
                    # /Ff bit 1 (value 2) means required
                    if field.get("/Ff", 0) & 2:
                        value = field.get("/V", "")
                        if not value:
                            required_unfilled.append(name)
            if required_unfilled:
                messages.error(
                    request,
                    "Required PDF fields not filled: " + ", ".join(required_unfilled)
                )
                return redirect('poe_submission')
        except Exception as e:
            messages.error(request, f"Could not process PDF: {e}")
            return redirect('poe_submission')

        # Allow resubmission: delete previous POE for this learner/module/qualification
        LearnerModulePOE.objects.filter(
            learner=learner,
            learner_qualification=qualification,
            module=module
        ).delete()

        poe = LearnerModulePOE.objects.create(
            learner=learner,
            learner_qualification=qualification,
            module=module,
            poe_file=poe_file
        )

        if template:
            for key, file in request.FILES.items():
                if key.startswith('annexture_'):
                    annexture_id = key.split('_')[1]
                    template_annexture = get_object_or_404(ModulePOEAnnexture, id=annexture_id)
                    POEAnnexture.objects.create(
                        poe=poe,
                        annexture_template=template_annexture,
                        file=file
                    )

        messages.success(request, 'POE submitted successfully')
        return redirect('poe_submission')
    
from django.views.generic import TemplateView
from django.core.mail import EmailMessage
from django.db.models import Q
from .models import (
    LearnerModulePOE, Group, Learner, LearnerQualification, Module, ProjectPlan, ModulePOETemplate
)
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.urls import reverse_lazy
from django.contrib import messages

class AdminPOESubmissionsDashboardView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    template_name = "core/admin_poe_dashboard.html"
    login_url = 'login'

    def test_func(self):
        # Only staff/admins allowed
        return self.request.user.is_staff

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Filters
        group_id = self.request.GET.get('group')
        qual_id = self.request.GET.get('qualification')
        module_id = self.request.GET.get('module')
        status = self.request.GET.get('status')
        search = self.request.GET.get('search', '').strip()

        groups = Group.objects.all().order_by('name')
        qualifications = LearnerQualification.objects.select_related('sla_qualification').all()
        modules = Module.objects.all().order_by('name')

        poes = LearnerModulePOE.objects.select_related(
            'learner', 'learner_qualification', 'module'
        ).all()

        if group_id:
            group = Group.objects.get(id=group_id)
            # Filter POEs by learners in this group (via qualification)
            group_quals = group.sla_qualifications.all()
            poes = poes.filter(learner_qualification__sla_qualification__in=group_quals)
        if qual_id:
            poes = poes.filter(learner_qualification_id=qual_id)
        if module_id:
            poes = poes.filter(module_id=module_id)
        if status:
            poes = poes.filter(status=status)
        if search:
            poes = poes.filter(
                Q(learner__FirstName__icontains=search) |
                Q(learner__Surname__icontains=search) |
                Q(learner__EmailAddress__icontains=search)
            )

        poes = poes.order_by('-submission_date')

        context.update({
            'groups': groups,
            'qualifications': qualifications,
            'modules': modules,
            'poes': poes,
            'filter_group': group_id,
            'filter_qualification': qual_id,
            'filter_module': module_id,
            'filter_status': status,
            'filter_search': search,
            'status_choices': LearnerModulePOE.STATUS_CHOICES,
        })
        return context

    def post(self, request, *args, **kwargs):
        action = request.POST.get('action')
        selected_poe_ids = request.POST.getlist('selected_poes')
        poes = LearnerModulePOE.objects.filter(id__in=selected_poe_ids).select_related(
            'learner', 'learner_qualification__sla_qualification', 'module'
        )

        if action == 'send_email':
            recipient = request.POST.get('recipient_email')
            if not recipient:
                messages.error(request, "Recipient email required.")
                return redirect(request.path)

            # Group by group and qualification for subject
            if not poes:
                messages.error(request, "No POE submissions selected.")
                return redirect(request.path)

            # Use the first POE for subject info
            first_poe = poes.first()
            group = None
            qualification = None
            # Find group via qualification
            if hasattr(first_poe.learner_qualification.sla_qualification, 'groups'):
                group = first_poe.learner_qualification.sla_qualification.groups.first()
            qualification = first_poe.learner_qualification.sla_qualification.service.name

            subject = f"POE Submission for {group.name if group else 'Group'} - {qualification}"

            # Build body
            body = "The following POE submissions are attached:\n\n"
            for poe in poes:
                learner = poe.learner
                module = poe.module.name
                submission_date = poe.submission_date.strftime("%Y-%m-%d %H:%M")
                body += f"- {learner.FirstName} {learner.Surname} | Module: {module} | Submitted: {submission_date}\n"

            email = EmailMessage(
                subject=subject,
                body=body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[recipient]
            )
            for poe in poes:
                if poe.poe_file:
                    poe.poe_file.open('rb')
                    email.attach(f"{poe.learner.FirstName}_{poe.learner.Surname}_{poe.module.name}.pdf", poe.poe_file.read())
                    poe.poe_file.close()
                for ann in poe.annextures.all():
                    ann.file.open('rb')
                    email.attach(f"{poe.learner.FirstName}_{poe.learner.Surname}_{ann.annexture_template.name}.pdf", ann.file.read())
                    ann.file.close()
            email.send()
            messages.success(request, f"Sent {poes.count()} POEs to {recipient}.")
            return redirect(request.path)

        elif action == 'assess':
            updated = 0
            for poe in poes:
                status = request.POST.get(f'status_{poe.id}')
                feedback = request.POST.get(f'feedback_{poe.id}', '').strip()
                if status == 'not_yet_competent' and not feedback:
                    messages.error(request, f"Feedback required for {poe.learner} if not yet competent.")
                    continue
                poe.status = status
                poe.feedback = feedback
                poe.review_date = timezone.now()
                poe.save()
                updated += 1

                # Notify learner by email
                if poe.learner.EmailAddress:
                    subject = f"POE Assessment Result for {poe.module.name}"
                    html_message = render_to_string('core/email_poe_assessment.html', {
                        'learner': poe.learner,
                        'module': poe.module,
                        'status': poe.get_status_display(),
                        'feedback': feedback,
                        'review_date': poe.review_date,
                    })
                    plain_message = strip_tags(html_message)
                    send_mail(
                        subject,
                        plain_message,
                        settings.DEFAULT_FROM_EMAIL,
                        [poe.learner.EmailAddress],
                        html_message=html_message,
                        fail_silently=True
                    )
            if updated:
                messages.success(request, f"Assessment saved for {updated} submission(s).")
            return redirect(request.path)

        return redirect(request.path)
    
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils import timezone
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from .models import LearnerModulePOE, Learner, Module, LearnerQualification
from django.conf import settings

def assess_poe(request, poe_id):
    poe = get_object_or_404(LearnerModulePOE, pk=poe_id)
    learner = poe.learner
    module = poe.module
    qualification = poe.learner_qualification.sla_qualification

    if request.method == 'POST':
        status = request.POST['status']
        feedback = request.POST.get('feedback', '')
        review_date = timezone.now()

        # Save assessment to LearnerModulePOE
        poe.status = status
        poe.feedback = feedback
        poe.review_date = review_date
        poe.save()

        # Send email to learner
        subject = f"POE Submission for {module.name} ({qualification.service.name}) - Assessment Result"
        message = render_to_string('core/email_poe_assessment.html', {
            'learner': learner,
            'module': module,
            'status': status.title(),
            'feedback': feedback,
            'review_date': review_date,
        })
        send_mail(
            subject,
            '',
            settings.DEFAULT_FROM_EMAIL,
            [learner.EmailAddress],
            html_message=message,
            fail_silently=True,
        )

        messages.success(request, "Assessment saved and learner notified by email.")
        return redirect('admin_poe_dashboard')

    return render(request, 'core/assess_poe.html', {'poe': poe, 'learner': learner, 'module': module})


@login_required
def learner_assessment_results(request):
    learner = request.user.learner_profile
    poes = LearnerModulePOE.objects.filter(learner=learner)
    return render(request, 'core/learner_assessment_results.html', {'poes': poes})

from django.views.generic import TemplateView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.shortcuts import redirect
from django.contrib import messages
from .models import Role, RolePermission

class RolePermissionManagementView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    template_name = "core/role_permission_management.html"

    def test_func(self):
        return self.request.user.is_staff

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        url_groups = [
            ("Dashboards", [
                ("sla_dashboard", "SLA Dashboard"),
                ("finance_dashboard", "Finance Dashboard"),
                ("admin_poe_dashboard", "PoE Dashboard"),
            ]),
            ("Attendance", [
                ("fingerprint_list", "Attendance Report"),
                ("upload_fingerprint", "Upload Fingerprint"),
                ("view_attendance", "Portal Attendance"),
                ("learner_attendance", "Learner Attendance"),
            ]),
            ("Learners", [
                ("learner_list", "Manage Learners"),
                ("add_learner", "Add Learner"),
                ("delete_learner", "Delete Learner"),
                ("edit_learner", "Edit Learner"),
                ("learner_details", "Learner Details"),
                ("learner_role_assignment", "Assign Roles"),
                ("learner_schedule_list", "Edit Learner Schedules"),
                ("set_learner_password", "Set Learner Password"),
                ("learner_portal", "Learner Portal"),
                ("learner_home", "Learner Home"),
                ("learner_groups", "Learner Groups"),
                ("learner_project_plans", "Learner Project Plans"),
                ("learner-autocomplete", "Learner Autocomplete"),
            ]),
            ("Groups", [
                ("group_management", "Manage Groups"),
                ("group_create", "Create Group"),
                ("group-edit", "Edit Group"),
                ("group-delete", "Delete Group"),
                ("group_detail", "Group Details"),
                ("upload_group_excel", "Import Groups (Excel)"),
                ("group_qualification_assignment", "Assign Group Qualifications"),
                ("assign_learner_qualification_to_group", "Assign Learner Qualification to Group"),
                ("generate_bulk_admin_pack_zip", "Generate Bulk Admin Pack ZIP"),
            ]),
            ("Qualifications", [
                ("qualification_times_list", "Qualification Times"),
                ("edit_qualification_times", "Edit Qualification Times"),
                ("add_learner_qualification", "Add Learner Qualification"),
                ("get_sla_qualifications", "Get SLA Qualifications (AJAX)"),
                ("get_learner_qualifications", "Get Learner Qualifications (AJAX)"),
            ]),
            ("Projects", [
                ("projectplan_list", "Project Plans"),
                ("projectplan_add", "Add Project Plan"),
                ("projectplan_edit", "Edit Project Plan"),
                ("projectplan_delete", "Delete Project Plan"),
                ("projectplan_detail", "Project Plan Details"),
                ("upload_project_plan_excel", "Import Project Plans (Excel)"),
                ("external_project_list", "External Project List"),
                ("sessiondate_list", "Session Dates"),
                ("sessiondate_add", "Add Session Date"),
                ("sessiondate_edit", "Edit Session Date"),
                ("sessiondate_delete", "Delete Session Date"),
                ("upload_session_date_excel", "Import Session Dates (Excel)"),
            ]),
            ("Venues", [
                ("venue_list", "Venues"),
                ("venue_add", "Add Venue"),
                ("venuebooking_list", "Venue Bookings"),
                ("venuebooking_add", "Book Venue"),
                ("venuebooking_edit", "Edit Venue Booking"),
                ("venuebooking_delete", "Delete Venue Booking"),
                ("venuebooking_calendar", "Venue Calendar"),
                ("venuebooking_events_api", "Venue Booking Events API"),
                ("venuebooking_modal_form", "Venue Booking Modal Form"),
                ("venuebooking_switch", "Switch Venue Booking"),
                ("venue_availability_api", "Venue Availability API"),
            ]),
            ("Services & Modules", [
                ("service-list", "Manage Services"),
                ("service-add", "Add Service"),
                ("service-edit", "Edit Service"),
                ("service-delete", "Delete Service"),
                ("module-list", "Manage Modules"),
                ("module-add", "Add Module"),
                ("module-edit", "Edit Module"),
                ("module-delete", "Delete Module"),
                ("upload_module_excel", "Import Modules (Excel)"),
                ("service-module-list", "Service-Module Links"),
                ("service-module-add", "Link Service & Module"),
                ("service-module-edit", "Edit Service-Module Link"),
                ("service-module-delete", "Delete Service-Module Link"),
            ]),
            ("PoE & Assessments", [
                ("poe_submission", "Submit PoE"),
                ("poe_list", "PoE List"),
                ("poe_template_list", "PoE Templates"),
                ("poe_template_add", "Add PoE Template"),
                ("poe_annextures", "Manage PoE Annexures"),
                ("facilitator_poe_list", "Facilitator PoE List"),
                ("facilitator_poe_feedback", "Facilitator PoE Feedback"),
                ("admin_poe_dashboard", "Admin PoE Dashboard"),
                ("assess_poe", "Assess PoE"),
                ("learner_assessment_results", "Assessment Results"),
            ]),
            ("Adobe Forms", [
                ("adobe_form_list", "Adobe Forms"),
                ("adobe_form_upload", "Upload Adobe Form"),
                ("adobe_form_download", "Download Adobe Form"),
                ("adobe_form_submit", "Submit Adobe Form"),
            ]),
            ("Finance & Billing", [
                ("billing_export", "Export Billing"),
                ("billing_payments", "Manage Payments"),
                ("add_sla_wizard", "Add SLA"),
                ("add_sla_qualifications", "Add SLA Qualifications"),
                ("add_sla_learners", "Add SLA Learners"),
                ("add_sla_billing", "Add SLA Billing"),
            ]),
            ("LIF & Templates", [
                ("lif_form", "LIF Form"),
                ("lif_update", "Update LIF Form"),
                ("lif_delete", "Delete LIF Form"),
                ("upload_lif_template", "Upload LIF Template"),
                ("map_lif_template_fields", "Map LIF Template Fields"),
                ("lif_template_list", "LIF Template List"),
                ("lif_template_delete", "Delete LIF Template"),
                ("generate_lif_word", "Generate LIF Word"),
                ("generate_bulk_lif_zip", "Generate Bulk LIF ZIP"),
                ("cognito_lif_entries", "Cognito LIF Entries"),
                ("export_cognito_to_lif", "Export Cognito to LIF"),
                ("lif_list_json", "LIF List JSON"),
            ]),
            ("Other", [
                ("home_redirect", "Home Redirect"),
                ("view_upcoming_dates", "Upcoming Dates"),
                ("submit_poe", "Portal Submit PoE"),
                ("sit_summative_exam", "Sit Summative Exam"),
                ("role_permission_management", "Role Permission Management"),
                ("switch_role", "Switch Role"),
                ("facilitator-autocomplete", "Facilitator Autocomplete"),
            ]),
        ]
        roles = Role.objects.all()
        role_permissions = {
            role.id: set(RolePermission.objects.filter(role=role, has_access=True).values_list('url_name', flat=True))
            for role in roles
        }
        context.update({
            "roles": roles,
            "url_groups": url_groups,
            "role_permissions": role_permissions,
        })
        return context

    def post(self, request, *args, **kwargs):
        roles = Role.objects.all()
        for role in roles:
            url_names = request.POST.getlist(f'permissions_{role.id}')
            RolePermission.objects.filter(role=role).update(has_access=False)
            for url_name in url_names:
                RolePermission.objects.update_or_create(
                    role=role,
                    url_name=url_name,
                    defaults={'has_access': True}
                )
        messages.success(request, "Permissions updated successfully.")
        return redirect("role_permission_management")


from django.shortcuts import redirect
from django.contrib import messages

class SwitchRoleView(LoginRequiredMixin, View):
    def get(self, request):
        role = request.GET.get('role', '').lower()
        
        try:
            learner = request.user.learner_profile
            available_roles = ['learner'] if hasattr(request.user, 'learner_profile') else []
            available_roles.extend([r.role.name.lower() for r in LearnerRole.objects.filter(learner=learner)])
            
            if role not in available_roles:
                messages.error(request, f"Invalid role selection: {role}")
                return redirect('home_redirect')
            
            if role == 'learner':
                request.session['current_role'] = 'Learner'
                messages.success(request, "Switched to Learner view")
            else:
                # Get proper case for role name
                role_obj = LearnerRole.objects.get(
                    learner=learner,
                    role__name__iexact=role
                )
                request.session['current_role'] = role_obj.role.name
                messages.success(request, f"Switched to {role_obj.role.name} role")
                
        except (AttributeError, LearnerRole.DoesNotExist):
            messages.error(request, "Invalid role or no learner profile found")
            
        return redirect(request.META.get('HTTP_REFERER', 'home_redirect'))

from .models import LearnerRole

def role_context(request):
    """
    Adds role-related context to all templates.
    """
    context = {
        'current_role': None,
        'has_learner_role': False,
        'admin_roles': []
    }
    
    if request.user.is_authenticated:
        try:
            learner = request.user.learner_profile
            # Set current role, defaulting to first available role or None
            current_role = request.session.get('current_role')
            if current_role:
                # Verify the current role is valid
                if current_role == 'Learner':
                    context['has_learner_role'] = True
                else:
                    has_role = LearnerRole.objects.filter(
                        learner=learner,
                        role__name=current_role
                    ).exists()
                    if not has_role:
                        current_role = None
                        request.session.pop('current_role', None)
            
            # If no valid current role, set to first available role
            if not current_role:
                admin_roles = LearnerRole.objects.filter(learner=learner)
                if admin_roles.exists():
                    current_role = admin_roles.first().role.name
                    request.session['current_role'] = current_role
                elif hasattr(request.user, 'learner_profile'):
                    current_role = 'Learner'
                    request.session['current_role'] = current_role
            
            context['current_role'] = current_role
            context['has_learner_role'] = hasattr(request.user, 'learner_profile')
            context['admin_roles'] = LearnerRole.objects.filter(
                learner=learner
            ).select_related('role')
            
        except AttributeError:
            pass
            
    return context


class LearnerAttendanceView(LoginRequiredMixin, TemplateView):
    template_name = 'portal/learner_attendance.html'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            learner = self.request.user.learner_profile
            if not learner:
                messages.error(self.request, "No learner profile found.")
                return context
            
            # Get active qualification
            active_qual = learner.learnerqualification_set.filter(status='active').first()
            context['active_qual'] = active_qual

            # Get filter parameters
            date_filter = self.request.GET.get('date')
            start_date = self.request.GET.get('start_date')
            end_date = self.request.GET.get('end_date')
            clock_in_status = self.request.GET.get('clock_in_status')
            clock_out_status = self.request.GET.get('clock_out_status')

            # Get fingerprint records
            queryset = Fingerprint.objects.filter(user_id=learner.UserID).order_by('-date', 'time')

            if date_filter:
                queryset = queryset.filter(date=date_filter)
            elif start_date and end_date:
                queryset = queryset.filter(date__range=[start_date, end_date])
            elif start_date:
                queryset = queryset.filter(date__gte=start_date)
            elif end_date:
                queryset = queryset.filter(date__lte=end_date)

            # Process fingerprint data
            fingerprint_data = []
            grouped_records = {}
            
            for record in queryset:
                if record.date not in grouped_records:
                    grouped_records[record.date] = []
                grouped_records[record.date].append(record)

            for date, records in grouped_records.items():
                clock_in = min((r.time for r in records), default=None)
                clock_out = max((r.time for r in records), default=None)

                # Get expected times from qualification
                expected_in = time(8, 0)  # Default
                expected_out = time(16, 30)  # Default
                if active_qual and active_qual.sla_qualification:
                    expected_in = active_qual.sla_qualification.expected_clock_in or expected_in
                    expected_out = active_qual.sla_qualification.expected_clock_out or expected_out

                record_data = {
                    'date': date,
                    'user_id': learner.UserID,
                    'learner_name': f"{learner.FirstName} {learner.Surname}",
                    'qualification': active_qual.sla_qualification.service.name if active_qual else 'N/A',
                    'sla_group': active_qual.sla_qualification.sla.sla_reference if active_qual else 'N/A',
                    'expected_clock_in': expected_in,
                    'clock_in_time': clock_in,
                    'clock_in_category': get_clock_in_category(clock_in, expected_in),
                    'expected_clock_out': expected_out,
                    'clock_out_time': clock_out,
                    'clock_out_category': get_clock_out_category(clock_out, expected_out)
                }

                # Apply status filters
                if clock_in_status and record_data['clock_in_category'] != clock_in_status:
                    continue
                if clock_out_status and record_data['clock_out_category'] != clock_out_status:
                    continue

                fingerprint_data.append(record_data)

            # Calculate summary
            summary = {
                'clock_in': {'early': 0, 'on_time': 0, 'late': 0, 'missing': 0},
                'clock_out': {'early_leave': 0, 'on_time': 0, 'overtime': 0, 'missing': 0}
            }

            for record in fingerprint_data:
                ci_cat = record['clock_in_category']
                co_cat = record['clock_out_category']
                if ci_cat in summary['clock_in']:
                    summary['clock_in'][ci_cat] += 1
                if co_cat in summary['clock_out']:
                    summary['clock_out'][co_cat] += 1

            # Paginate
            paginator = Paginator(fingerprint_data, self.paginate_by)
            page = self.request.GET.get('page')
            page_obj = paginator.get_page(page)

            context.update({
                'learner': learner,
                'fingerprint_data': page_obj,
                'is_paginated': page_obj.has_other_pages(),
                'page_obj': page_obj,
                'summary': summary,
                'filters': {
                    'date': date_filter,
                    'start_date': start_date,
                    'end_date': end_date,
                    'clock_in_status': clock_in_status,
                    'clock_out_status': clock_out_status,
                }
            })

        except Exception as e:
            messages.error(self.request, f"Error retrieving attendance data: {str(e)}")

        return context

class LearnerGroupsView(LoginRequiredMixin, TemplateView):
    template_name = 'portal/learner_groups.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            learner = self.request.user.learner_profile
        except Learner.DoesNotExist:
            messages.error(self.request, "No learner profile found.")
            return context

        # Get all groups for the learner
        groups = Group.objects.filter(
            sla_qualifications__in=learner.learnerqualification_set.values_list('sla_qualification', flat=True)
        ).distinct()

        context.update({
            'learner': learner,
            'groups': groups,
        })
        return context

class LearnerProjectPlansView(LoginRequiredMixin, TemplateView):
    template_name = 'portal/learner_project_plans.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            learner = self.request.user.learner_profile
        except Learner.DoesNotExist:
            messages.error(self.request, "No learner profile found.")
            return context

        # Get filter parameters
        group_id = self.request.GET.get('group_id')
        project_plan_id = self.request.GET.get('project_plan_id')

        # Get all groups for the learner
        groups = Group.objects.filter(
            sla_qualifications__in=learner.learnerqualification_set.values_list('sla_qualification', flat=True)
        ).distinct()

        # Get project plans
        project_plans = ProjectPlan.objects.filter(group__in=groups)
        if group_id:
            project_plans = project_plans.filter(group_id=group_id)
        if project_plan_id:
            project_plans = project_plans.filter(id=project_plan_id)

        # Get sessions for each project plan
        plan_sessions = {}
        for plan in project_plans:
            plan_sessions[plan.id] = plan.sessiondate_set.all()

        context.update({
            'learner': learner,
            'groups': groups,
            'project_plans': project_plans,
            'plan_sessions': plan_sessions,
            'filters': {
                'group_id': group_id,
                'project_plan_id': project_plan_id,
            }
        })
        return context
    
from django.views.generic import UpdateView, DeleteView
from .models import Group
from django.urls import reverse_lazy

class GroupUpdateView(RolePermissionRequiredMixin, UpdateView):
    model = Group
    fields = ['name', 'projectcode', 'service', 'seta', 'start_date', 'end_date']
    template_name = 'core/group_form.html'
    success_url = reverse_lazy('group_management')

class GroupDeleteView(RolePermissionRequiredMixin, DeleteView):
    model = Group
    template_name = 'core/group_confirm_delete.html'
    success_url = reverse_lazy('group_management')

from django.views.decorators.http import require_GET

@require_GET
def venue_availability_api(request):
    from .models import VenueBooking, Venue
    import json
    venue_id = request.GET.get('venue_id')
    date_str = request.GET.get('date')
    if not venue_id or not date_str:
        return JsonResponse({'available': False, 'reason': 'Missing venue or date'}, status=400)
    bookings = VenueBooking.objects.filter(
        venue_id=venue_id,
        start_datetime__date=date_str,
        status__in=['booked', 'rescheduled']
    ).order_by('start_datetime')
    available = not bookings.exists()
    booking_list = [
        {
            'start': b.start_datetime.strftime('%H:%M'),
            'end': b.end_datetime.strftime('%H:%M'),
            'purpose': b.booking_purpose,
            'status': b.status
        }
        for b in bookings
    ]
    return JsonResponse({'available': available, 'bookings': booking_list})

from .models import LIF
from .forms import LIFForm

from django.views.generic import CreateView
from django.urls import reverse_lazy

from django.views.generic import CreateView
from django.urls import reverse_lazy
from django.contrib import messages
from django.shortcuts import render
from .models import LIF
from .forms import LIFForm

from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import CreateView
from django.urls import reverse_lazy
from django.contrib import messages
from django.shortcuts import render
from .models import LIF, Learner, LearnerRole
from .forms import LIFForm

from django.views.generic import CreateView, UpdateView
from django.urls import reverse_lazy
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.forms import modelformset_factory
from django.http import HttpResponse
from .models import LIF, Learner, LearnerRole, LIFTemplate, LIFTemplateFieldMap
from .forms import LIFForm, LIFTemplateUploadForm, LIFTemplateFieldMapForm
from docx import Document
from datetime import date

from django.views.generic import CreateView, UpdateView
from django.urls import reverse_lazy
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.forms import modelformset_factory
from django.http import HttpResponse
from .models import LIF, Learner, LearnerRole, LIFTemplate, LIFTemplateFieldMap
from .forms import LIFForm, LIFTemplateUploadForm, LIFTemplateFieldMapForm
from docx import Document
from datetime import date

class LIFCreateView(LoginRequiredMixin, CreateView):
    model = LIF
    form_class = LIFForm
    template_name = 'core/lif_form.html'
    success_url = reverse_lazy('lif_form')

    def dispatch(self, request, *args, **kwargs):
        user = request.user
        is_learner = hasattr(user, 'learner_profile')
        is_admin = user.is_staff or user.is_superuser
        has_projects_role = LearnerRole.objects.filter(learner__user=user, role__name__iexact='Projects').exists()

        if is_learner and not (is_admin or has_projects_role):
            learner = user.learner_profile
            if LIF.objects.filter(learner=learner).exists():
                return redirect('lif_update')
        return super().dispatch(request, *args, **kwargs)

    def get_initial(self):
        initial = super().get_initial()
        learner_id = self.request.GET.get('learner_id')
        if learner_id:
            initial['learner'] = learner_id
        return initial

    def form_valid(self, form):
        user = self.request.user
        learner_id = self.request.GET.get('learner_id')
        is_learner = hasattr(user, 'learner_profile')
        is_admin = user.is_staff or user.is_superuser
        has_projects_role = LearnerRole.objects.filter(learner__user=user, role__name__iexact='Projects').exists()

        if is_learner and not (is_admin or has_projects_role):
            learner = user.learner_profile
            form.instance.learner = learner
        else:
            if learner_id:
                learner = Learner.objects.filter(id=learner_id).first()
                form.instance.learner = learner
            else:
                form.instance.learner = None

        self.object = form.save()
        messages.success(self.request, "Your LIF form was submitted successfully.", extra_tags='lif')
        return redirect(self.success_url)

    def form_invalid(self, form):
        messages.error(self.request, "Please correct the errors below.", extra_tags='lif')
        return self.render_to_response(self.get_context_data(form=form))


class LIFUpdateView(LoginRequiredMixin, UpdateView):
    model = LIF
    form_class = LIFForm
    template_name = 'core/lif_form.html'
    success_url = reverse_lazy('lif_form')

    def get_object(self, queryset=None):
        user = self.request.user
        return user.learner_profile.lif_form

    def form_valid(self, form):
        messages.success(self.request, "Your LIF form was updated successfully.", extra_tags='lif')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "Please correct the errors below.", extra_tags='lif')
        return self.render_to_response(self.get_context_data(form=form))

def upload_lif_template(request):
    if request.method == 'POST':
        form = LIFTemplateUploadForm(request.POST, request.FILES)
        if form.is_valid():
            template = form.save()
            # Auto-map placeholders to fields
            from .forms import PLACEHOLDER_CHOICES, LIF_FIELD_CHOICES
            # Define mapping rules for cross-based placeholders
            placeholder_to_field = {
                # Direct field mappings
                '{{learner}}': 'learner',
                '{{national_id}}': 'national_id',
                '{{alternative_id}}': 'alternative_id',
                '{{alternative_id_type}}': 'alternative_id_type',
                '{{equity_code}}': 'equity_code',
                '{{nationality_code}}': 'nationality_code',
                '{{gender_code}}': 'gender_code',
                '{{citizen_resident_status_code}}': 'citizen_resident_status_code',
                '{{home_language_code}}': 'home_language_code',
                '{{province_code}}': 'province_code',
                '{{disability_status_code}}': 'disability_status_code',
                '{{socio_economic_status_code}}': 'socio_economic_status_code',
                '{{learner_title}}': 'learner_title',
                '{{learner_birth_date}}': 'learner_birth_date',
                '{{learner_first_name}}': 'learner_first_name',
                '{{learner_middle_name}}': 'learner_middle_name',
                '{{learner_last_name}}': 'learner_last_name',
                '{{learner_previous_last_name}}': 'learner_previous_last_name',
                '{{learner_current_occupation}}': 'learner_current_occupation',
                '{{years_in_occupation}}': 'years_in_occupation',
                '{{employer}}': 'employer',
                '{{highest_secondary_education}}': 'highest_secondary_education',
                '{{secondary_school_name}}': 'secondary_school_name',
                '{{secondary_year_completed}}': 'secondary_year_completed',
                '{{percentage_maths}}': 'percentage_maths',
                '{{percentage_first_language}}': 'percentage_first_language',
                '{{percentage_second_language}}': 'percentage_second_language',
                '{{highest_tertiary_education}}': 'highest_tertiary_education',
                '{{tertiary_school_name}}': 'tertiary_school_name',
                '{{tertiary_year_completed}}': 'tertiary_year_completed',
                '{{address_line1}}': 'address_line1',
                '{{address_line2}}': 'address_line2',
                '{{city}}': 'city',
                '{{state_province}}': 'state_province',
                '{{phone_number}}': 'phone_number',
                '{{alt_contact_number}}': 'alt_contact_number',
                '{{email_address}}': 'email_address',
                '{{provider_etqa_id}}': 'provider_etqa_id',
                '{{provider_code}}': 'provider_code',
                '{{programme_title}}': 'programme_title',
                '{{qualification_code}}': 'qualification_code',
                '{{nqf_level}}': 'nqf_level',
                '{{sponsor}}': 'sponsor',
                '{{duration_start_date}}': 'duration_start_date',
                '{{duration_end_date}}': 'duration_end_date',
                '{{consent_to_process}}': 'consent_to_process',
                '{{postal_code}}': 'postal_code',
                # Cross-based placeholders (mapped to their parent field)
                '{{gender_male}}': 'gender_code',
                '{{gender_female}}': 'gender_code',
                '{{below_35_yes}}': 'learner_birth_date',
                '{{below_35_no}}': 'learner_birth_date',
                '{{equity_african}}': 'equity_code',
                '{{equity_coloured}}': 'equity_code',
                '{{equity_indian}}': 'equity_code',
                '{{equity_white}}': 'equity_code',
                '{{disability_sight}}': 'disability_status_code',
                '{{disability_hearing}}': 'disability_status_code',
                '{{disability_communication}}': 'disability_status_code',
                '{{disability_physical}}': 'disability_status_code',
                '{{disability_intellectual}}': 'disability_status_code',
                '{{disability_emotional}}': 'disability_status_code',
                '{{disability_multiple}}': 'disability_status_code',
                '{{disability_unspecified}}': 'disability_status_code',
                '{{disability_none}}': 'disability_status_code',
                '{{citizen_sa}}': 'citizen_resident_status_code',
                '{{citizen_dual}}': 'citizen_resident_status_code',
                '{{citizen_other}}': 'citizen_resident_status_code',
                '{{citizen_permanent}}': 'citizen_resident_status_code',
                '{{citizen_unknown}}': 'citizen_resident_status_code',
                '{{nationality_afrikaans}}': 'nationality_code',
                '{{nationality_english}}': 'nationality_code',
                '{{nationality_ndebele}}': 'nationality_code',
                '{{nationality_sepedi}}': 'nationality_code',
                '{{nationality_sesotho}}': 'nationality_code',
                '{{nationality_setswana}}': 'nationality_code',
                '{{nationality_siswati}}': 'nationality_code',
                '{{nationality_tshivenda}}': 'nationality_code',
                '{{nationality_isixhosa}}': 'nationality_code',
                '{{nationality_xitsonga}}': 'nationality_code',
                '{{nationality_isizulu}}': 'nationality_code',
                '{{nationality_sasl}}': 'nationality_code',
                '{{nationality_other}}': 'nationality_code',
                '{{nationality_unknown}}': 'nationality_code',
                '{{home_language_afrikaans}}': 'home_language_code',
                '{{home_language_english}}': 'home_language_code',
                '{{home_language_ndebele}}': 'home_language_code',
                '{{home_language_sepedi}}': 'home_language_code',
                '{{home_language_sesotho}}': 'home_language_code',
                '{{home_language_setswana}}': 'home_language_code',
                '{{home_language_siswati}}': 'home_language_code',
                '{{home_language_tshivenda}}': 'home_language_code',
                '{{home_language_isixhosa}}': 'home_language_code',
                '{{home_language_xitsonga}}': 'home_language_code',
                '{{home_language_isizulu}}': 'home_language_code',
                '{{home_language_sasl}}': 'home_language_code',
                '{{home_language_other}}': 'home_language_code',
                '{{home_language_unknown}}': 'home_language_code',
                '{{province_western_cape}}': 'province_code',
                '{{province_eastern_cape}}': 'province_code',
                '{{province_northern_cape}}': 'province_code',
                '{{province_free_state}}': 'province_code',
                '{{province_kwazulu_natal}}': 'province_code',
                '{{province_north_west}}': 'province_code',
                '{{province_gauteng_jhb}}': 'province_code',
                '{{province_gauteng_pta}}': 'province_code',
                '{{province_mpumalanga}}': 'province_code',
                '{{province_limpopo}}': 'province_code',
                '{{province_outside_sa}}': 'province_code',
                '{{province_national}}': 'province_code',
                '{{socio_employed}}': 'socio_economic_status_code',
                '{{socio_unemployed_seeking}}': 'socio_economic_status_code',
                '{{socio_not_working_not_looking}}': 'socio_economic_status_code',
                '{{socio_homemaker}}': 'socio_economic_status_code',
                '{{socio_student}}': 'socio_economic_status_code',
                '{{socio_pensioner}}': 'socio_economic_status_code',
                '{{socio_disabled}}': 'socio_economic_status_code',
                '{{socio_no_wish_to_work}}': 'socio_economic_status_code',
                '{{socio_not_working_nec}}': 'socio_economic_status_code',
                '{{socio_aged_under_15}}': 'socio_economic_status_code',
                '{{socio_institution}}': 'socio_economic_status_code',
                '{{socio_unspecified}}': 'socio_economic_status_code',
                '{{alt_id_saqa}}': 'alternative_id_type',
                '{{alt_id_passport}}': 'alternative_id_type',
                '{{alt_id_driver}}': 'alternative_id_type',
                '{{alt_id_temp_id}}': 'alternative_id_type',
                '{{alt_id_none}}': 'alternative_id_type',
                '{{alt_id_unknown}}': 'alternative_id_type',
                '{{alt_id_student}}': 'alternative_id_type',
                '{{alt_id_work_permit}}': 'alternative_id_type',
                '{{alt_id_employee}}': 'alternative_id_type',
                '{{alt_id_birth_cert}}': 'alternative_id_type',
                '{{alt_id_hsrc}}': 'alternative_id_type',
                '{{alt_id_etqa}}': 'alternative_id_type',
                '{{alt_id_refugee}}': 'alternative_id_type',
                '{{tertiary_national_certificate}}': 'highest_tertiary_education',
                '{{tertiary_national_diploma}}': 'highest_tertiary_education',
                '{{tertiary_first_degree}}': 'highest_tertiary_education',
                '{{tertiary_post_doctoral}}': 'highest_tertiary_education',
                '{{tertiary_doctoral}}': 'highest_tertiary_education',
                '{{tertiary_professional}}': 'highest_tertiary_education',
                '{{tertiary_honours}}': 'highest_tertiary_education',
                '{{tertiary_higher_diploma}}': 'highest_tertiary_education',
                '{{tertiary_masters_diploma}}': 'highest_tertiary_education',
                '{{tertiary_national_higher}}': 'highest_tertiary_education',
                '{{tertiary_further_diploma}}': 'highest_tertiary_education',
                '{{tertiary_post_graduate}}': 'highest_tertiary_education',
                '{{tertiary_senior_certificate}}': 'highest_tertiary_education',
                '{{tertiary_qual_nat_sen_cert}}': 'highest_tertiary_education',
                '{{tertiary_apprenticeship}}': 'highest_tertiary_education',
                '{{tertiary_post_grad_b_degree}}': 'highest_tertiary_education',
                '{{tertiary_post_diploma_diploma}}': 'highest_tertiary_education',
                '{{tertiary_post_basic_diploma}}': 'highest_tertiary_education',
            }
            # Create mappings for all placeholders
            for placeholder, _ in PLACEHOLDER_CHOICES:
                lif_field = placeholder_to_field.get(placeholder)
                if lif_field:
                    LIFTemplateFieldMap.objects.get_or_create(
                        template=template,
                        placeholder=placeholder,
                        defaults={'lif_field': lif_field}
                    )
            messages.success(request, "Template uploaded and fields auto-mapped. You can review or edit mappings.", extra_tags='lif')
            return redirect('map_lif_template_fields', template_id=template.id)
    else:
        form = LIFTemplateUploadForm()
    return render(request, 'core/upload_lif_template.html', {'form': form})
# Custom formset to prevent duplicate placeholders
from django.forms import BaseModelFormSet

class LIFTemplateFieldMapFormSet(BaseModelFormSet):
    def clean(self):
        super().clean()
        placeholders = []
        existing_mappings = {}
        template_id = None

        # Get the template ID from the first form's instance
        for form in self.forms:
            if form.instance and form.instance.template_id:
                template_id = form.instance.template_id
                break
        if not template_id and 'template_id' in self.data:
            template_id = self.data.get('template_id')

        # Load existing mappings from the database
        if template_id:
            existing_mappings = {
                mapping.placeholder: mapping.lif_field
                for mapping in LIFTemplateFieldMap.objects.filter(template_id=template_id)
            }

        for form in self.forms:
            if form.cleaned_data and not form.cleaned_data.get('DELETE', False):
                placeholder = form.cleaned_data.get('placeholder')
                lif_field = form.cleaned_data.get('lif_field')
                if placeholder and lif_field:  # Only validate non-empty rows
                    if placeholder in placeholders:
                        form.add_error('placeholder', f"Duplicate placeholder: {placeholder}")
                    elif placeholder in existing_mappings and not form.instance.pk:
                        form.add_error('placeholder', f"Placeholder {placeholder} is already mapped to {existing_mappings[placeholder]}")
                    placeholders.append(placeholder)


def map_lif_template_fields(request, template_id):
    template = get_object_or_404(LIFTemplate, id=template_id)
    extra = int(request.GET.get('extra', 1))
    FieldMapFormSet = modelformset_factory(
        LIFTemplateFieldMap,
        form=LIFTemplateFieldMapForm,
        formset=LIFTemplateFieldMapFormSet,
        extra=extra,
        can_delete=True
    )
    if request.method == 'POST':
        formset = FieldMapFormSet(request.POST, queryset=LIFTemplateFieldMap.objects.filter(template=template), initial=[{'template_id': template_id}])
        if formset.is_valid():
            for form in formset:
                cleaned = form.cleaned_data
                if cleaned.get('DELETE', False) and form.instance.pk:
                    form.instance.delete()
                    continue
                placeholder = cleaned.get('placeholder')
                lif_field = cleaned.get('lif_field')
                if not (placeholder and lif_field):  # Skip empty rows
                    continue
                mapping, created = LIFTemplateFieldMap.objects.get_or_create(
                    template=template,
                    placeholder=placeholder,
                    defaults={'lif_field': lif_field}
                )
                if not created and mapping.lif_field != lif_field:
                    mapping.lif_field = lif_field
                    mapping.save()
            messages.success(request, "Field mappings saved successfully.", extra_tags='lif')
            return redirect('map_lif_template_fields', template_id=template.id)
        else:
            # Provide specific error messages
            error_messages = []
            for form in formset:
                if form.errors:
                    for field, errors in form.errors.items():
                        for error in errors:
                            error_messages.append(f"Row {form.prefix}: {error}")
            if error_messages:
                messages.error(request, "Please correct the following errors: " + "; ".join(error_messages), extra_tags='lif')
            else:
                messages.error(request, "Please correct the errors below.", extra_tags='lif')
    else:
        formset = FieldMapFormSet(queryset=LIFTemplateFieldMap.objects.filter(template=template))
    return render(request, 'core/map_lif_template_fields.html', {'formset': formset, 'template': template})


def lif_template_list(request):
    templates = LIFTemplate.objects.all()
    return render(request, 'core/lif_template_list.html', {'templates': templates})


def generate_lif_word(request):
    import io
    import zipfile
    import re
    from django.core.paginator import Paginator
    from django.db.models import Q
    from docx import Document
    from datetime import date
    from django.http import HttpResponse
    from django.shortcuts import render, get_object_or_404
    from core.models import LIF, LIFTemplate, LIFTemplateFieldMap

    # --- Filtering ---
    search = request.GET.get('search', '').strip()
    template_id = request.GET.get('template_id', '')
    gender = request.GET.get('gender', '')
    equity = request.GET.get('equity', '')
    page = request.GET.get('page', 1)

    lifs = LIF.objects.all()
    if search:
        lifs = lifs.filter(
            Q(learner_first_name__icontains=search) |
            Q(learner_last_name__icontains=search) |
            Q(national_id__icontains=search) |
            Q(email_address__icontains=search)
        )
    if gender:
        lifs = lifs.filter(gender_code=gender)
    if equity:
        lifs = lifs.filter(equity_code=equity)
    lifs = lifs.order_by('-id')

    paginator = Paginator(lifs, 15)
    lifs_page = paginator.get_page(page)

    templates = LIFTemplate.objects.all().order_by('name')

    filters = {
        'search': search,
        'template_id': template_id,
        'gender': gender,
        'equity': equity,
    }

    allow_zip = lifs.count() > 1

    # --- POST: Generate LIF doc(s) ---
    if request.method == 'POST':
        lif_ids = request.POST.getlist('lif_ids')
        template_id = request.POST.get('template_id')
        errors = []
        if not template_id:
            errors.append("Please select a template.")
        if not lif_ids:
            errors.append("Please select at least one LIF.")
        if errors:
            return render(request, 'core/generate_lif_word.html', {
                'lifs': lifs_page,
                'templates': templates,
                'filters': filters,
                'allow_zip': allow_zip,
                'errors': errors,
            })

        template = get_object_or_404(LIFTemplate, id=template_id)
        mappings = LIFTemplateFieldMap.objects.filter(template=template)

        def get_choice_display(choices, code):
            return dict(choices).get(code, '')

        table_placeholders = [
            '{{gender_male}}', '{{gender_female}}', '{{below_35_yes}}', '{{below_35_no}}',
            '{{equity_african}}', '{{equity_coloured}}', '{{equity_indian}}', '{{equity_white}}',
            '{{disability_sight}}', '{{disability_hearing}}', '{{disability_communication}}',
            '{{disability_physical}}', '{{disability_intellectual}}', '{{disability_emotional}}',
            '{{disability_multiple}}', '{{disability_unspecified}}', '{{disability_none}}',
            '{{citizen_sa}}', '{{citizen_dual}}', '{{citizen_other}}', '{{citizen_permanent}}',
            '{{citizen_unknown}}',
            '{{socio_employed}}',                # 01
            '{{socio_unemployed_seeking}}',      # 02
            '{{socio_not_working_not_looking}}', # 03
            '{{socio_homemaker}}',               # 04
            '{{socio_student}}',                 # 06
            '{{socio_pensioner}}',               # 07
            '{{socio_disabled}}',                # 08
            '{{socio_no_wish_to_work}}',         # 09
            '{{socio_not_working_nec}}',         # 10
            '{{socio_aged_under_15}}',           # 97
            '{{socio_institution}}',             # 98
            '{{socio_unspecified}}',             # U
            '{{tertiary_national_certificate}}',
            '{{tertiary_national_diploma}}',
            '{{tertiary_first_degree}}',
            '{{tertiary_post_doctoral}}',
            '{{tertiary_doctoral}}',
            '{{tertiary_professional}}',
            '{{tertiary_honours}}',
            '{{tertiary_higher_diploma}}',
            '{{tertiary_masters_diploma}}',
            '{{tertiary_national_higher}}',
            '{{tertiary_further_diploma}}',
            '{{tertiary_post_graduate}}',
            '{{tertiary_senior_certificate}}',
            '{{tertiary_qual_nat_sen_cert}}',
            '{{tertiary_apprenticeship}}',
            '{{tertiary_post_grad_b_degree}}',
            '{{tertiary_post_diploma_diploma}}',
            '{{tertiary_post_basic_diploma}}',
            '{{province_western_cape}}',
            '{{province_eastern_cape}}',
            '{{province_northern_cape}}',
            '{{province_free_state}}',
            '{{province_kwazulu_natal}}',
            '{{province_north_west}}',
            '{{province_gauteng_jhb}}',
            '{{province_gauteng_pta}}',
            '{{province_mpumalanga}}',
            '{{province_limpopo}}',
            '{{province_outside_sa}}',
            '{{province_national}}',
            '{{home_language_afrikaans}}',
            '{{home_language_english}}',
            '{{home_language_ndebele}}',
            '{{home_language_sepedi}}',
            '{{home_language_sesotho}}',
            '{{home_language_setswana}}',
            '{{home_language_siswati}}',
            '{{home_language_tshivenda}}',
            '{{home_language_isixhosa}}',
            '{{home_language_xitsonga}}',
            '{{home_language_isizulu}}',
            '{{home_language_sasl}}',
            '{{home_language_other}}',
            '{{home_language_unknown}}',
            '{{tertiary_national_certificate}}',
            '{{tertiary_national_diploma}}',
            '{{tertiary_first_degree}}',
            '{{tertiary_post_doctoral}}',
            '{{tertiary_doctoral}}',
            '{{tertiary_professional}}',
            '{{tertiary_honours}}',
            '{{tertiary_higher_diploma}}',
            '{{tertiary_masters_diploma}}',
            '{{tertiary_national_higher}}',
            '{{tertiary_further_diploma}}',
            '{{tertiary_post_graduate}}',
            '{{tertiary_senior_certificate}}',
            '{{tertiary_qual_nat_sen_cert}}',
            '{{tertiary_apprenticeship}}',
            '{{tertiary_post_grad_b_degree}}',
            '{{tertiary_post_diploma_diploma}}',
            '{{tertiary_post_basic_diploma}}',
            '{{secondary_grade_8}}',
            '{{secondary_grade_9}}',
            '{{secondary_grade_10}}',
            '{{secondary_grade_11}}',
            '{{secondary_grade_12}}',
            '{{nationality_sa}}',
            '{{nationality_sdc}}',
            '{{nationality_ang}}',
            '{{nationality_bot}}',
            '{{nationality_les}}',
            '{{nationality_mal}}',
            '{{nationality_mau}}',
            '{{nationality_moz}}',
            '{{nationality_nam}}',
            '{{nationality_sey}}',
            '{{nationality_swa}}',
            '{{nationality_tan}}',
            '{{nationality_zai}}',
            '{{nationality_zam}}',
            '{{nationality_zim}}',
            '{{nationality_ais}}',
            '{{nationality_aus}}',
            '{{nationality_eur}}',
            '{{nationality_nor}}',
            '{{nationality_sou}}',
            '{{nationality_roa}}',
            '{{nationality_ooc}}',
            '{{nationality_u}}',
            '{{nationality_not}}',
        ]

        def build_context(lif):
            context = {
                '{{learner_first_name}}': lif.learner_first_name or '',
                '{{learner_middle_name}}': lif.learner_middle_name or '',
                '{{learner_last_name}}': lif.learner_last_name or '',
                '{{national_id}}': lif.national_id or '',
                '{{alternative_id}}': lif.alternative_id or '',
                '{{alternative_id_type}}': lif.alternative_id_type or '',
                '{{alternative_id_type_name}}': get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, lif.alternative_id_type) or '',
                '{{learner_title}}': lif.learner_title or '',
                '{{learner_birth_date}}': lif.learner_birth_date.strftime('%Y-%m-%d') if lif.learner_birth_date else '',
                '{{learner_previous_last_name}}': lif.learner_previous_last_name or '',
                '{{learner_current_occupation}}': lif.learner_current_occupation or '',
                '{{years_in_occupation}}': str(lif.years_in_occupation) if lif.years_in_occupation else '',
                '{{employer}}': lif.employer or '',
                '{{highest_secondary_education}}': lif.highest_secondary_education or '',
                '{{secondary_school_name}}': lif.secondary_school_name or '',
                '{{secondary_year_completed}}': str(lif.secondary_year_completed) if lif.secondary_year_completed else '',
                '{{percentage_maths}}': str(lif.percentage_maths) if lif.percentage_maths else '',
                '{{percentage_first_language}}': str(lif.percentage_first_language) if lif.percentage_first_language else '',
                '{{percentage_second_language}}': str(lif.percentage_second_language) if lif.percentage_second_language else '',
                '{{highest_tertiary_education}}': lif.highest_tertiary_education or '',
                '{{highest_tertiary_education_name}}': get_choice_display(lif.TERTIARY_CHOICES, lif.highest_tertiary_education) or '',
                '{{tertiary_school_name}}': lif.tertiary_school_name or '',
                '{{tertiary_year_completed}}': str(lif.tertiary_year_completed) if lif.tertiary_year_completed else '',
                '{{address_line1}}': lif.address_line1 or '',
                '{{address_line2}}': lif.address_line2 or '',
                '{{city}}': lif.city or '',
                '{{state_province}}': lif.state_province or '',
                '{{postal_code}}': lif.postal_code or '',
                '{{phone_number}}': lif.phone_number or '',
                '{{alt_contact_number}}': lif.alt_contact_number or '',
                '{{email_address}}': lif.email_address or '',
                '{{provider_etqa_id}}': lif.provider_etqa_id or '',
                '{{provider_code}}': lif.provider_code or '',
                '{{programme_title}}': lif.programme_title or '',
                '{{qualification_code}}': lif.qualification_code or '',
                '{{nqf_level}}': lif.nqf_level or '',
                '{{sponsor}}': lif.sponsor or '',
                '{{commencement_date}}': lif.duration_start_date.strftime('%d/%m/%Y') if lif.duration_start_date else '',
                '{{termination_date}}': lif.duration_end_date.strftime('%d/%m/%Y') if lif.duration_end_date else '',
                '{{consent_to_process}}': ' Agree' if lif.consent_to_process else ' Disagree',
                '{{equity_code}}': lif.equity_code or '',
                '{{equity_name}}': get_choice_display(lif.EQUITY_CHOICES, lif.equity_code) or '',
                '{{nationality_code}}': lif.nationality_code or '',
                '{{nationality_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.nationality_code) or '',
                '{{gender_code}}': lif.gender_code or '',
                '{{gender_name}}': dict([('F', 'Female'), ('M', 'Male')]).get(lif.gender_code, ''),
                '{{citizen_resident_status_code}}': lif.citizen_resident_status_code or '',
                '{{citizen_resident_status_name}}': get_choice_display(lif.CITIZEN_STATUS_CHOICES, lif.citizen_resident_status_code) or '',
                '{{home_language_code}}': lif.home_language_code or '',
                '{{home_language_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.home_language_code) or '',
                '{{province_code}}': lif.province_code or '',
                '{{province_name}}': get_choice_display(lif.PROVINCE_CHOICES, lif.province_code) or '',
                '{{disability_status_code}}': lif.disability_status_code or '',
                '{{disability_status_name}}': get_choice_display(lif.DISABILITY_CHOICES, lif.disability_status_code) or '',
                '{{socio_economic_status_code}}': lif.socio_economic_status_code or '',
                '{{socio_economic_status_name}}': get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, lif.socio_economic_status_code) or '',
                
            }
            # Add split fields
            if lif.learner_birth_date:
                birth_str = lif.learner_birth_date.strftime('%Y%m%d')
                for i in range(8):
                    context[f'{{{{learner_birth_date_{i}}}}}'] = birth_str[i] if i < len(birth_str) else ''
            else:
                for i in range(8):
                    context[f'{{{{learner_birth_date_{i}}}}}'] = ''
            if lif.national_id:
                for i in range(13):
                    context[f'{{{{national_id_{i}}}}}'] = lif.national_id[i] if i < len(lif.national_id) else ''
            else:
                for i in range(13):
                    context[f'{{{{national_id_{i}}}}}'] = ''
            if lif.duration_start_date:
                start_str = lif.duration_start_date.strftime('%d%m%Y')
                for i in range(8):
                    context[f'{{{{commencement_date_{i}}}}}'] = start_str[i] if i < len(start_str) else ''
            else:
                for i in range(8):
                    context[f'{{{{commencement_date_{i}}}}}'] = ''
            if lif.duration_end_date:
                end_str = lif.duration_end_date.strftime('%d%m%Y')
                for i in range(8):
                    context[f'{{{{termination_date_{i}}}}}'] = end_str[i] if i < len(end_str) else ''
            else:
                for i in range(8):
                    context[f'{{{{termination_date_{i}}}}}'] = ''
            # Table-based "checkbox" fields (X for selected)
            context.update({
                '{{gender_male}}': 'X' if lif.gender_code == 'M' else '',
                '{{gender_female}}': 'X' if lif.gender_code == 'F' else '',
                '{{below_35_yes}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year < 35) else '',
                '{{below_35_no}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year >= 35) else '',
                '{{equity_african}}': 'X' if lif.equity_code == 'BA' else '',
                '{{equity_coloured}}': 'X' if lif.equity_code == 'BC' else '',
                '{{equity_indian}}': 'X' if lif.equity_code == 'BI' else '',
                '{{equity_white}}': 'X' if lif.equity_code == 'Wh' else '',
                '{{disability_sight}}': 'X' if lif.disability_status_code == '01' else '',
                '{{disability_hearing}}': 'X' if lif.disability_status_code == '02' else '',
                '{{disability_communication}}': 'X' if lif.disability_status_code == '03' else '',
                '{{disability_physical}}': 'X' if lif.disability_status_code == '04' else '',
                '{{disability_intellectual}}': 'X' if lif.disability_status_code == '05' else '',
                '{{disability_emotional}}': 'X' if lif.disability_status_code == '06' else '',
                '{{disability_multiple}}': 'X' if lif.disability_status_code == '07' else '',
                '{{disability_unspecified}}': 'X' if lif.disability_status_code == '09' else '',
                '{{disability_none}}': 'X' if lif.disability_status_code == 'N' else '',
                '{{citizen_sa}}': 'X' if lif.citizen_resident_status_code == 'SA' else '',
                '{{citizen_dual}}': 'X' if lif.citizen_resident_status_code == 'D' else '',
                '{{citizen_other}}': 'X' if lif.citizen_resident_status_code == 'O' else '',
                '{{citizen_permanent}}': 'X' if lif.citizen_resident_status_code == 'PR' else '',
                '{{citizen_unknown}}': 'X' if lif.citizen_resident_status_code == 'U' else '',
                '{{socio_employed}}': 'X' if lif.socio_economic_status_code == '01' else '',
                '{{socio_unemployed_seeking}}': 'X' if lif.socio_economic_status_code == '02' else '',
                '{{socio_not_working_not_looking}}': 'X' if lif.socio_economic_status_code == '03' else '',
                '{{socio_homemaker}}': 'X' if lif.socio_economic_status_code == '04' else '',
                '{{socio_student}}': 'X' if lif.socio_economic_status_code == '06' else '',
                '{{socio_pensioner}}': 'X' if lif.socio_economic_status_code == '07' else '',
                '{{socio_disabled}}': 'X' if lif.socio_economic_status_code == '08' else '',
                '{{socio_no_wish_to_work}}': 'X' if lif.socio_economic_status_code == '09' else '',
                '{{socio_not_working_nec}}': 'X' if lif.socio_economic_status_code == '10' else '',
                '{{socio_aged_under_15}}': 'X' if lif.socio_economic_status_code == '97' else '',
                '{{socio_institution}}': 'X' if lif.socio_economic_status_code == '98' else '',
                '{{socio_unspecified}}': 'X' if lif.socio_economic_status_code == 'U' else '',
                '{{tertiary_national_certificate}}': 'X' if lif.highest_tertiary_education == 'National Certificate' else '',
                '{{tertiary_national_diploma}}': 'X' if lif.highest_tertiary_education == 'National Diploma' else '',
                '{{tertiary_first_degree}}': 'X' if lif.highest_tertiary_education == 'National First Degree' else '',
                '{{tertiary_post_doctoral}}': 'X' if lif.highest_tertiary_education == 'Post-doctoral Degree' else '',
                '{{tertiary_doctoral}}': 'X' if lif.highest_tertiary_education == 'Doctoral Degree' else '',
                '{{tertiary_professional}}': 'X' if lif.highest_tertiary_education == 'Professional Qualification' else '',
                '{{tertiary_honours}}': 'X' if lif.highest_tertiary_education == 'Honours Degree' else '',
                '{{tertiary_higher_diploma}}': 'X' if lif.highest_tertiary_education == 'National Higher Diploma' else '',
                '{{tertiary_masters_diploma}}': 'X' if lif.highest_tertiary_education == 'National Masters Diploma' else '',
                '{{tertiary_national_higher}}': 'X' if lif.highest_tertiary_education == 'National Higher' else '',
                '{{tertiary_further_diploma}}': 'X' if lif.highest_tertiary_education == 'Further Diploma' else '',
                '{{tertiary_post_graduate}}': 'X' if lif.highest_tertiary_education == 'Post Graduate' else '',
                '{{tertiary_senior_certificate}}': 'X' if lif.highest_tertiary_education == 'Senior Certificate' else '',
                '{{tertiary_qual_nat_sen_cert}}': 'X' if lif.highest_tertiary_education == 'Qual at Nat Sen Cert' else '',
                '{{tertiary_apprenticeship}}': 'X' if lif.highest_tertiary_education == 'Apprenticeship' else '',
                '{{tertiary_post_grad_b_degree}}': 'X' if lif.highest_tertiary_education == 'Post Grad B Degree' else '',
                '{{tertiary_post_diploma_diploma}}': 'X' if lif.highest_tertiary_education == 'Post Diploma Diploma' else '',
                '{{tertiary_post_basic_diploma}}': 'X' if lif.highest_tertiary_education == 'Post-basic Diploma' else '',
                '{{province_western_cape}}': 'X' if lif.province_code == '1' else '',
                '{{province_eastern_cape}}': 'X' if lif.province_code == '2' else '',
                '{{province_northern_cape}}': 'X' if lif.province_code == '3' else '',
                '{{province_free_state}}': 'X' if lif.province_code == '4' else '',
                '{{province_kwazulu_natal}}': 'X' if lif.province_code == '5' else '',
                '{{province_north_west}}': 'X' if lif.province_code == '6' else '',
                '{{province_gauteng_jhb}}': 'X' if lif.province_code == '7' else '',
                '{{province_gauteng_pta}}': 'X' if lif.province_code == '7b' else '',
                '{{province_mpumalanga}}': 'X' if lif.province_code == '8' else '',
                '{{province_limpopo}}': 'X' if lif.province_code == '9' else '',
                '{{province_outside_sa}}': 'X' if lif.province_code == 'X' else '',
                '{{province_national}}': 'X' if lif.province_code == 'N' else '',
                '{{national_id_checkbox}}': 'X' if (lif.alternative_id_type == '' or lif.alternative_id_type is None) else '',
                '{{alt_id_saqa}}':      'X' if lif.alternative_id_type == '521' else '',
                '{{alt_id_passport}}':  'X' if lif.alternative_id_type == '527' else '',
                '{{alt_id_driver}}':    'X' if lif.alternative_id_type == '529' else '',
                '{{alt_id_temp_id}}':   'X' if lif.alternative_id_type == '531' else '',
                '{{alt_id_none}}':      'X' if lif.alternative_id_type == '533' else '',
                '{{alt_id_unknown}}':   'X' if lif.alternative_id_type == '535' else '',
                '{{alt_id_student}}':   'X' if lif.alternative_id_type == '537' else '',
                '{{alt_id_work_permit}}':'X' if lif.alternative_id_type == '538' else '',
                '{{alt_id_employee}}':  'X' if lif.alternative_id_type == '539' else '',
                '{{alt_id_birth_cert}}':'X' if lif.alternative_id_type == '540' else '',
                '{{alt_id_hsrc}}':      'X' if lif.alternative_id_type == '541' else '',
                '{{alt_id_etqa}}':      'X' if lif.alternative_id_type == '561' else '',
                '{{alt_id_refugee}}':   'X' if lif.alternative_id_type == '565' else '',
                '{{home_language_afrikaans}}': 'X' if lif.home_language_code == 'Afr' else '',
                '{{home_language_english}}': 'X' if lif.home_language_code == 'Eng' else '',
                '{{home_language_ndebele}}': 'X' if lif.home_language_code == 'Nde' else '',
                '{{home_language_sepedi}}': 'X' if lif.home_language_code == 'Sep' else '',
                '{{home_language_sesotho}}': 'X' if lif.home_language_code == 'Ses' else '',
                '{{home_language_setswana}}': 'X' if lif.home_language_code == 'Set' else '',
                '{{home_language_siswati}}': 'X' if lif.home_language_code == 'Swa' else '',
                '{{home_language_tshivenda}}': 'X' if lif.home_language_code == 'Tsh' else '',
                '{{home_language_isixhosa}}': 'X' if lif.home_language_code == 'Xho' else '',
                '{{home_language_xitsonga}}': 'X' if lif.home_language_code == 'Xit' else '',
                '{{home_language_isizulu}}': 'X' if lif.home_language_code == 'Zul' else '',
                '{{home_language_sasl}}': 'X' if lif.home_language_code == 'SASL' else '',
                '{{home_language_other}}': 'X' if lif.home_language_code == 'Oth' else '',
                '{{home_language_unknown}}': 'X' if lif.home_language_code == 'U' else '',
                '{{tertiary_national_certificate}}': 'X' if lif.highest_tertiary_education == 'National Certificate' else '',
                '{{tertiary_national_diploma}}': 'X' if lif.highest_tertiary_education == 'National Diploma' else '',
                '{{tertiary_first_degree}}': 'X' if lif.highest_tertiary_education == 'National First Degree' else '',
                '{{tertiary_post_doctoral}}': 'X' if lif.highest_tertiary_education == 'Post-doctoral Degree' else '',
                '{{tertiary_doctoral}}': 'X' if lif.highest_tertiary_education == 'Doctoral Degree' else '',
                '{{tertiary_professional}}': 'X' if lif.highest_tertiary_education == 'Professional Qualification' else '',
                '{{tertiary_honours}}': 'X' if lif.highest_tertiary_education == 'Honours Degree' else '',
                '{{tertiary_higher_diploma}}': 'X' if lif.highest_tertiary_education == 'National Higher Diploma' else '',
                '{{tertiary_masters_diploma}}': 'X' if lif.highest_tertiary_education == 'National Masters Diploma' else '',
                '{{tertiary_national_higher}}': 'X' if lif.highest_tertiary_education == 'National Higher' else '',
                '{{tertiary_further_diploma}}': 'X' if lif.highest_tertiary_education == 'Further Diploma' else '',
                '{{tertiary_post_graduate}}': 'X' if lif.highest_tertiary_education == 'Post Graduate' else '',
                '{{tertiary_senior_certificate}}': 'X' if lif.highest_tertiary_education == 'Senior Certificate' else '',
                '{{tertiary_qual_nat_sen_cert}}': 'X' if lif.highest_tertiary_education == 'Qual at Nat Sen Cert' else '',
                '{{tertiary_apprenticeship}}': 'X' if lif.highest_tertiary_education == 'Apprenticeship' else '',
                '{{tertiary_post_grad_b_degree}}': 'X' if lif.highest_tertiary_education == 'Post Grad B Degree' else '',
                '{{tertiary_post_diploma_diploma}}': 'X' if lif.highest_tertiary_education == 'Post Diploma Diploma' else '',
                '{{tertiary_post_basic_diploma}}': 'X' if lif.highest_tertiary_education == 'Post-basic Diploma' else '',
                '{{secondary_grade_8}}':  'X' if lif.highest_secondary_education == 'Grade 8' else '',
                '{{secondary_grade_9}}':  'X' if lif.highest_secondary_education == 'Grade 9' else '',
                '{{secondary_grade_10}}': 'X' if lif.highest_secondary_education == 'Grade 10' else '',
                '{{secondary_grade_11}}': 'X' if lif.highest_secondary_education == 'Grade 11' else '',
                '{{secondary_grade_12}}': 'X' if lif.highest_secondary_education == 'Grade 12' else '',
                '{{nationality_sa}}':   'X' if lif.nationality_code == 'SA' else '',
                '{{nationality_sdc}}':  'X' if lif.nationality_code == 'SDC' else '',
                '{{nationality_ang}}':  'X' if lif.nationality_code == 'ANG' else '',
                '{{nationality_bot}}':  'X' if lif.nationality_code == 'BOT' else '',
                '{{nationality_les}}':  'X' if lif.nationality_code == 'LES' else '',
                '{{nationality_mal}}':  'X' if lif.nationality_code == 'MAL' else '',
                '{{nationality_mau}}':  'X' if lif.nationality_code == 'MAU' else '',
                '{{nationality_moz}}':  'X' if lif.nationality_code == 'MOZ' else '',
                '{{nationality_nam}}':  'X' if lif.nationality_code == 'NAM' else '',
                '{{nationality_sey}}':  'X' if lif.nationality_code == 'SEY' else '',
                '{{nationality_swa}}':  'X' if lif.nationality_code == 'SWA' else '',
                '{{nationality_tan}}':  'X' if lif.nationality_code == 'TAN' else '',
                '{{nationality_zai}}':  'X' if lif.nationality_code == 'ZAI' else '',
                '{{nationality_zam}}':  'X' if lif.nationality_code == 'ZAM' else '',
                '{{nationality_zim}}':  'X' if lif.nationality_code == 'ZIM' else '',
                '{{nationality_ais}}':  'X' if lif.nationality_code == 'AIS' else '',
                '{{nationality_aus}}':  'X' if lif.nationality_code == 'AUS' else '',
                '{{nationality_eur}}':  'X' if lif.nationality_code == 'EUR' else '',
                '{{nationality_nor}}':  'X' if lif.nationality_code == 'NOR' else '',
                '{{nationality_sou}}':  'X' if lif.nationality_code == 'SOU' else '',
                '{{nationality_roa}}':  'X' if lif.nationality_code == 'ROA' else '',
                '{{nationality_ooc}}':  'X' if lif.nationality_code == 'OOC' else '',
                '{{nationality_u}}':    'X' if lif.nationality_code == 'U' else '',
                '{{nationality_not}}':  'X' if lif.nationality_code == 'NOT' else '',
            })
            # Mapped fields
            for mapping in mappings:
                if mapping.placeholder in table_placeholders:
                    continue
                value = getattr(lif, mapping.lif_field, '')
                if mapping.lif_field == 'alternative_id_type':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, value) or ''
                elif mapping.lif_field == 'equity_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.EQUITY_CHOICES, value) or ''
                elif mapping.lif_field == 'nationality_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value) or ''
                elif mapping.lif_field == 'gender_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = dict([('F', 'Female'), ('M', 'Male')]).get(value, '')
                elif mapping.lif_field == 'citizen_resident_status_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.CITIZEN_STATUS_CHOICES, value) or ''
                elif mapping.lif_field == 'home_language_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value) or ''
                elif mapping.lif_field == 'province_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.PROVINCE_CHOICES, value) or ''
                elif mapping.lif_field == 'disability_status_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.DISABILITY_CHOICES, value) or ''
                elif mapping.lif_field == 'socio_economic_status_code':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, value) or ''
                elif mapping.lif_field == 'highest_tertiary_education':
                    context[mapping.placeholder] = value or ''
                    context[f'{mapping.placeholder}_name'] = get_choice_display(lif.TERTIARY_CHOICES, value) or ''
                elif mapping.lif_field in ['duration_start_date', 'duration_end_date'] and value:
                    context[mapping.placeholder] = value.strftime('%d/%m/%Y')
                elif mapping.lif_field == 'consent_to_process':
                    context[mapping.placeholder] = ' Agree' if lif.consent_to_process else ' Disagree'
                elif mapping.lif_field in ['years_in_occupation', 'secondary_year_completed', 'tertiary_year_completed']:
                    context[mapping.placeholder] = str(value) if value else ''
                elif mapping.lif_field in ['percentage_maths', 'percentage_first_language', 'percentage_second_language']:
                    context[mapping.placeholder] = str(value) if value else ''
                else:
                    context[mapping.placeholder] = value or ''
                if isinstance(value, str) and value and mapping.lif_field not in ['duration_start_date', 'duration_end_date'] and mapping.placeholder not in table_placeholders:
                    base_match = re.match(r"\{\{(\w+)\}\}", mapping.placeholder)
                    if base_match:
                        base = base_match.group(1)
                        for i, char in enumerate(value):
                            context[f"{{{{{base}_{i}}}}}"] = char
                        for i in range(len(value), 20):
                            context[f"{{{{{base}_{i}}}}}"] = ''
            return context
        def replace_placeholders_in_paragraph(paragraph, context):
            full_text = ''.join(run.text for run in paragraph.runs)
            replaced = False
            for placeholder, value in context.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
                    replaced = True
            if replaced and paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ''

        def replace_placeholders_in_table(table, context):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholders_in_paragraph(p, context)
                    # Recursively process nested tables
                    for nested_table in cell.tables:
                        replace_placeholders_in_table(nested_table, context)

        # --- Single or Bulk ---
        if len(lif_ids) == 1:
            lif = get_object_or_404(LIF, id=lif_ids[0])
            doc = Document(template.template_file.path)
            context = build_context(lif)
            for p in doc.paragraphs:
                replace_placeholders_in_paragraph(p, context)
            for table in doc.tables:
                replace_placeholders_in_table(table, context)
            safe_template = "".join(x for x in template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
            safe_first = "".join(x for x in (lif.learner_first_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
            safe_last = "".join(x for x in (lif.learner_last_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
            filename = f"{safe_template}_{safe_first}_{safe_last}.docx"
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            doc.save(response)
            return response
        else:
            # Bulk ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for lif_id in lif_ids:
                    lif = get_object_or_404(LIF, id=lif_id)
                    doc = Document(template.template_file.path)
                    context = build_context(lif)
                    for p in doc.paragraphs:
                        replace_placeholders_in_paragraph(p, context)
                    for table in doc.tables:
                        replace_placeholders_in_table(table, context)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    safe_template = "".join(x for x in template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
                    safe_first = "".join(x for x in (lif.learner_first_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
                    safe_last = "".join(x for x in (lif.learner_last_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
                    filename = f"{safe_template}_{safe_first}_{safe_last}.docx"
                    zip_file.writestr(filename, doc_io.read())
            zip_buffer.seek(0)
            safe_template = "".join(x for x in template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
            response = HttpResponse(zip_buffer, content_type="application/zip")
            response["Content-Disposition"] = f"attachment; filename={safe_template}.zip"
            return response

    return render(request, 'core/generate_lif_word.html', {
        'lifs': lifs_page,
        'templates': templates,
        'filters': filters,
        'allow_zip': allow_zip,
    })


from django.http import JsonResponse
from core.models import LIF

from django.http import JsonResponse
from core.models import LIF
from django.db.models import Q

def lif_list_json(request):
    q = request.GET.get('q', '').strip()
    lifs = LIF.objects.all()
    if q:
        lifs = lifs.filter(
            Q(learner_first_name__icontains=q) |
            Q(learner_last_name__icontains=q) |
            Q(national_id__icontains=q) |
            Q(email_address__icontains=q)
        )
    data = []
    for lif in lifs[:100]:  # Limit to 100 results for speed
        data.append({
            'id': lif.id,
            'first_name': lif.learner_first_name,
            'last_name': lif.learner_last_name,
            'national_id': lif.national_id,
            'gender': lif.get_gender_code_display(),
            'equity': lif.get_equity_code_display(),
            'email': lif.email_address,
            'birth_date': lif.learner_birth_date.strftime('%Y-%m-%d') if lif.learner_birth_date else '',
        })
    return JsonResponse({'lifs': data})

from django.shortcuts import render, redirect, get_object_or_404
from django.forms import modelformset_factory
from django.contrib import messages
from django.http import HttpResponse
from .models import LIFTemplate, LIFTemplateFieldMap, LIF
from .forms import LIFTemplateUploadForm, LIFTemplateFieldMapForm
from docx import Document
from django.shortcuts import render, get_object_or_404, redirect
from django.forms import modelformset_factory
from django.contrib import messages
from .models import LIFTemplate, LIFTemplateFieldMap
from .forms import LIFTemplateFieldMapForm

from django.views.decorators.csrf import csrf_exempt
from django.db import transaction
from django.http import JsonResponse
import json

@csrf_exempt
def venuebooking_switch(request):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Invalid method"}, status=405)
    try:
        data = json.loads(request.body)
        booking_id = data.get("booking_id")
        new_venue_id = data.get("new_venue_id")
        if not booking_id or not new_venue_id:
            return JsonResponse({"success": False, "error": "Missing data"}, status=400)
        from .models import VenueBooking, Venue
        booking = VenueBooking.objects.get(pk=booking_id)
        new_venue = Venue.objects.get(pk=new_venue_id)
        if "virtual venue" in new_venue.name.lower():
            return JsonResponse({"success": False, "error": "Cannot switch to a virtual venue."}, status=400)
        with transaction.atomic():
            # Find conflicting booking (if any)
            conflict = VenueBooking.objects.filter(
                venue=new_venue,
                start_datetime=booking.start_datetime,
                end_datetime=booking.end_datetime,
                status__in=['booked', 'rescheduled']
            ).exclude(pk=booking.pk).first()
            if conflict:
                # Save conflict data
                conflict_data = {
                    'session_date': conflict.session_date,
                    'venue': booking.venue,  # swap to old venue
                    'start_datetime': conflict.start_datetime,
                    'end_datetime': conflict.end_datetime,
                    'booking_purpose': conflict.booking_purpose,
                    'facilitator': conflict.facilitator,
                    'status': conflict.status,
                    'num_learners': conflict.num_learners,
                    'num_learners_lunch': conflict.num_learners_lunch,
                }
                conflict.delete()
                booking.venue = new_venue
                booking.save()
                VenueBooking.objects.create(**conflict_data)
            else:
                booking.venue = new_venue
                booking.save()
        return JsonResponse({"success": True})
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=500)
    

from .models import SessionDate, ProjectPlan, Group, SLA_Qualifications

def get_default_num_learners_for_session(session_date_id):
    try:
        session = SessionDate.objects.select_related('project_plan__group').get(id=session_date_id)
        group = session.project_plan.group
        # Get all SLA_Qualifications linked to this group
        sla_quals = group.sla_qualifications.all()
        # Sum their learner_count
        total = sum(q.learner_count for q in sla_quals if q.learner_count)
        return total if total > 0 else None
    except Exception:
        return None
    
from .models import ExternalProject
from django.views.generic import ListView

from django.db import connections
from django.shortcuts import render

def external_project_list(request):
    with connections['mssql'].cursor() as cursor:
        cursor.execute("""
            SELECT TOP (100)
                [Id], [Title], [ProjectCode], [SETA], [StartDate], [EndDate], [ProjectLeadId], [Created], [Modified],
                [CreatedBy], [ModifiedBy], [LANumber], [ProviderCode], [ProviderETQAID], [QualificationId],
                [AssessorOfAllRemediationDate], [AssessorOfAllRemediationPersonId], [FileUploadToSETADate],
                [FileUploadToSETAPersonId], [FinalRemidiationBriefingSessionEndDate], [FinalRemidiationBriefingSessionPersonId],
                [FinalRemidiationBriefingSessionStartDate], [IntroductionSessionDate], [ModerationOfFilesDate],
                [ModerationOfFilesPersonId], [ProjectCloseOutReportDate], [ProjectCloseOutReportPersonId],
                [ProjectPlanDocumentId], [ProjectPlanPdfId], [AttendanceOrCompetence], [CertificateProvince], [IssueDate],
                [ProjectPlanVersion], [Status], [CurrentStatus], [CurrentStatusColour], [ProjectInductionDate],
                [ETQAAdministratorId], [ETQAManagerId], [LeadFacilitatorId], [OpsManagerId], [ProjectManagerId], [ProjectType],
                [AssessmentBookInDate], [ConfirmationOfRegistrationDate], [FinalInternalModerationBookInDate],
                [FinalInternalModerationBookOutDate], [FinalInternalModerationDueDate], [FinalInternalModerationPersonId],
                [FinalModerationReportReviewDate], [FinalModerationReportReviewPersonId], [FinalRemediationBookOutDate],
                [FinalRemediationSubmissionDueDate], [FinalResultsReleasedDueDate], [InternalModerationBookInDate],
                [InternalModerationBookOutDate], [InternalModerationDueDate], [InternalModerationPersonId],
                [ModerationReportReviewDate], [ModerationReportReviewPersonId], [ProvisionalResultsLetterDueDate],
                [QualificationRegistrationUploadDueDate], [SETACertificateDueDate], [SETACertificateDueDate2],
                [SETAExternalVerificationDueDate], [SETARegistrationConfirmationDueDate], [SETARegistrationSubmissionDueDate],
                [SETAVerificationDueDate], [TLOCertificateDueDate], [TLOCertificateDueDate2], [TLOProvisionalResultsLetterDueDate],
                [FinalRemediationBookInDate]
            FROM [dbo].[Projects]
        """)
        columns = [col[0] for col in cursor.description]
        projects = [dict(zip(columns, row)) for row in cursor.fetchall()]
    return render(request, 'core/external_project_list.html', {'projects': projects})
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from .forms import VenueBookingForm
from .models import SessionDate, VenueBooking, ProjectPlan

from django.views import View
from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from .forms import VenueBookingForm
from .models import SessionDate, VenueBooking, ProjectPlan, Venue

from django.views import View
from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from .forms import VenueBookingForm
from .models import SessionDate, VenueBooking, ProjectPlan, Venue
import uuid
from django.utils import timezone

from datetime import datetime, time
from django.utils import timezone
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.http import JsonResponse, HttpResponse
import uuid

class VenueBookingModalFormView(RolePermissionRequiredMixin, View):
    def get(self, request):
        booking_id = request.GET.get('booking')
        session_id = request.GET.get('session')
        date = request.GET.get('date')
        session_ids = request.GET.get('session_ids')  # For multi-booking
        form = None
        multi_booking_mode = False
        selected_session_ids = []

        if booking_id:
            booking = get_object_or_404(VenueBooking, pk=booking_id)
            obj = booking
            # Try combined_booking_id first, else match on start/end/venue
            if getattr(booking, 'combined_booking_id', None):
                bookings = VenueBooking.objects.filter(combined_booking_id=booking.combined_booking_id)
            else:
                bookings = VenueBooking.objects.filter(
                    venue=booking.venue,
                    start_datetime=booking.start_datetime,
                    end_datetime=booking.end_datetime,
                    status__in=['booked', 'rescheduled']
                )
            session_ids_list = list(bookings.values_list('session_date_id', flat=True))
            session_dates = SessionDate.objects.filter(id__in=session_ids_list).select_related('project_plan', 'project_plan__group', 'project_plan__module')
            form = VenueBookingForm(instance=booking)
            if 'session_dates' in form.fields:
                form.fields['session_dates'].queryset = session_dates
            # Pass session_ids to context so template/POST works for multi-booking
            session_ids = ",".join(str(sid) for sid in session_ids_list)
            if len(session_ids_list) > 1:
                multi_booking_mode = True
                selected_session_ids = [str(sid) for sid in session_ids_list]
        else:
            initial = {}
            session_dates = SessionDate.objects.select_related('project_plan', 'project_plan__group', 'project_plan__module').all()
            obj = None

            # Single session booking
            if session_id:
                try:
                    session = SessionDate.objects.get(id=session_id)
                    initial['session_date'] = session
                    initial['start_datetime'] = timezone.make_aware(datetime.combine(session.start_date, time(8, 0)))
                    initial['end_datetime'] = timezone.make_aware(datetime.combine(session.end_date, time(17, 0)))
                    initial['booking_purpose'] = f"{session.project_plan.group.name} - {session.project_plan.module.name}"
                    # Try to get default num_learners
                    default_learners = get_default_num_learners_for_session(session_id)
                    if default_learners:
                        initial['num_learners'] = default_learners
                except SessionDate.DoesNotExist:
                    pass

            if date and not initial.get('start_datetime'):
                try:
                    date_obj = datetime.strptime(date, '%Y-%m-%d').date()
                    initial['start_datetime'] = timezone.make_aware(datetime.combine(date_obj, time(8, 0)))
                    initial['end_datetime'] = timezone.make_aware(datetime.combine(date_obj, time(17, 0)))
                except ValueError:
                    pass

            form = VenueBookingForm(initial=initial)

            # Multi-booking: pre-select session_dates and filter options
            if session_ids:
                ids = [int(sid) for sid in session_ids.split(',') if sid.isdigit()]
                session_dates = SessionDate.objects.filter(id__in=ids).select_related('project_plan', 'project_plan__group', 'project_plan__module')
                if 'session_dates' in form.fields:
                    form.fields['session_dates'].queryset = session_dates
                multi_booking_mode = True
                selected_session_ids = [str(sid) for sid in ids]
            else:
                if 'session_dates' in form.fields:
                    form.fields['session_dates'].queryset = session_dates

        project_plans = {pp.id: str(pp) for pp in ProjectPlan.objects.all().order_by('id')}
        
        # Convert all sessions to JSON-serializable format for search functionality
        all_sessions_queryset = SessionDate.objects.select_related(
            'project_plan', 'project_plan__group', 'project_plan__module'
        ).all().order_by('-start_date')
        
        # Convert to JSON-serializable format for JavaScript
        all_sessions = []
        for session in all_sessions_queryset:
            all_sessions.append({
                'id': session.id,
                'text': str(session),
                'searchText': str(session).lower()
            })
        
        # IMPORTANT: Import json to properly serialize the data
        import json
        
        html = render_to_string('core/venuebooking_modal_form.html', {
            'form': form,
            'session_dates': session_dates,
            'all_sessions': json.dumps(all_sessions),  # Properly serialize for JavaScript
            'project_plans': project_plans,
            'object': obj,
            'session_ids': session_ids if booking_id else session_ids if session_ids else "",
            'multi_booking_mode': multi_booking_mode,
            'selected_session_ids': selected_session_ids,
        }, request=request)
        return HttpResponse(html)

    def post(self, request):
        booking_id = request.GET.get('booking')
        session_id = request.GET.get('session')
        session_ids = request.GET.get('session_ids')  # For multi-booking
        confirm_override = request.POST.get('confirm_override')

        if booking_id:
            booking = get_object_or_404(VenueBooking, pk=booking_id)
            form = VenueBookingForm(request.POST, instance=booking)
        else:
            form = VenueBookingForm(request.POST)

        # Always filter the queryset for session_dates field if multi-booking
        if session_ids and 'session_dates' in form.fields:
            ids = [int(sid) for sid in session_ids.split(',') if sid.isdigit()]
            session_dates = SessionDate.objects.filter(id__in=ids).select_related('project_plan', 'project_plan__group', 'project_plan__module')
            form.fields['session_dates'].queryset = session_dates
        else:
            session_dates = SessionDate.objects.select_related('project_plan', 'project_plan__group', 'project_plan__module').all()
            if 'session_dates' in form.fields:
                form.fields['session_dates'].queryset = session_dates

        # --- STRICT CONFLICT CHECK LOGIC ---
        # Only for single booking (not multi-booking)
        if not session_ids:
            venue = form.data.get('venue')
            session_date = form.data.get('session_date')
            start_datetime = form.data.get('start_datetime')
            end_datetime = form.data.get('end_datetime')
            # Parse datetimes if needed
            try:
                start_dt = timezone.make_aware(datetime.fromisoformat(start_datetime)) if start_datetime else None
                end_dt = timezone.make_aware(datetime.fromisoformat(end_datetime)) if end_datetime else None
            except Exception:
                start_dt = end_dt = None

            # Find existing booking for this venue, start, end (regardless of session), exclude self if editing
            conflict_qs = VenueBooking.objects.filter(
                venue_id=venue,
                start_datetime=start_dt,
                end_datetime=end_dt,
            )
            if booking_id:
                conflict_qs = conflict_qs.exclude(pk=booking_id)
            existing_booking = conflict_qs.first()

            if existing_booking and not confirm_override:
                # Render only the conflict info (not the form)
                conflict_html = render_to_string('core/venuebooking_conflict_popup.html', {
                    'existing_booking': existing_booking,
                }, request=request)
                return JsonResponse({
                    'conflict': True,
                    'html': conflict_html
                })

        if form.is_valid():
            # Handle override after form is valid so cleaned_data is available
            if not session_ids:
                venue = form.data.get('venue')
                session_date = form.data.get('session_date')
                start_datetime = form.data.get('start_datetime')
                end_datetime = form.data.get('end_datetime')
                try:
                    start_dt = timezone.make_aware(datetime.fromisoformat(start_datetime)) if start_datetime else None
                    end_dt = timezone.make_aware(datetime.fromisoformat(end_datetime)) if end_datetime else None
                except Exception:
                    start_dt = end_dt = None
                conflict_qs = VenueBooking.objects.filter(
                    venue_id=venue,
                    start_datetime=start_dt,
                    end_datetime=end_dt,
                )
                if booking_id:
                    conflict_qs = conflict_qs.exclude(pk=booking_id)
                existing_booking = conflict_qs.first()
                if existing_booking and confirm_override == "yes":
                    previous_user = existing_booking.user
                    previous_user_email = previous_user.email if previous_user else "brendonmandlandlovu@gmail.com"
                    previous_user_name = f"{previous_user.first_name} {previous_user.last_name}".strip() if previous_user else "Unknown"
                    overridden_by = request.user
                    overridden_by_name = f"{overridden_by.first_name} {overridden_by.last_name}".strip() or overridden_by.username
                    venue_name = str(existing_booking.venue)
                    session_date_str = str(existing_booking.session_date)
                    start_str = existing_booking.start_datetime.strftime("%Y-%m-%d %H:%M")
                    end_str = existing_booking.end_datetime.strftime("%Y-%m-%d %H:%M")
                    purpose = str(existing_booking.booking_purpose)
                    facilitator = str(existing_booking.facilitator) if existing_booking.facilitator else "N/A"
                    new_start = form.cleaned_data.get('start_datetime')
                    new_end = form.cleaned_data.get('end_datetime')
                    new_purpose = form.cleaned_data.get('booking_purpose')
                    new_facilitator = str(form.cleaned_data.get('facilitator')) if form.cleaned_data.get('facilitator') else "N/A"
                    new_session_date = ""
                    if form.cleaned_data.get('session_date'):
                        try:
                            new_session_obj = SessionDate.objects.get(pk=form.cleaned_data['session_date'].id if hasattr(form.cleaned_data['session_date'], 'id') else form.cleaned_data['session_date'])
                            new_session_date = str(new_session_obj)
                        except Exception:
                            new_session_date = str(form.cleaned_data.get('session_date'))
                    existing_booking.delete()

                    subject = f"Venue Booking Overridden: {venue_name} ({session_date_str})"
                    message = (
                        f"Dear {previous_user_name},\n\n"
                        f"Your venue booking has been overridden by another user. Please see the details below.\n\n"
                        f"--- Previous Booking (Yours) ---\n"
                        f"Venue: {venue_name}\n"
                        f"Session: {session_date_str}\n"
                        f"Start: {start_str}\n"
                        f"End: {end_str}\n"
                        f"Purpose: {purpose}\n"
                        f"Facilitator: {facilitator}\n\n"
                        f"--- New Booking (Override) ---\n"
                        f"Booked By: {overridden_by_name} (username: {overridden_by.username}, email: {overridden_by.email})\n"
                        f"Session: {new_session_date or session_date_str}\n"
                        f"Start: {new_start}\n"
                        f"End: {new_end}\n"
                        f"Purpose: {new_purpose}\n"
                        f"Facilitator: {new_facilitator or 'N/A'}\n\n"
                        f"If you have any questions, please contact your administrator.\n"
                    )

                    send_mail(
                        subject=subject,
                        message=message,
                        from_email="noreply@ensemble.com",
                        recipient_list=["brendonmandlandlovu@gmail.com"],  # [previous_user_email]
                        fail_silently=True,
                    )
                    # To notify the user directly, uncomment below:
                    # send_mail(
                    #     subject=subject,
                    #     message=message,
                    #     from_email="noreply@ensemble.com",
                    #     recipient_list=[previous_user_email],
                    #     fail_silently=True,
                    # )

            session_dates_selected = form.cleaned_data.get('session_dates')
            # MULTI-BOOKING: create a booking for each session
            if session_dates_selected and len(session_dates_selected) > 1:
                combined_id = str(uuid.uuid4())
                for session in session_dates_selected:
                    VenueBooking.objects.create(
                        session_date=session,
                        venue=form.cleaned_data['venue'],
                        start_datetime=timezone.make_aware(datetime.combine(session.start_date, time(8, 0))),
                        end_datetime=timezone.make_aware(datetime.combine(session.end_date, time(17, 0))),
                        booking_purpose=form.cleaned_data['booking_purpose'],
                        facilitator=form.cleaned_data['facilitator'],
                        status=form.cleaned_data['status'],
                        num_learners=form.cleaned_data['num_learners'],
                        num_learners_lunch=form.cleaned_data['num_learners_lunch'],
                        combined_booking_id=combined_id,
                        user=request.user  # <-- Ensure user is set for multi-booking!
                    )
                return JsonResponse({'success': True})
            # SINGLE BOOKING
            booking = form.save(commit=False)
            if not booking.venue and session_id:
                try:
                    session = SessionDate.objects.get(pk=session_id)
                    if getattr(session, 'preferred_training_methodology', None) and "virtual session" in session.preferred_training_methodology.lower():
                        virtual_venue = Venue.objects.filter(name__istartswith="Virtual Session").order_by('name').first()
                        if virtual_venue:
                            booking.venue = virtual_venue
                except SessionDate.DoesNotExist:
                    pass
            booking.user = request.user  # <-- Ensure user is set for single booking!
            booking.save()
            return JsonResponse({'success': True})
        else:
            # For error redisplay, filter session_dates if multi-booking
            if session_ids:
                ids = [int(sid) for sid in session_ids.split(',') if sid.isdigit()]
                session_dates = SessionDate.objects.filter(id__in=ids).select_related('project_plan', 'project_plan__group', 'project_plan__module')
                if 'session_dates' in form.fields:
                    form.fields['session_dates'].queryset = session_dates
            else:
                session_dates = SessionDate.objects.select_related('project_plan', 'project_plan__group', 'project_plan__module').all()
                if 'session_dates' in form.fields:
                    form.fields['session_dates'].queryset = session_dates
            project_plans = {pp.id: str(pp) for pp in ProjectPlan.objects.all().order_by('id')}
            # Determine multi_booking_mode and selected_session_ids for error redisplay
            multi_booking_mode = bool(session_ids and len(session_ids.split(',')) > 1)
            selected_session_ids = [sid for sid in session_ids.split(',') if sid] if session_ids else []
            html = render_to_string('core/venuebooking_modal_form.html', {
                'form': form,
                'session_dates': session_dates,
                'project_plans': project_plans,
                'object': booking if booking_id else None,
                'multi_booking_mode': multi_booking_mode,
                'selected_session_ids': selected_session_ids,
            }, request=request)
            return JsonResponse({'success': False, 'html': html})

import io
import zipfile
import re
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from docx import Document
from datetime import date
from core.models import LIF, LIFTemplate, LIFTemplateFieldMap

def generate_bulk_lif_zip(request):
    import re
    from datetime import date

    if request.method == "POST":
        learner_ids = request.POST.getlist("learner_ids")
        template_id = request.POST.get("template_id")
        if not learner_ids or not template_id:
            return HttpResponse("No learners or template selected.", status=400)

        template = LIFTemplate.objects.get(id=template_id)
        mappings = LIFTemplateFieldMap.objects.filter(template=template)
        zip_buffer = io.BytesIO()

        def get_choice_display(choices, code):
            return dict(choices).get(code, '')

        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            for learner_id in learner_ids:
                try:
                    lif = LIF.objects.get(learner_id=learner_id)
                except LIF.DoesNotExist:
                    continue

                doc = Document(template.template_file.path)

                # Build context (all fields, code and display, checkboxes, date splits)
                context = {
                    '{{learner_first_name}}': lif.learner_first_name or '',
                    '{{learner_middle_name}}': lif.learner_middle_name or '',
                    '{{learner_last_name}}': lif.learner_last_name or '',
                    '{{national_id}}': lif.national_id or '',
                    '{{alternative_id}}': lif.alternative_id or '',
                    '{{alternative_id_type}}': lif.alternative_id_type or '',
                    '{{alternative_id_type_name}}': get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, lif.alternative_id_type) or '',
                    '{{learner_title}}': lif.learner_title or '',
                    '{{learner_birth_date}}': lif.learner_birth_date.strftime('%Y-%m-%d') if lif.learner_birth_date else '',
                    '{{learner_previous_last_name}}': lif.learner_previous_last_name or '',
                    '{{learner_current_occupation}}': lif.learner_current_occupation or '',
                    '{{years_in_occupation}}': str(lif.years_in_occupation) if lif.years_in_occupation else '',
                    '{{employer}}': lif.employer or '',
                    '{{highest_secondary_education}}': lif.highest_secondary_education or '',
                    '{{secondary_school_name}}': lif.secondary_school_name or '',
                    '{{secondary_year_completed}}': str(lif.secondary_year_completed) if lif.secondary_year_completed else '',
                    '{{percentage_maths}}': str(lif.percentage_maths) if lif.percentage_maths else '',
                    '{{percentage_first_language}}': str(lif.percentage_first_language) if lif.percentage_first_language else '',
                    '{{percentage_second_language}}': str(lif.percentage_second_language) if lif.percentage_second_language else '',
                    '{{highest_tertiary_education}}': lif.highest_tertiary_education or '',
                    '{{highest_tertiary_education_name}}': get_choice_display(lif.TERTIARY_CHOICES, lif.highest_tertiary_education) or '',
                    '{{tertiary_school_name}}': lif.tertiary_school_name or '',
                    '{{tertiary_year_completed}}': str(lif.tertiary_year_completed) if lif.tertiary_year_completed else '',
                    '{{address_line1}}': lif.address_line1 or '',
                    '{{address_line2}}': lif.address_line2 or '',
                    '{{city}}': lif.city or '',
                    '{{state_province}}': lif.state_province or '',
                    '{{postal_code}}': lif.postal_code or '',
                    '{{phone_number}}': lif.phone_number or '',
                    '{{alt_contact_number}}': lif.alt_contact_number or '',
                    '{{email_address}}': lif.email_address or '',
                    '{{provider_etqa_id}}': lif.provider_etqa_id or '',
                    '{{provider_code}}': lif.provider_code or '',
                    '{{programme_title}}': lif.programme_title or '',
                    '{{qualification_code}}': lif.qualification_code or '',
                    '{{nqf_level}}': lif.nqf_level or '',
                    '{{sponsor}}': lif.sponsor or '',
                    '{{commencement_date}}': lif.duration_start_date.strftime('%d/%m/%Y') if lif.duration_start_date else '',
                    '{{termination_date}}': lif.duration_end_date.strftime('%d/%m/%Y') if lif.duration_end_date else '',
                    '{{consent_to_process}}': ' Agree' if lif.consent_to_process else ' Disagree',
                    '{{equity_code}}': lif.equity_code or '',
                    '{{equity_name}}': get_choice_display(lif.EQUITY_CHOICES, lif.equity_code) or '',
                    '{{nationality_code}}': lif.nationality_code or '',
                    '{{nationality_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.nationality_code) or '',
                    '{{gender_code}}': lif.gender_code or '',
                    '{{gender_name}}': dict([('F', 'Female'), ('M', 'Male')]).get(lif.gender_code, ''),
                    '{{citizen_resident_status_code}}': lif.citizen_resident_status_code or '',
                    '{{citizen_resident_status_name}}': get_choice_display(lif.CITIZEN_STATUS_CHOICES, lif.citizen_resident_status_code) or '',
                    '{{home_language_code}}': lif.home_language_code or '',
                    '{{home_language_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.home_language_code) or '',
                    '{{province_code}}': lif.province_code or '',
                    '{{province_name}}': get_choice_display(lif.PROVINCE_CHOICES, lif.province_code) or '',
                    '{{disability_status_code}}': lif.disability_status_code or '',
                    '{{disability_status_name}}': get_choice_display(lif.DISABILITY_CHOICES, lif.disability_status_code) or '',
                    '{{socio_economic_status_code}}': lif.socio_economic_status_code or '',
                    '{{socio_economic_status_name}}': get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, lif.socio_economic_status_code) or '',
                }

                # Add split fields
                if lif.learner_birth_date:
                    birth_str = lif.learner_birth_date.strftime('%Y%m%d')
                    for i in range(8):
                        context[f'{{{{learner_birth_date_{i}}}}}'] = birth_str[i] if i < len(birth_str) else ''
                else:
                    for i in range(8):
                        context[f'{{{{learner_birth_date_{i}}}}}'] = ''
                if lif.national_id:
                    for i in range(13):
                        context[f'{{{{national_id_{i}}}}}'] = lif.national_id[i] if i < len(lif.national_id) else ''
                else:
                    for i in range(13):
                        context[f'{{{{national_id_{i}}}}}'] = ''
                if lif.duration_start_date:
                    start_str = lif.duration_start_date.strftime('%d%m%Y')
                    for i in range(8):
                        context[f'{{{{commencement_date_{i}}}}}'] = start_str[i] if i < len(start_str) else ''
                else:
                    for i in range(8):
                        context[f'{{{{commencement_date_{i}}}}}'] = ''
                if lif.duration_end_date:
                    end_str = lif.duration_end_date.strftime('%d%m%Y')
                    for i in range(8):
                        context[f'{{{{termination_date_{i}}}}}'] = end_str[i] if i < len(end_str) else ''
                else:
                    for i in range(8):
                        context[f'{{{{termination_date_{i}}}}}'] = ''

                # Table-based "checkbox" fields (X for selected)
                context.update({
                    '{{gender_male}}': 'X' if lif.gender_code == 'M' else '',
                    '{{gender_female}}': 'X' if lif.gender_code == 'F' else '',
                    '{{below_35_yes}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year < 35) else '',
                    '{{below_35_no}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year >= 35) else '',
                    '{{equity_african}}': 'X' if lif.equity_code == 'BA' else '',
                    '{{equity_coloured}}': 'X' if lif.equity_code == 'BC' else '',
                    '{{equity_indian}}': 'X' if lif.equity_code == 'BI' else '',
                    '{{equity_white}}': 'X' if lif.equity_code == 'Wh' else '',
                    '{{disability_sight}}': 'X' if lif.disability_status_code == '01' else '',
                    '{{disability_hearing}}': 'X' if lif.disability_status_code == '02' else '',
                    '{{disability_communication}}': 'X' if lif.disability_status_code == '03' else '',
                    '{{disability_physical}}': 'X' if lif.disability_status_code == '04' else '',
                    '{{disability_intellectual}}': 'X' if lif.disability_status_code == '05' else '',
                    '{{disability_emotional}}': 'X' if lif.disability_status_code == '06' else '',
                    '{{disability_multiple}}': 'X' if lif.disability_status_code == '07' else '',
                    '{{disability_unspecified}}': 'X' if lif.disability_status_code == '09' else '',
                    '{{disability_none}}': 'X' if lif.disability_status_code == 'N' else '',
                    '{{citizen_sa}}': 'X' if lif.citizen_resident_status_code == 'SA' else '',
                    '{{citizen_dual}}': 'X' if lif.citizen_resident_status_code == 'D' else '',
                    '{{citizen_other}}': 'X' if lif.citizen_resident_status_code == 'O' else '',
                    '{{citizen_permanent}}': 'X' if lif.citizen_resident_status_code == 'PR' else '',
                    '{{citizen_unknown}}': 'X' if lif.citizen_resident_status_code == 'U' else '',
                    '{{socio_employed}}': 'X' if lif.socio_economic_status_code == '01' else '',
                    '{{socio_unemployed_seeking}}': 'X' if lif.socio_economic_status_code == '02' else '',
                    '{{socio_not_working_not_looking}}': 'X' if lif.socio_economic_status_code == '03' else '',
                    '{{socio_homemaker}}': 'X' if lif.socio_economic_status_code == '04' else '',
                    '{{socio_student}}': 'X' if lif.socio_economic_status_code == '06' else '',
                    '{{socio_pensioner}}': 'X' if lif.socio_economic_status_code == '07' else '',
                    '{{socio_disabled}}': 'X' if lif.socio_economic_status_code == '08' else '',
                    '{{socio_no_wish_to_work}}': 'X' if lif.socio_economic_status_code == '09' else '',
                    '{{socio_not_working_nec}}': 'X' if lif.socio_economic_status_code == '10' else '',
                    '{{socio_aged_under_15}}': 'X' if lif.socio_economic_status_code == '97' else '',
                    '{{socio_institution}}': 'X' if lif.socio_economic_status_code == '98' else '',
                    '{{socio_unspecified}}': 'X' if lif.socio_economic_status_code == 'U' else '',
                    '{{tertiary_national_certificate}}': 'X' if lif.highest_tertiary_education == 'National Certificate' else '',
                    '{{tertiary_national_diploma}}': 'X' if lif.highest_tertiary_education == 'National Diploma' else '',
                    '{{tertiary_first_degree}}': 'X' if lif.highest_tertiary_education == 'National First Degree' else '',
                    '{{tertiary_post_doctoral}}': 'X' if lif.highest_tertiary_education == 'Post-doctoral Degree' else '',
                    '{{tertiary_doctoral}}': 'X' if lif.highest_tertiary_education == 'Doctoral Degree' else '',
                    '{{tertiary_professional}}': 'X' if lif.highest_tertiary_education == 'Professional Qualification' else '',
                    '{{tertiary_honours}}': 'X' if lif.highest_tertiary_education == 'Honours Degree' else '',
                    '{{tertiary_higher_diploma}}': 'X' if lif.highest_tertiary_education == 'National Higher Diploma' else '',
                    '{{tertiary_masters_diploma}}': 'X' if lif.highest_tertiary_education == 'National Masters Diploma' else '',
                    '{{tertiary_national_higher}}': 'X' if lif.highest_tertiary_education == 'National Higher' else '',
                    '{{tertiary_further_diploma}}': 'X' if lif.highest_tertiary_education == 'Further Diploma' else '',
                    '{{tertiary_post_graduate}}': 'X' if lif.highest_tertiary_education == 'Post Graduate' else '',
                    '{{tertiary_senior_certificate}}': 'X' if lif.highest_tertiary_education == 'Senior Certificate' else '',
                    '{{tertiary_qual_nat_sen_cert}}': 'X' if lif.highest_tertiary_education == 'Qual at Nat Sen Cert' else '',
                    '{{tertiary_apprenticeship}}': 'X' if lif.highest_tertiary_education == 'Apprenticeship' else '',
                    '{{tertiary_post_grad_b_degree}}': 'X' if lif.highest_tertiary_education == 'Post Grad B Degree' else '',
                    '{{tertiary_post_diploma_diploma}}': 'X' if lif.highest_tertiary_education == 'Post Diploma Diploma' else '',
                    '{{tertiary_post_basic_diploma}}': 'X' if lif.highest_tertiary_education == 'Post-basic Diploma' else '',
                    '{{province_western_cape}}': 'X' if lif.province_code == '1' else '',
                    '{{province_eastern_cape}}': 'X' if lif.province_code == '2' else '',
                    '{{province_northern_cape}}': 'X' if lif.province_code == '3' else '',
                    '{{province_free_state}}': 'X' if lif.province_code == '4' else '',
                    '{{province_kwazulu_natal}}': 'X' if lif.province_code == '5' else '',
                    '{{province_north_west}}': 'X' if lif.province_code == '6' else '',
                    '{{province_gauteng_jhb}}': 'X' if lif.province_code == '7' else '',
                    '{{province_gauteng_pta}}': 'X' if lif.province_code == '7b' else '',
                    '{{province_mpumalanga}}': 'X' if lif.province_code == '8' else '',
                    '{{province_limpopo}}': 'X' if lif.province_code == '9' else '',
                    '{{province_outside_sa}}': 'X' if lif.province_code == 'X' else '',
                    '{{province_national}}': 'X' if lif.province_code == 'N' else '',
                    '{{national_id_checkbox}}': 'X' if (lif.alternative_id_type == '' or lif.alternative_id_type is None) else '',
                    '{{alt_id_saqa}}':      'X' if lif.alternative_id_type == '521' else '',
                    '{{alt_id_passport}}':  'X' if lif.alternative_id_type == '527' else '',
                    '{{alt_id_driver}}':    'X' if lif.alternative_id_type == '529' else '',
                    '{{alt_id_temp_id}}':   'X' if lif.alternative_id_type == '531' else '',
                    '{{alt_id_none}}':      'X' if lif.alternative_id_type == '533' else '',
                    '{{alt_id_unknown}}':   'X' if lif.alternative_id_type == '535' else '',
                    '{{alt_id_student}}':   'X' if lif.alternative_id_type == '537' else '',
                    '{{alt_id_work_permit}}':'X' if lif.alternative_id_type == '538' else '',
                    '{{alt_id_employee}}':  'X' if lif.alternative_id_type == '539' else '',
                    '{{alt_id_birth_cert}}':'X' if lif.alternative_id_type == '540' else '',
                    '{{alt_id_hsrc}}':      'X' if lif.alternative_id_type == '541' else '',
                    '{{alt_id_etqa}}':      'X' if lif.alternative_id_type == '561' else '',
                    '{{alt_id_refugee}}':   'X' if lif.alternative_id_type == '565' else '',
                    '{{home_language_afrikaans}}': 'X' if lif.home_language_code == 'Afr' else '',
                    '{{home_language_english}}': 'X' if lif.home_language_code == 'Eng' else '',
                    '{{home_language_ndebele}}': 'X' if lif.home_language_code == 'Nde' else '',
                    '{{home_language_sepedi}}': 'X' if lif.home_language_code == 'Sep' else '',
                    '{{home_language_sesotho}}': 'X' if lif.home_language_code == 'Ses' else '',
                    '{{home_language_setswana}}': 'X' if lif.home_language_code == 'Set' else '',
                    '{{home_language_siswati}}': 'X' if lif.home_language_code == 'Swa' else '',
                    '{{home_language_tshivenda}}': 'X' if lif.home_language_code == 'Tsh' else '',
                    '{{home_language_isixhosa}}': 'X' if lif.home_language_code == 'Xho' else '',
                    '{{home_language_xitsonga}}': 'X' if lif.home_language_code == 'Xit' else '',
                    '{{home_language_isizulu}}': 'X' if lif.home_language_code == 'Zul' else '',
                    '{{home_language_sasl}}': 'X' if lif.home_language_code == 'SASL' else '',
                    '{{home_language_other}}': 'X' if lif.home_language_code == 'Oth' else '',
                    '{{home_language_unknown}}': 'X' if lif.home_language_code == 'U' else '',
                    '{{tertiary_national_certificate}}': 'X' if lif.highest_tertiary_education == 'National Certificate' else '',
                    '{{tertiary_national_diploma}}': 'X' if lif.highest_tertiary_education == 'National Diploma' else '',
                    '{{tertiary_first_degree}}': 'X' if lif.highest_tertiary_education == 'National First Degree' else '',
                    '{{tertiary_post_doctoral}}': 'X' if lif.highest_tertiary_education == 'Post-doctoral Degree' else '',
                    '{{tertiary_doctoral}}': 'X' if lif.highest_tertiary_education == 'Doctoral Degree' else '',
                    '{{tertiary_professional}}': 'X' if lif.highest_tertiary_education == 'Professional Qualification' else '',
                    '{{tertiary_honours}}': 'X' if lif.highest_tertiary_education == 'Honours Degree' else '',
                    '{{tertiary_higher_diploma}}': 'X' if lif.highest_tertiary_education == 'National Higher Diploma' else '',
                    '{{tertiary_masters_diploma}}': 'X' if lif.highest_tertiary_education == 'National Masters Diploma' else '',
                    '{{tertiary_national_higher}}': 'X' if lif.highest_tertiary_education == 'National Higher' else '',
                    '{{tertiary_further_diploma}}': 'X' if lif.highest_tertiary_education == 'Further Diploma' else '',
                    '{{tertiary_post_graduate}}': 'X' if lif.highest_tertiary_education == 'Post Graduate' else '',
                    '{{tertiary_senior_certificate}}': 'X' if lif.highest_tertiary_education == 'Senior Certificate' else '',
                    '{{tertiary_qual_nat_sen_cert}}': 'X' if lif.highest_tertiary_education == 'Qual at Nat Sen Cert' else '',
                    '{{tertiary_apprenticeship}}': 'X' if lif.highest_tertiary_education == 'Apprenticeship' else '',
                    '{{tertiary_post_grad_b_degree}}': 'X' if lif.highest_tertiary_education == 'Post Grad B Degree' else '',
                    '{{tertiary_post_diploma_diploma}}': 'X' if lif.highest_tertiary_education == 'Post Diploma Diploma' else '',
                    '{{tertiary_post_basic_diploma}}': 'X' if lif.highest_tertiary_education == 'Post-basic Diploma' else '',
                    '{{secondary_grade_8}}':  'X' if lif.highest_secondary_education == 'Grade 8' else '',
                    '{{secondary_grade_9}}':  'X' if lif.highest_secondary_education == 'Grade 9' else '',
                    '{{secondary_grade_10}}': 'X' if lif.highest_secondary_education == 'Grade 10' else '',
                    '{{secondary_grade_11}}': 'X' if lif.highest_secondary_education == 'Grade 11' else '',
                    '{{secondary_grade_12}}': 'X' if lif.highest_secondary_education == 'Grade 12' else '',
                    '{{nationality_sa}}':   'X' if lif.nationality_code == 'SA' else '',
                    '{{nationality_sdc}}':  'X' if lif.nationality_code == 'SDC' else '',
                    '{{nationality_ang}}':  'X' if lif.nationality_code == 'ANG' else '',
                    '{{nationality_bot}}':  'X' if lif.nationality_code == 'BOT' else '',
                    '{{nationality_les}}':  'X' if lif.nationality_code == 'LES' else '',
                    '{{nationality_mal}}':  'X' if lif.nationality_code == 'MAL' else '',
                    '{{nationality_mau}}':  'X' if lif.nationality_code == 'MAU' else '',
                    '{{nationality_moz}}':  'X' if lif.nationality_code == 'MOZ' else '',
                    '{{nationality_nam}}':  'X' if lif.nationality_code == 'NAM' else '',
                    '{{nationality_sey}}':  'X' if lif.nationality_code == 'SEY' else '',
                    '{{nationality_swa}}':  'X' if lif.nationality_code == 'SWA' else '',
                    '{{nationality_tan}}':  'X' if lif.nationality_code == 'TAN' else '',
                    '{{nationality_zai}}':  'X' if lif.nationality_code == 'ZAI' else '',
                    '{{nationality_zam}}':  'X' if lif.nationality_code == 'ZAM' else '',
                    '{{nationality_zim}}':  'X' if lif.nationality_code == 'ZIM' else '',
                    '{{nationality_ais}}':  'X' if lif.nationality_code == 'AIS' else '',
                    '{{nationality_aus}}':  'X' if lif.nationality_code == 'AUS' else '',
                    '{{nationality_eur}}':  'X' if lif.nationality_code == 'EUR' else '',
                    '{{nationality_nor}}':  'X' if lif.nationality_code == 'NOR' else '',
                    '{{nationality_sou}}':  'X' if lif.nationality_code == 'SOU' else '',
                    '{{nationality_roa}}':  'X' if lif.nationality_code == 'ROA' else '',
                    '{{nationality_ooc}}':  'X' if lif.nationality_code == 'OOC' else '',
                    '{{nationality_u}}':    'X' if lif.nationality_code == 'U' else '',
                    '{{nationality_not}}':  'X' if lif.nationality_code == 'NOT' else '',
                })

                table_placeholders = [
                    '{{gender_male}}', '{{gender_female}}', '{{below_35_yes}}', '{{below_35_no}}',
                    '{{equity_african}}', '{{equity_coloured}}', '{{equity_indian}}', '{{equity_white}}',
                    '{{disability_sight}}', '{{disability_hearing}}', '{{disability_communication}}',
                    '{{disability_physical}}', '{{disability_intellectual}}', '{{disability_emotional}}',
                    '{{disability_multiple}}', '{{disability_unspecified}}', '{{disability_none}}',
                    '{{citizen_sa}}', '{{citizen_dual}}', '{{citizen_other}}', '{{citizen_permanent}}',
                    '{{citizen_unknown}}' , '{{socio_employed}}',                # 01
                    '{{socio_unemployed_seeking}}',      # 02
                    '{{socio_not_working_not_looking}}', # 03
                    '{{socio_homemaker}}',               # 04
                    '{{socio_student}}',                 # 06
                    '{{socio_pensioner}}',               # 07
                    '{{socio_disabled}}',                # 08
                    '{{socio_no_wish_to_work}}',         # 09
                    '{{socio_not_working_nec}}',         # 10
                    '{{socio_aged_under_15}}',           # 97
                    '{{socio_institution}}',             # 98
                    '{{socio_unspecified}}',             # U
                    '{{tertiary_national_certificate}}',
                    '{{tertiary_national_diploma}}',
                    '{{tertiary_first_degree}}',
                    '{{tertiary_post_doctoral}}',
                    '{{tertiary_doctoral}}',
                    '{{tertiary_professional}}',
                    '{{tertiary_honours}}',
                    '{{tertiary_higher_diploma}}',
                    '{{tertiary_masters_diploma}}',
                    '{{tertiary_national_higher}}',
                    '{{tertiary_further_diploma}}',
                    '{{tertiary_post_graduate}}',
                    '{{tertiary_senior_certificate}}',
                    '{{tertiary_qual_nat_sen_cert}}',
                    '{{tertiary_apprenticeship}}',
                    '{{tertiary_post_grad_b_degree}}',
                    '{{tertiary_post_diploma_diploma}}',
                    '{{tertiary_post_basic_diploma}}',
                    '{{province_western_cape}}',
                    '{{province_eastern_cape}}',
                    '{{province_northern_cape}}',
                    '{{province_free_state}}',
                    '{{province_kwazulu_natal}}',
                    '{{province_north_west}}',
                    '{{province_gauteng_jhb}}',
                    '{{province_gauteng_pta}}',
                    '{{province_mpumalanga}}',
                    '{{province_limpopo}}',
                    '{{province_outside_sa}}',
                    '{{province_national}}',
                    '{{national_id_checkbox}}',  # <-- Add this line for National ID checkbox
                    '{{alt_id_saqa}}',
                    '{{alt_id_passport}}',
                    '{{alt_id_driver}}',
                    '{{alt_id_temp_id}}',
                    '{{alt_id_none}}',
                    '{{alt_id_unknown}}',
                    '{{alt_id_student}}',
                    '{{alt_id_work_permit}}',
                    '{{alt_id_employee}}',
                    '{{alt_id_birth_cert}}',
                    '{{alt_id_hsrc}}',
                    '{{alt_id_etqa}}',
                    '{{alt_id_refugee}}',
                    '{{national_id_checkbox}}',  # <-- Add this line for National ID checkbox
                    '{{alt_id_saqa}}',
                    '{{alt_id_passport}}',
                    '{{alt_id_driver}}',
                    '{{alt_id_temp_id}}',
                    '{{alt_id_none}}',
                    '{{alt_id_unknown}}',
                    '{{alt_id_student}}',
                    '{{alt_id_work_permit}}',
                    '{{alt_id_employee}}',
                    '{{alt_id_birth_cert}}',
                    '{{alt_id_hsrc}}',
                    '{{alt_id_etqa}}',
                    '{{alt_id_refugee}}',
                    '{{home_language_afrikaans}}',
                    '{{home_language_english}}',
                    '{{home_language_ndebele}}',
                    '{{home_language_sepedi}}',
                    '{{home_language_sesotho}}',
                    '{{home_language_setswana}}',
                    '{{home_language_siswati}}',
                    '{{home_language_tshivenda}}',
                    '{{home_language_isixhosa}}',
                    '{{home_language_xitsonga}}',
                    '{{home_language_isizulu}}',
                    '{{home_language_sasl}}',
                    '{{home_language_other}}',
                    '{{home_language_unknown}}',
                    '{{tertiary_national_certificate}}',
                    '{{tertiary_national_diploma}}',
                    '{{tertiary_first_degree}}',
                    '{{tertiary_post_doctoral}}',
                    '{{tertiary_doctoral}}',
                    '{{tertiary_professional}}',
                    '{{tertiary_honours}}',
                    '{{tertiary_higher_diploma}}',
                    '{{tertiary_masters_diploma}}',
                    '{{tertiary_national_higher}}',
                    '{{tertiary_further_diploma}}',
                    '{{tertiary_post_graduate}}',
                    '{{tertiary_senior_certificate}}',
                    '{{tertiary_qual_nat_sen_cert}}',
                    '{{tertiary_apprenticeship}}',
                    '{{tertiary_post_grad_b_degree}}',
                    '{{tertiary_post_diploma_diploma}}',
                    '{{tertiary_post_basic_diploma}}',
                    '{{secondary_grade_8}}',
                    '{{secondary_grade_9}}',
                    '{{secondary_grade_10}}',
                    '{{secondary_grade_11}}',
                    '{{secondary_grade_12}}',
                    '{{nationality_sa}}',
                    '{{nationality_sdc}}',
                    '{{nationality_ang}}',
                    '{{nationality_bot}}',
                    '{{nationality_les}}',
                    '{{nationality_mal}}',
                    '{{nationality_mau}}',
                    '{{nationality_moz}}',
                    '{{nationality_nam}}',
                    '{{nationality_sey}}',
                    '{{nationality_swa}}',
                    '{{nationality_tan}}',
                    '{{nationality_zai}}',
                    '{{nationality_zam}}',
                    '{{nationality_zim}}',
                    '{{nationality_ais}}',
                    '{{nationality_aus}}',
                    '{{nationality_eur}}',
                    '{{nationality_nor}}',
                    '{{nationality_sou}}',
                    '{{nationality_roa}}',
                    '{{nationality_ooc}}',
                    '{{nationality_u}}',
                    '{{nationality_not}}',
                ]

                for mapping in mappings:
                    if mapping.placeholder in table_placeholders:
                        continue
                    value = getattr(lif, mapping.lif_field, '')
                    if mapping.lif_field == 'alternative_id_type':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, value) or ''
                    elif mapping.lif_field == 'equity_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.EQUITY_CHOICES, value) or ''
                    elif mapping.lif_field == 'nationality_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value) or ''
                    elif mapping.lif_field == 'gender_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = dict([('F', 'Female'), ('M', 'Male')]).get(value, '')
                    elif mapping.lif_field == 'citizen_resident_status_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.CITIZEN_STATUS_CHOICES, value) or ''
                    elif mapping.lif_field == 'home_language_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value) or ''
                    elif mapping.lif_field == 'province_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.PROVINCE_CHOICES, value) or ''
                    elif mapping.lif_field == 'disability_status_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.DISABILITY_CHOICES, value) or ''
                    elif mapping.lif_field == 'socio_economic_status_code':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, value) or ''
                    elif mapping.lif_field == 'highest_tertiary_education':
                        context[mapping.placeholder] = value or ''
                        context[f'{mapping.placeholder}_name'] = get_choice_display(lif.TERTIARY_CHOICES, value) or ''
                    elif mapping.lif_field in ['duration_start_date', 'duration_end_date'] and value:
                        context[mapping.placeholder] = value.strftime('%d/%m/%Y')
                    elif mapping.lif_field == 'consent_to_process':
                        context[mapping.placeholder] = ' Agree' if lif.consent_to_process else ' Disagree'
                    elif mapping.lif_field in ['years_in_occupation', 'secondary_year_completed', 'tertiary_year_completed']:
                        context[mapping.placeholder] = str(value) if value else ''
                    elif mapping.lif_field in ['percentage_maths', 'percentage_first_language', 'percentage_second_language']:
                        context[mapping.placeholder] = str(value) if value else ''
                    else:
                        context[mapping.placeholder] = value or ''
                    if isinstance(value, str) and value and mapping.lif_field not in ['duration_start_date', 'duration_end_date'] and mapping.placeholder not in table_placeholders:
                        base_match = re.match(r"\{\{(\w+)\}\}", mapping.placeholder)
                        if base_match:
                            base = base_match.group(1)
                            for i, char in enumerate(value):
                                context[f"{{{{{base}_{i}}}}}"] = char
                            for i in range(len(value), 20):
                                context[f"{{{{{base}_{i}}}}}"] = ''

                def replace_placeholders_in_paragraph(paragraph, context):
                    full_text = ''.join(run.text for run in paragraph.runs)
                    replaced = False
                    for placeholder, value in context.items():
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, str(value))
                            replaced = True
                    if replaced and paragraph.runs:
                        paragraph.runs[0].text = full_text
                        for run in paragraph.runs[1:]:
                            run.text = ''

                def replace_placeholders_in_table(table, context):
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_placeholders_in_paragraph(p, context)
                            for nested_table in cell.tables:
                                replace_placeholders_in_table(nested_table, context)

                for p in doc.paragraphs:
                    replace_placeholders_in_paragraph(p, context)
                for table in doc.tables:
                    replace_placeholders_in_table(table, context)

                # --- Filename: <TemplateName>_<FirstName>_<LastName>.docx ---
                safe_template = "".join(x for x in template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
                safe_first = "".join(x for x in (lif.learner_first_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
                safe_last = "".join(x for x in (lif.learner_last_name or "") if x.isalnum() or x in (' ', '_', '-')).rstrip()
                lif_filename = f"{safe_template}_{safe_first}_{safe_last}.docx"
                lif_buffer = io.BytesIO()
                doc.save(lif_buffer)
                zip_file.writestr(lif_filename, lif_buffer.getvalue())

        zip_buffer.seek(0)
        safe_template = "".join(x for x in template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
        response = HttpResponse(zip_buffer, content_type="application/zip")
        response["Content-Disposition"] = f"attachment; filename={safe_template}.zip"
        return response

    return HttpResponse("Invalid request", status=405)

import os
import json
from django.conf import settings
from django.shortcuts import render, redirect
from django.contrib import messages
from django.views import View

COGNITO_DATA_DIR = os.path.join(settings.BASE_DIR, 'ensemble_app', 'core', 'cognito_data')

class CognitoLIFEntriesView(View):
    template_name = 'core/cognito_lif_entries.html'
    
    def get(self, request):
        files = []
        if os.path.exists(COGNITO_DATA_DIR):
            files = sorted([f for f in os.listdir(COGNITO_DATA_DIR) if f.endswith('.json')], reverse=True)
        selected_file = request.GET.get('file')
        entries = []
        page_obj = None
        if selected_file and selected_file in files:
            with open(os.path.join(COGNITO_DATA_DIR, selected_file), 'r') as f:
                entries = json.load(f)
            # --- FILTERING ---
            filter_id = request.GET.get('filter_id', '').strip()
            filter_name = request.GET.get('filter_name', '').strip().lower()
            if filter_id:
                entries = [e for e in entries if filter_id in (e.get('NationalID') or '')]
            if filter_name:
                def normalize(s):
                    return unicodedata.normalize('NFKD', s or '').encode('ascii', 'ignore').decode('ascii').lower()
                entries = [
                    e for e in entries
                    if filter_name in normalize(
                        f"{e.get('LearnerTitle','')} {e.get('LearnerFirstName','')} {e.get('LearnerMiddleName','')} {e.get('LearnerLastName','')}"
                    )
                ]
            # --- SORTING ---
            sort_entry = request.GET.get('sort_entry')
            if sort_entry == 'asc':
                entries = sorted(entries, key=lambda e: (e.get('Entry', {}).get('Number') or 0))
            elif sort_entry == 'desc':
                entries = sorted(entries, key=lambda e: (e.get('Entry', {}).get('Number') or 0), reverse=True)
            # --- PAGINATION ---
            paginator = Paginator(entries, 5)
            page_number = request.GET.get('page')
            try:
                page_obj = paginator.page(page_number)
            except PageNotAnInteger:
                page_obj = paginator.page(1)
            except EmptyPage:
                page_obj = paginator.page(paginator.num_pages)
        return render(request, self.template_name, {
            'files': files,
            'selected_file': selected_file,
            'entries': page_obj if page_obj else [],
            'page_obj': page_obj,
            'is_paginated': bool(page_obj and page_obj.has_other_pages()),
        })

    def post(self, request):
        # Trigger fetch via management command
        form_id = request.POST.get('form_id')
        entry_id = request.POST.get('entry_id')
        max_entry_id = request.POST.get('max_entry_id')
        args = ['fetch_cognito_lif', '--form', form_id]
        if entry_id:
            args += ['--entry', entry_id]
        elif max_entry_id:
            args += ['--max', max_entry_id]
        else:
            messages.error(request, "Provide Entry ID or Max Entry ID.")
            return redirect(request.path)
        # Run management command
        from django.core.management import call_command
        call_command(*args)
        messages.success(request, "Cognito fetch complete. Refresh to see new files.")
        return redirect(request.path)
    
import os
import json
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.views import View
from .models import LIF, Learner

@method_decorator(login_required, name='dispatch')
class ExportCognitoToLIFView(View):
    def post(self, request):
        selected_file = request.POST.get('file')
        if not selected_file:
            messages.error(request, "No file selected.")
            return redirect('cognito_lif_entries')
        file_path = os.path.join(COGNITO_DATA_DIR, selected_file)
        if not os.path.exists(file_path):
            messages.error(request, "File not found.")
            return redirect('cognito_lif_entries')

        with open(file_path, 'r', encoding='utf-8') as f:
            entries = json.load(f)

        linked_learners = []
        created, updated, skipped = 0, 0, 0

        # Helper to clamp decimals
        def clamp_decimal(val, maxval=999.99):
            try:
                v = float(val)
                return min(v, maxval)
            except Exception:
                return None

        # Helper to truncate strings to 255 chars
        def truncate_str(val, maxlen=255):
            return val[:maxlen] if isinstance(val, str) and len(val) > maxlen else val

        for entry in entries:
            # Extract code part for coded fields
            def code(val):
                if val is None:
                    return None
                return val.split(' ')[0] if ' ' in val else val

            national_id = entry.get('NationalID')
            if not national_id:
                continue

            # Try to find existing LIF
            lif = LIF.objects.filter(national_id=national_id).first()
            # Try to find matching Learner
            learner = Learner.objects.filter(IDNumber=national_id).first()

            # Prepare LIF data
            lif_data = {
                'national_id': entry.get('NationalID'),
                'alternative_id': entry.get('AlternativeID'),
                'alternative_id_type': entry.get('AlternativeIDType'),
                'equity_code': code(entry.get('EquityCode')),
                'nationality_code': code(entry.get('NationalityCode')),
                'gender_code': code(entry.get('GenderCode')),
                'citizen_resident_status_code': code(entry.get('CitizenResidentStatusCode')),
                'home_language_code': code(entry.get('HomeLanguageCode')),
                'province_code': code(entry.get('ProvinceCode')),
                'disability_status_code': code(entry.get('DisabiltyStatusCode')),
                'socio_economic_status_code': code(entry.get('SocioeconomicStatusCode')),
                'learner_title': entry.get('LearnerTitle'),
                'learner_birth_date': entry.get('LearnerBirthDate') or None,
                'learner_first_name': entry.get('LearnerFirstName'),
                'learner_middle_name': entry.get('LearnerMiddleName'),
                'learner_last_name': entry.get('LearnerLastName'),
                'learner_previous_last_name': entry.get('LearnerPreviousLastName'),
                'learner_current_occupation': entry.get('LearnerCurrentOccupation'),
                'years_in_occupation': entry.get('YearsInOccupation') if entry.get('YearsInOccupation') and str(entry.get('YearsInOccupation')).isdigit() else None,
                'employer': entry.get('Employer'),
                'highest_secondary_education': entry.get('HighestSecondaryEducation'),
                'secondary_school_name': entry.get('SchoolName'),
                'secondary_year_completed': entry.get('YearCompleted') if entry.get('YearCompleted') and str(entry.get('YearCompleted')).isdigit() else None,
                'percentage_maths': entry.get('PercentageMaths') or None,
                'percentage_first_language': entry.get('Percentage1stLanguage') or None,
                'percentage_second_language': entry.get('Percentage2ndLanguage') or None,
                'highest_tertiary_education': entry.get('HighestTertiaryEducation'),
                'tertiary_school_name': entry.get('SchoolName2'),
                'tertiary_year_completed': entry.get('YearCompleted2') if entry.get('YearCompleted2') and str(entry.get('YearCompleted2')).isdigit() else None,
                'address_line1': entry.get('LearnerHomeAddress', {}).get('Line1') if entry.get('LearnerHomeAddress') else None,
                'address_line2': entry.get('LearnerHomeAddress', {}).get('Line2') if entry.get('LearnerHomeAddress') else None,
                'city': entry.get('LearnerHomeAddress', {}).get('City') if entry.get('LearnerHomeAddress') else None,
                'state_province': entry.get('LearnerHomeAddress', {}).get('State') if entry.get('LearnerHomeAddress') else None,
                'postal_code': entry.get('LearnerHomeAddress', {}).get('PostalCode') if entry.get('LearnerHomeAddress') else None,
                'phone_number': entry.get('LearnerPhoneNumber'),
                'alt_contact_number': entry.get('AltenativeContactNumber'),
                'email_address': entry.get('LearnerEmailAddress'),
                'provider_etqa_id': entry.get('ProviderETQAID'),
                'provider_code': entry.get('ProviderCode'),
                'programme_title': entry.get('USLearningProgrammeCourseQualificationTitle'),
                'qualification_code': entry.get('QualificationCode'),
                'nqf_level': entry.get('NQFLevel'),
                'sponsor': entry.get('Sponsor'),
                'duration_start_date': entry.get('DurationStartDate') or None,
                'duration_end_date': entry.get('DurationEndDate') or None,
                'consent_to_process': entry.get('IHerebyConsentToTheLearningOrganisationCollectingStoringAndProcessingMyInformationForThePurposeOfRegisteringMeWithTheRelevantSETA', False),
            }

            # Convert dates
            from datetime import datetime
            for date_field in ['learner_birth_date', 'duration_start_date', 'duration_end_date']:
                if lif_data[date_field]:
                    try:
                        lif_data[date_field] = datetime.strptime(lif_data[date_field], "%Y-%m-%d").date()
                    except Exception:
                        lif_data[date_field] = None

            # Clamp decimals
            for dec_field in ['percentage_maths', 'percentage_first_language', 'percentage_second_language']:
                if lif_data[dec_field] is not None:
                    lif_data[dec_field] = clamp_decimal(lif_data[dec_field])

            # Truncate all CharFields to 255 chars (if not None)
            char_fields = [
                'national_id', 'alternative_id', 'learner_title', 'learner_first_name', 'learner_middle_name',
                'learner_last_name', 'learner_previous_last_name', 'learner_current_occupation', 'employer',
                'highest_secondary_education', 'secondary_school_name', 'highest_tertiary_education',
                'tertiary_school_name', 'address_line1', 'address_line2', 'city', 'state_province', 'postal_code',
                'phone_number', 'alt_contact_number', 'email_address', 'provider_etqa_id', 'provider_code',
                'programme_title', 'qualification_code', 'nqf_level', 'sponsor'
            ]
            for field in char_fields:
                if lif_data.get(field):
                    lif_data[field] = truncate_str(lif_data[field], 255)

            # Link learner if found
            if learner:
                lif_data['learner'] = learner
                linked_learners.append(f"{learner.FirstName} {learner.Surname} ({learner.IDNumber})")

            # Create or update LIF
            if lif:
                changed = False
                for k, v in lif_data.items():
                    if getattr(lif, k, None) != v:
                        setattr(lif, k, v)
                        changed = True
                if changed:
                    lif.save()
                    updated += 1
                else:
                    skipped += 1
            else:
                lif = LIF.objects.create(**lif_data)
                created += 1

        messages.success(request, f"Export complete: {created} created, {updated} updated, {skipped} skipped.")
        if linked_learners:
            messages.info(request, "Linked learners:<br>" + "<br>".join(linked_learners))
        return redirect('cognito_lif_entries')
    
from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import LIFTemplate

class LIFTemplateDeleteView(DeleteView):
    model = LIFTemplate
    template_name = 'core/lif_template_confirm_delete.html'
    success_url = reverse_lazy('lif_template_list')
import io
import zipfile
import re
from docx import Document
from datetime import date
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from .models import Group, Learner, LIF, LIFTemplate, LIFTemplateFieldMap, LearnerQualification, Service
from copy import deepcopy
import logging

logger = logging.getLogger(__name__)

def insert_lif_content_after_first_section(doc, lif_doc):
    """
    Inserts the content of lif_doc into doc after the first section break (cover page).
    If no section break is found, inserts after the first paragraph.
    This version is robust for most real-world DOCX files.
    """
    body = doc.element.body
    insert_idx = None
    for idx, p in enumerate(body):
        if p.tag.endswith('}p'):
            if p.findall('.//w:sectPr', namespaces=body.nsmap):
                insert_idx = idx + 1
                break
    if insert_idx is None:
        insert_idx = 1
    for i, element in enumerate(list(lif_doc.element.body)):
        body.insert(insert_idx + i, deepcopy(element))

@csrf_exempt
def generate_bulk_admin_pack_zip(request):
    if request.method != "POST":
        return HttpResponse("Invalid request method", status=405)

    group_id = request.POST.get("group_id")
    learner_ids = request.POST.getlist("learner_ids")
    template_id = request.POST.get("template_id")
    admin_pack_id = request.POST.get("admin_pack_id")

    if not all([group_id, learner_ids, template_id, admin_pack_id]):
        return HttpResponse("Missing required parameters: group_id, learner_ids, template_id, or admin_pack_id", status=400)

    try:
        group = get_object_or_404(Group, id=group_id)
        lif_template = get_object_or_404(LIFTemplate, id=template_id)
        service = get_object_or_404(Service, id=admin_pack_id)
        mappings = LIFTemplateFieldMap.objects.filter(template=lif_template)
    except Exception as e:
        logger.error(f"Error fetching group, template, or service: {str(e)}")
        return HttpResponse("Error fetching group, template, or service", status=400)

    def get_choice_display(choices, code):
        return dict(choices).get(code, '')

    table_placeholders = [
        '{{gender_male}}', '{{gender_female}}', '{{below_35_yes}}', '{{below_35_no}}',
        '{{equity_african}}', '{{equity_coloured}}', '{{equity_indian}}', '{{equity_white}}',
        '{{disability_sight}}', '{{disability_hearing}}', '{{disability_communication}}',
        '{{disability_physical}}', '{{disability_intellectual}}', '{{disability_emotional}}',
        '{{disability_multiple}}', '{{disability_unspecified}}', '{{disability_none}}',
        '{{citizen_sa}}', '{{citizen_dual}}', '{{citizen_other}}', '{{citizen_permanent}}',
        '{{citizen_unknown}}', '{{socio_employed}}', '{{socio_unemployed_seeking}}',
        '{{socio_not_working_not_looking}}', '{{socio_homemaker}}', '{{socio_student}}',
        '{{socio_pensioner}}', '{{socio_disabled}}', '{{socio_no_wish_to_work}}',
        '{{socio_not_working_nec}}', '{{socio_aged_under_15}}', '{{socio_institution}}',
        '{{socio_unspecified}}', '{{tertiary_national_certificate}}', '{{tertiary_national_diploma}}',
        '{{tertiary_first_degree}}', '{{tertiary_post_doctoral}}', '{{tertiary_doctoral}}',
        '{{tertiary_professional}}', '{{tertiary_honours}}', '{{tertiary_higher_diploma}}',
        '{{tertiary_masters_diploma}}', '{{tertiary_national_higher}}', '{{tertiary_further_diploma}}',
        '{{tertiary_post_graduate}}', '{{tertiary_senior_certificate}}', '{{tertiary_qual_nat_sen_cert}}',
        '{{tertiary_apprenticeship}}', '{{tertiary_post_grad_b_degree}}', '{{tertiary_post_diploma_diploma}}',
        '{{tertiary_post_basic_diploma}}', '{{province_western_cape}}', '{{province_eastern_cape}}',
        '{{province_northern_cape}}', '{{province_free_state}}', '{{province_kwazulu_natal}}',
        '{{province_north_west}}', '{{province_gauteng_jhb}}', '{{province_gauteng_pta}}',
        '{{province_mpumalanga}}', '{{province_limpopo}}', '{{province_outside_sa}}',
        '{{province_national}}', '{{national_id_checkbox}}', '{{alt_id_saqa}}', '{{alt_id_passport}}',
        '{{alt_id_driver}}', '{{alt_id_temp_id}}', '{{alt_id_none}}', '{{alt_id_unknown}}',
        '{{alt_id_student}}', '{{alt_id_work_permit}}', '{{alt_id_employee}}', '{{alt_id_birth_cert}}',
        '{{alt_id_hsrc}}', '{{alt_id_etqa}}', '{{alt_id_refugee}}', '{{home_language_afrikaans}}',
        '{{home_language_english}}', '{{home_language_ndebele}}', '{{home_language_sepedi}}',
        '{{home_language_sesotho}}', '{{home_language_setswana}}', '{{home_language_siswati}}',
        '{{home_language_tshivenda}}', '{{home_language_isixhosa}}', '{{home_language_xitsonga}}',
        '{{home_language_isizulu}}', '{{home_language_sasl}}', '{{home_language_other}}',
        '{{home_language_unknown}}', '{{secondary_grade_8}}', '{{secondary_grade_9}}',
        '{{secondary_grade_10}}', '{{secondary_grade_11}}', '{{secondary_grade_12}}',
        '{{nationality_sa}}', '{{nationality_sdc}}', '{{nationality_ang}}', '{{nationality_bot}}',
        '{{nationality_les}}', '{{nationality_mal}}', '{{nationality_mau}}', '{{nationality_moz}}',
        '{{nationality_nam}}', '{{nationality_sey}}', '{{nationality_swa}}', '{{nationality_tan}}',
        '{{nationality_zai}}', '{{nationality_zam}}', '{{nationality_zim}}', '{{nationality_ais}}',
        '{{nationality_aus}}', '{{nationality_eur}}', '{{nationality_nor}}', '{{nationality_sou}}',
        '{{nationality_roa}}', '{{nationality_ooc}}', '{{nationality_u}}', '{{nationality_not}}',
    ]

    def build_context(lif):
        context = {
            '{{learner_first_name}}': lif.learner_first_name or '',
            '{{learner_middle_name}}': lif.learner_middle_name or '',
            '{{learner_last_name}}': lif.learner_last_name or '',
            '{{national_id}}': lif.national_id or '',
            '{{alternative_id}}': lif.alternative_id or '',
            '{{alternative_id_type}}': lif.alternative_id_type or '',
            '{{alternative_id_type_name}}': get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, lif.alternative_id_type) or '',
            '{{learner_title}}': lif.learner_title or '',
            '{{learner_birth_date}}': lif.learner_birth_date.strftime('%Y-%m-%d') if lif.learner_birth_date else '',
            '{{learner_previous_last_name}}': lif.learner_previous_last_name or '',
            '{{learner_current_occupation}}': lif.learner_current_occupation or '',
            '{{years_in_occupation}}': str(lif.years_in_occupation) if lif.years_in_occupation else '',
            '{{employer}}': lif.employer or '',
            '{{highest_secondary_education}}': lif.highest_secondary_education or '',
            '{{secondary_school_name}}': lif.secondary_school_name or '',
            '{{secondary_year_completed}}': str(lif.secondary_year_completed) if lif.secondary_year_completed else '',
            '{{percentage_maths}}': str(lif.percentage_maths) if lif.percentage_maths else '',
            '{{percentage_first_language}}': str(lif.percentage_first_language) if lif.percentage_first_language else '',
            '{{percentage_second_language}}': str(lif.percentage_second_language) if lif.percentage_second_language else '',
            '{{highest_tertiary_education}}': lif.highest_tertiary_education or '',
            '{{highest_tertiary_education_name}}': get_choice_display(lif.TERTIARY_CHOICES, lif.highest_tertiary_education) or '',
            '{{tertiary_school_name}}': lif.tertiary_school_name or '',
            '{{tertiary_year_completed}}': str(lif.tertiary_year_completed) if lif.tertiary_year_completed else '',
            '{{address_line1}}': lif.address_line1 or '',
            '{{address_line2}}': lif.address_line2 or '',
            '{{city}}': lif.city or '',
            '{{state_province}}': lif.state_province or '',
            '{{postal_code}}': lif.postal_code or '',
            '{{phone_number}}': lif.phone_number or '',
            '{{alt_contact_number}}': lif.alt_contact_number or '',
            '{{email_address}}': lif.email_address or '',
            '{{provider_etqa_id}}': lif.provider_etqa_id or '',
            '{{provider_code}}': lif.provider_code or '',
            '{{programme_title}}': lif.programme_title or '',
            '{{qualification_code}}': lif.qualification_code or '',
            '{{nqf_level}}': lif.nqf_level or '',
            '{{sponsor}}': lif.sponsor or '',
            '{{commencement_date}}': lif.duration_start_date.strftime('%d/%m/%Y') if lif.duration_start_date else '',
            '{{termination_date}}': lif.duration_end_date.strftime('%d/%m/%Y') if lif.duration_end_date else '',
            '{{consent_to_process}}': ' Agree' if lif.consent_to_process else ' Disagree',
            '{{equity_code}}': lif.equity_code or '',
            '{{equity_name}}': get_choice_display(lif.EQUITY_CHOICES, lif.equity_code) or '',
            '{{nationality_code}}': lif.nationality_code or '',
            '{{nationality_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.nationality_code) or '',
            '{{gender_code}}': lif.gender_code or '',
            '{{gender_name}}': dict([('F', 'Female'), ('M', 'Male')]).get(lif.gender_code, ''),
            '{{citizen_resident_status_code}}': lif.citizen_resident_status_code or '',
            '{{citizen_resident_status_name}}': get_choice_display(lif.CITIZEN_STATUS_CHOICES, lif.citizen_resident_status_code) or '',
            '{{home_language_code}}': lif.home_language_code or '',
            '{{home_language_name}}': get_choice_display(lif.NATIONALITY_CHOICES, lif.home_language_code) or '',
            '{{province_code}}': lif.province_code or '',
            '{{province_name}}': get_choice_display(lif.PROVINCE_CHOICES, lif.province_code) or '',
            '{{disability_status_code}}': lif.disability_status_code or '',
            '{{disability_status_name}}': get_choice_display(lif.DISABILITY_CHOICES, lif.disability_status_code) or '',
            '{{socio_economic_status_code}}': lif.socio_economic_status_code or '',
            '{{socio_economic_status_name}}': get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, lif.socio_economic_status_code) or '',
        }
        # Split fields
        if lif.learner_birth_date:
            birth_str = lif.learner_birth_date.strftime('%Y%m%d')
            for i in range(8):
                context[f'{{{{learner_birth_date_{i}}}}}'] = birth_str[i] if i < len(birth_str) else ''
        else:
            for i in range(8):
                context[f'{{{{learner_birth_date_{i}}}}}'] = ''
        if lif.national_id:
            for i in range(13):
                context[f'{{{{national_id_{i}}}}}'] = lif.national_id[i] if i < len(lif.national_id) else ''
        else:
            for i in range(13):
                context[f'{{{{national_id_{i}}}}}'] = ''
        if lif.duration_start_date:
            start_str = lif.duration_start_date.strftime('%d%m%Y')
            for i in range(8):
                context[f'{{{{commencement_date_{i}}}}}'] = start_str[i] if i < len(start_str) else ''
        else:
            for i in range(8):
                context[f'{{{{commencement_date_{i}}}}}'] = ''
        if lif.duration_end_date:
            end_str = lif.duration_end_date.strftime('%d%m%Y')
            for i in range(8):
                context[f'{{{{termination_date_{i}}}}}'] = end_str[i] if i < len(end_str) else ''
        else:
            for i in range(8):
                context[f'{{{{termination_date_{i}}}}}'] = ''
        # Table-based "checkbox" fields (X for selected)
        context.update({
            '{{gender_male}}': 'X' if lif.gender_code == 'M' else '',
            '{{gender_female}}': 'X' if lif.gender_code == 'F' else '',
            '{{below_35_yes}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year < 35) else '',
            '{{below_35_no}}': 'X' if lif.learner_birth_date and (date.today().year - lif.learner_birth_date.year >= 35) else '',
            '{{equity_african}}': 'X' if lif.equity_code == 'BA' else '',
            '{{equity_coloured}}': 'X' if lif.equity_code == 'BC' else '',
            '{{equity_indian}}': 'X' if lif.equity_code == 'BI' else '',
            '{{equity_white}}': 'X' if lif.equity_code == 'Wh' else '',
            '{{disability_sight}}': 'X' if lif.disability_status_code == '01' else '',
            '{{disability_hearing}}': 'X' if lif.disability_status_code == '02' else '',
            '{{disability_communication}}': 'X' if lif.disability_status_code == '03' else '',
            '{{disability_physical}}': 'X' if lif.disability_status_code == '04' else '',
            '{{disability_intellectual}}': 'X' if lif.disability_status_code == '05' else '',
            '{{disability_emotional}}': 'X' if lif.disability_status_code == '06' else '',
            '{{disability_multiple}}': 'X' if lif.disability_status_code == '07' else '',
            '{{disability_unspecified}}': 'X' if lif.disability_status_code == '09' else '',
            '{{disability_none}}': 'X' if lif.disability_status_code == 'N' else '',
            '{{citizen_sa}}': 'X' if lif.citizen_resident_status_code == 'SA' else '',
            '{{citizen_dual}}': 'X' if lif.citizen_resident_status_code == 'D' else '',
            '{{citizen_other}}': 'X' if lif.citizen_resident_status_code == 'O' else '',
            '{{citizen_permanent}}': 'X' if lif.citizen_resident_status_code == 'PR' else '',
            '{{citizen_unknown}}': 'X' if lif.citizen_resident_status_code == 'U' else '',
            '{{socio_employed}}': 'X' if lif.socio_economic_status_code == '01' else '',
            '{{socio_unemployed_seeking}}': 'X' if lif.socio_economic_status_code == '02' else '',
            '{{socio_not_working_not_looking}}': 'X' if lif.socio_economic_status_code == '03' else '',
            '{{socio_homemaker}}': 'X' if lif.socio_economic_status_code == '04' else '',
            '{{socio_student}}': 'X' if lif.socio_economic_status_code == '06' else '',
            '{{socio_pensioner}}': 'X' if lif.socio_economic_status_code == '07' else '',
            '{{socio_disabled}}': 'X' if lif.socio_economic_status_code == '08' else '',
            '{{socio_no_wish_to_work}}': 'X' if lif.socio_economic_status_code == '09' else '',
            '{{socio_not_working_nec}}': 'X' if lif.socio_economic_status_code == '10' else '',
            '{{socio_aged_under_15}}': 'X' if lif.socio_economic_status_code == '97' else '',
            '{{socio_institution}}': 'X' if lif.socio_economic_status_code == '98' else '',
            '{{socio_unspecified}}': 'X' if lif.socio_economic_status_code == 'U' else '',
            '{{tertiary_national_certificate}}': 'X' if lif.highest_tertiary_education == 'National Certificate' else '',
            '{{tertiary_national_diploma}}': 'X' if lif.highest_tertiary_education == 'National Diploma' else '',
            '{{tertiary_first_degree}}': 'X' if lif.highest_tertiary_education == 'National First Degree' else '',
            '{{tertiary_post_doctoral}}': 'X' if lif.highest_tertiary_education == 'Post-doctoral Degree' else '',
            '{{tertiary_doctoral}}': 'X' if lif.highest_tertiary_education == 'Doctoral Degree' else '',
            '{{tertiary_professional}}': 'X' if lif.highest_tertiary_education == 'Professional Qualification' else '',
            '{{tertiary_honours}}': 'X' if lif.highest_tertiary_education == 'Honours Degree' else '',
            '{{tertiary_higher_diploma}}': 'X' if lif.highest_tertiary_education == 'National Higher Diploma' else '',
            '{{tertiary_masters_diploma}}': 'X' if lif.highest_tertiary_education == 'National Masters Diploma' else '',
            '{{tertiary_national_higher}}': 'X' if lif.highest_tertiary_education == 'National Higher' else '',
            '{{tertiary_further_diploma}}': 'X' if lif.highest_tertiary_education == 'Further Diploma' else '',
            '{{tertiary_post_graduate}}': 'X' if lif.highest_tertiary_education == 'Post Graduate' else '',
            '{{tertiary_senior_certificate}}': 'X' if lif.highest_tertiary_education == 'Senior Certificate' else '',
            '{{tertiary_qual_nat_sen_cert}}': 'X' if lif.highest_tertiary_education == 'Qual at Nat Sen Cert' else '',
            '{{tertiary_apprenticeship}}': 'X' if lif.highest_tertiary_education == 'Apprenticeship' else '',
            '{{tertiary_post_grad_b_degree}}': 'X' if lif.highest_tertiary_education == 'Post Grad B Degree' else '',
            '{{tertiary_post_diploma_diploma}}': 'X' if lif.highest_tertiary_education == 'Post Diploma Diploma' else '',
            '{{tertiary_post_basic_diploma}}': 'X' if lif.highest_tertiary_education == 'Post-basic Diploma' else '',
            '{{province_western_cape}}': 'X' if lif.province_code == '1' else '',
            '{{province_eastern_cape}}': 'X' if lif.province_code == '2' else '',
            '{{province_northern_cape}}': 'X' if lif.province_code == '3' else '',
            '{{province_free_state}}': 'X' if lif.province_code == '4' else '',
            '{{province_kwazulu_natal}}': 'X' if lif.province_code == '5' else '',
            '{{province_north_west}}': 'X' if lif.province_code == '6' else '',
            '{{province_gauteng_jhb}}': 'X' if lif.province_code == '7' else '',
            '{{province_gauteng_pta}}': 'X' if lif.province_code == '7b' else '',
            '{{province_mpumalanga}}': 'X' if lif.province_code == '8' else '',
            '{{province_limpopo}}': 'X' if lif.province_code == '9' else '',
            '{{province_outside_sa}}': 'X' if lif.province_code == 'X' else '',
            '{{province_national}}': 'X' if lif.province_code == 'N' else '',
            '{{national_id_checkbox}}': 'X' if (lif.alternative_id_type == '' or lif.alternative_id_type is None) else '',
            '{{alt_id_saqa}}':      'X' if lif.alternative_id_type == '521' else '',
            '{{alt_id_passport}}':  'X' if lif.alternative_id_type == '527' else '',
            '{{alt_id_driver}}':    'X' if lif.alternative_id_type == '529' else '',
            '{{alt_id_temp_id}}':   'X' if lif.alternative_id_type == '531' else '',
            '{{alt_id_none}}':      'X' if lif.alternative_id_type == '533' else '',
            '{{alt_id_unknown}}':   'X' if lif.alternative_id_type == '535' else '',
            '{{alt_id_student}}':   'X' if lif.alternative_id_type == '537' else '',
            '{{alt_id_work_permit}}':'X' if lif.alternative_id_type == '538' else '',
            '{{alt_id_employee}}':  'X' if lif.alternative_id_type == '539' else '',
            '{{alt_id_birth_cert}}':'X' if lif.alternative_id_type == '540' else '',
            '{{alt_id_hsrc}}':      'X' if lif.alternative_id_type == '541' else '',
            '{{alt_id_etqa}}':      'X' if lif.alternative_id_type == '561' else '',
            '{{alt_id_refugee}}':   'X' if lif.alternative_id_type == '565' else '',
            '{{home_language_afrikaans}}': 'X' if lif.home_language_code == 'Afr' else '',
            '{{home_language_english}}': 'X' if lif.home_language_code == 'Eng' else '',
            '{{home_language_ndebele}}': 'X' if lif.home_language_code == 'Nde' else '',
            '{{home_language_sepedi}}': 'X' if lif.home_language_code == 'Sep' else '',
            '{{home_language_sesotho}}': 'X' if lif.home_language_code == 'Ses' else '',
            '{{home_language_setswana}}': 'X' if lif.home_language_code == 'Set' else '',
            '{{home_language_siswati}}': 'X' if lif.home_language_code == 'Swa' else '',
            '{{home_language_tshivenda}}': 'X' if lif.home_language_code == 'Tsh' else '',
            '{{home_language_isixhosa}}': 'X' if lif.home_language_code == 'Xho' else '',
            '{{home_language_xitsonga}}': 'X' if lif.home_language_code == 'Xit' else '',
            '{{home_language_isizulu}}': 'X' if lif.home_language_code == 'Zul' else '',
            '{{home_language_sasl}}': 'X' if lif.home_language_code == 'SASL' else '',
            '{{home_language_other}}': 'X' if lif.home_language_code == 'Oth' else '',
            '{{home_language_unknown}}': 'X' if lif.home_language_code == 'U' else '',
            '{{secondary_grade_8}}':  'X' if lif.highest_secondary_education == 'Grade 8' else '',
            '{{secondary_grade_9}}':  'X' if lif.highest_secondary_education == 'Grade 9' else '',
            '{{secondary_grade_10}}': 'X' if lif.highest_secondary_education == 'Grade 10' else '',
            '{{secondary_grade_11}}': 'X' if lif.highest_secondary_education == 'Grade 11' else '',
            '{{secondary_grade_12}}': 'X' if lif.highest_secondary_education == 'Grade 12' else '',
            '{{nationality_sa}}':   'X' if lif.nationality_code == 'SA' else '',
            '{{nationality_sdc}}':  'X' if lif.nationality_code == 'SDC' else '',
            '{{nationality_ang}}':  'X' if lif.nationality_code == 'ANG' else '',
            '{{nationality_bot}}':  'X' if lif.nationality_code == 'BOT' else '',
            '{{nationality_les}}':  'X' if lif.nationality_code == 'LES' else '',
            '{{nationality_mal}}':  'X' if lif.nationality_code == 'MAL' else '',
            '{{nationality_mau}}':  'X' if lif.nationality_code == 'MAU' else '',
            '{{nationality_moz}}':  'X' if lif.nationality_code == 'MOZ' else '',
            '{{nationality_nam}}':  'X' if lif.nationality_code == 'NAM' else '',
            '{{nationality_sey}}':  'X' if lif.nationality_code == 'SEY' else '',
            '{{nationality_swa}}':  'X' if lif.nationality_code == 'SWA' else '',
            '{{nationality_tan}}':  'X' if lif.nationality_code == 'TAN' else '',
            '{{nationality_zai}}':  'X' if lif.nationality_code == 'ZAI' else '',
            '{{nationality_zam}}':  'X' if lif.nationality_code == 'ZAM' else '',
            '{{nationality_zim}}':  'X' if lif.nationality_code == 'ZIM' else '',
            '{{nationality_ais}}':  'X' if lif.nationality_code == 'AIS' else '',
            '{{nationality_aus}}':  'X' if lif.nationality_code == 'AUS' else '',
            '{{nationality_eur}}':  'X' if lif.nationality_code == 'EUR' else '',
            '{{nationality_nor}}':  'X' if lif.nationality_code == 'NOR' else '',
            '{{nationality_sou}}':  'X' if lif.nationality_code == 'SOU' else '',
            '{{nationality_roa}}':  'X' if lif.nationality_code == 'ROA' else '',
            '{{nationality_ooc}}':  'X' if lif.nationality_code == 'OOC' else '',
            '{{nationality_u}}':    'X' if lif.nationality_code == 'U' else '',
            '{{nationality_not}}':  'X' if lif.nationality_code == 'NOT' else '',
        })
        for mapping in mappings:
            if mapping.placeholder in table_placeholders:
                continue
            value = getattr(lif, mapping.lif_field, '') or ''
            if mapping.lif_field == 'alternative_id_type':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.ALTERNATIVE_ID_TYPE_CHOICES, value)
            elif mapping.lif_field == 'equity_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.EQUITY_CHOICES, value)
            elif mapping.lif_field == 'nationality_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value)
            elif mapping.lif_field == 'gender_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = dict([('F', 'Female'), ('M', 'Male')]).get(value, '')
            elif mapping.lif_field == 'citizen_resident_status_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.CITIZEN_STATUS_CHOICES, value)
            elif mapping.lif_field == 'home_language_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.NATIONALITY_CHOICES, value)
            elif mapping.lif_field == 'province_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.PROVINCE_CHOICES, value)
            elif mapping.lif_field == 'disability_status_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.DISABILITY_CHOICES, value)
            elif mapping.lif_field == 'socio_economic_status_code':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.SOCIO_ECONOMIC_CHOICES, value)
            elif mapping.lif_field == 'highest_tertiary_education':
                context[mapping.placeholder] = value
                context[f'{mapping.placeholder}_name'] = get_choice_display(lif.TERTIARY_CHOICES, value)
            elif mapping.lif_field in ['duration_start_date', 'duration_end_date'] and value:
                context[mapping.placeholder] = value.strftime('%d/%m/%Y')
            elif mapping.lif_field == 'consent_to_process':
                context[mapping.placeholder] = ' Agree' if value else ' Disagree'
            elif mapping.lif_field in ['years_in_occupation', 'secondary_year_completed', 'tertiary_year_completed', 'percentage_maths', 'percentage_first_language', 'percentage_second_language']:
                context[mapping.placeholder] = str(value) if value else ''
            else:
                context[mapping.placeholder] = value
            if isinstance(value, str) and value and mapping.lif_field not in ['duration_start_date', 'duration_end_date'] and mapping.placeholder not in table_placeholders:
                base_match = re.match(r"\{\{(\w+)\}\}", mapping.placeholder)
                if base_match:
                    base = base_match.group(1)
                    for i, char in enumerate(value):
                        context[f"{{{{{base}_{i}}}}}"] = char
                    for i in range(len(value), 20):
                        context[f"{{{{{base}_{i}}}}}"] = ''
        return context

    def replace_placeholders_in_paragraph(paragraph, context):
        full_text = ''.join(run.text for run in paragraph.runs)
        replaced = False
        for placeholder, value in context.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                replaced = True
        if replaced and paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ''

    def replace_placeholders_in_table(table, context):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, context)
                for nested_table in cell.tables:
                    replace_placeholders_in_table(nested_table, context)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for learner_id in learner_ids:
            try:
                learner = get_object_or_404(Learner, id=learner_id)
                lif = LIF.objects.get(learner=learner)
                lq = LearnerQualification.objects.filter(
                    learner=learner,
                    sla_qualification__groups=group
                ).select_related('sla_qualification__service').first()

                if not lq or not lq.sla_qualification.service.admin_pack_document:
                    logger.warning(f"Skipping learner {learner.id}: No valid learner qualification or admin pack document")
                    continue

                admin_pack_template = lq.sla_qualification.service.admin_pack_document
                doc = Document(admin_pack_template.path)

                # Insert LIF template content after the first section break (cover page)
                lif_doc = Document(lif_template.template_file.path)
                insert_lif_content_after_first_section(doc, lif_doc)

                # Replace placeholders in the combined document
                context = build_context(lif)
                for paragraph in doc.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, context)
                for table in doc.tables:
                    replace_placeholders_in_table(table, context)

                learner_name = f"{learner.FirstName}_{learner.Surname}".replace(" ", "_").replace("/", "_").replace("\\", "_")
                file_buffer = io.BytesIO()
                doc.save(file_buffer)
                file_buffer.seek(0)
                zip_file.writestr(f"{learner_name}_AdminPack.docx", file_buffer.read())
                logger.info(f"Successfully generated document for learner {learner_id}")

            except Exception as e:
                logger.error(f"Error processing learner {learner_id}: {str(e)}")
                continue

    zip_buffer.seek(0)
    if zip_buffer.getvalue():
        safe_template = "".join(x for x in lif_template.name if x.isalnum() or x in (' ', '_', '-')).rstrip()
        response = HttpResponse(zip_buffer, content_type="application/zip")
        response["Content-Disposition"] = f'attachment; filename="AdminPack_{safe_template}_Group_{group.id}.zip"'
        return response
    else:
        logger.error("No documents were generated for the ZIP file")
        return HttpResponse("No valid documents generated", status=400)

from django.views.generic.edit import DeleteView
from django.urls import reverse_lazy
from .models import LIF

class LIFDeleteView(DeleteView):
    model = LIF
    template_name = 'core/lif_confirm_delete.html'
    success_url = reverse_lazy('generate_lif_word')  # or your LIF list view name

from django.views.generic.edit import UpdateView
from django.urls import reverse_lazy
from .models import LIF
from .forms import LIFForm  # You may need to create this ModelForm if not present
from django.utils.safestring import mark_safe

class LIFUpdateView(UpdateView):
    model = LIF
    form_class = LIFForm
    template_name = 'core/lif_update.html'  # You need to create this template

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # If not editing, make the form read-only
        context['is_edit'] = self.request.GET.get('edit') == '1'
        return context

    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        return form

    def get_success_url(self):
        # Redirect back to the LIF list or detail page after saving
        return reverse_lazy('generate_lif_word')



class VenueBookingSwitchView(RolePermissionRequiredMixin, View):
    """
    AJAX view to handle venue switching for existing bookings
    """
    
    def post(self, request):
        try:
            import json
            data = json.loads(request.body)
            booking_id = data.get('booking_id')
            new_venue_id = data.get('new_venue_id')
            
            if not booking_id or not new_venue_id:
                return JsonResponse({'success': False, 'message': 'Missing booking ID or venue ID'})
            
            booking = get_object_or_404(VenueBooking, pk=booking_id)
            
            # Handle virtual session group
            if new_venue_id == "virtual-session-group":
                # Find an available virtual session venue
                virtual_venues = Venue.objects.filter(name__istartswith="Virtual Session").order_by('name')
                new_venue = None
                
                for venue in virtual_venues:
                    # Check if this venue is available at the same time
                    conflict = VenueBooking.objects.filter(
                        venue=venue,
                        start_datetime=booking.start_datetime,
                        end_datetime=booking.end_datetime,
                        status__in=['booked', 'rescheduled']
                    ).exclude(pk=booking.id).exists()
                    
                    if not conflict:
                        new_venue = venue
                        break
                
                if not new_venue:
                    # Create a new virtual session venue
                    existing_names = [v.name for v in virtual_venues]
                    idx = 1
                    while f"Virtual Session {idx}" in existing_names:
                        idx += 1
                    new_venue = Venue.objects.create(name=f"Virtual Session {idx}")
            else:
                new_venue = get_object_or_404(Venue, pk=new_venue_id)
                
                # Check for conflicts
                conflict = VenueBooking.objects.filter(
                    venue=new_venue,
                    start_datetime=booking.start_datetime,
                    end_datetime=booking.end_datetime,
                    status__in=['booked', 'rescheduled']
                ).exclude(pk=booking.id).exists()
                
                if conflict:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Venue {new_venue.name} is not available at this time'
                    })
            
            # Store old venue for notification
            old_venue = booking.venue
            
            # Update the booking
            booking.venue = new_venue
            booking.status = 'rescheduled'
            booking.save()
            
            # If this is a combined booking, update all related bookings
            if hasattr(booking, 'combined_booking_id') and booking.combined_booking_id:
                VenueBooking.objects.filter(
                    combined_booking_id=booking.combined_booking_id
                ).exclude(pk=booking.id).update(
                    venue=new_venue,
                    status='rescheduled'
                )
            
            # Send notification email
            try:
                recipient_email = "brendonmandlandlovu@gmail.com"
                if booking.user and booking.user.email:
                    recipient_email = booking.user.email
                
                subject = f"Venue Changed - {booking.session_date if booking.session_date else 'Booking'}"
                
                message = f"""
Dear {booking.user.get_full_name() if booking.user else 'User'},

Your venue booking has been switched to a new venue.

Booking Details:
- Session: {booking.session_date if booking.session_date else 'N/A'}
- Date: {booking.start_datetime.strftime('%B %d, %Y')}
- Time: {booking.start_datetime.strftime('%I:%M %p')} - {booking.end_datetime.strftime('%I:%M %p')}
- Purpose: {booking.booking_purpose}

Venue Change:
- Previous Venue: {old_venue.name}
- New Venue: {new_venue.name}

Changed by: {request.user.get_full_name() if request.user.get_full_name() else request.user.username}
Changed on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}

Please make note of the new venue location.

Best regards,
The Learning Organisation
                """
                
                send_mail(
                    subject=subject,
                    message=message,
                    from_email="noreply@ensemble.com",
                    recipient_list=[recipient_email],
                    fail_silently=True,
                )
            except Exception as e:
                pass  # Email sending is not critical
            
            return JsonResponse({
                'success': True,
                'message': f'Venue switched from {old_venue.name} to {new_venue.name}'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error switching venue: {str(e)}'})



@login_required
def all_sessions_api(request):
    """API endpoint to get all sessions for search functionality."""
    from django.http import JsonResponse
    
    sessions = SessionDate.objects.select_related('project_plan__group', 'project_plan__module').all()
    
    sessions_data = []
    for session in sessions:
        group_name = ""
        module_name = ""
        if session.project_plan:
            if session.project_plan.group:
                group_name = str(session.project_plan.group)
            if session.project_plan.module:
                module_name = str(session.project_plan.module)
        
        session_text = f"{group_name} - {module_name} ({session.start_date} to {session.end_date})"
        
        sessions_data.append({
            'id': session.id,
            'text': session_text,
            'start_date': session.start_date.strftime('%Y-%m-%d'),
            'end_date': session.end_date.strftime('%Y-%m-%d'),
            'group': group_name,
            'module': module_name
        })
    
    return JsonResponse({'sessions': sessions_data})